from __future__ import annotations

from io import BytesIO
import json
import logging
from pathlib import Path
import sys
import tempfile
import unittest
from unittest.mock import patch
import zipfile

THIS_DIR = Path(__file__).resolve().parent
SRC_ROOT = THIS_DIR.parent / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from document_processor import (
    BoundingBox,
    CellStyleInfo,
    DocIR,
    DocumentInput,
    HwpxDocument,
    ImageIR,
    NativeAnchor,
    PageInfo,
    ParaStyleInfo,
    ParagraphIR,
    RunIR,
    RunStyleInfo,
    StyleMap,
    TableCellIR,
    TableIR,
    TableStyleInfo,
    build_doc_ir_from_mapping,
    configure_logging,
    get_logger,
    read_document,
)
from document_processor.builder import apply_style_map_to_doc_ir
from document_processor.core.hwpx_structured_exporter import export_hwpx_structured_mapping
import document_processor.models as document_models


class DocumentIRTests(unittest.TestCase):
    def _sample_mapping(self) -> dict[str, str]:
        return {
            "s1.p1.r1": "Hello ",
            "s1.p1.r2": "World",
            "s1.p2.r1.tbl1.tr1.tc1.p1.r1": "A1",
            "s1.p2.r1.tbl1.tr1.tc2.p1.r1": "B1",
            "s1.p2.r1.tbl1.tr2.tc1.p1.r1": "A2",
            "s1.p2.r1.tbl1.tr2.tc2.p1.r1": "B2",
        }

    def _sample_style_map(self) -> StyleMap:
        return StyleMap(runs={
                "s1.p1.r1": RunStyleInfo(bold=True, size_pt=11.0),
                "s1.p1.r2": RunStyleInfo(italic=True, size_pt=11.0),
            },
            paragraphs={
                "s1.p1": ParaStyleInfo(align="center"),
            },
            cells={
                "s1.p2.r1.tbl1.tr1.tc1": CellStyleInfo(background="#ffeeaa"),
            },
            tables={
                "s1.p2.r1.tbl1": TableStyleInfo(row_count=2, col_count=2),
            },
        )

    def test_hierarchy_construction(self) -> None:
        doc_ir = build_doc_ir_from_mapping(self._sample_mapping())

        self.assertEqual(len(doc_ir.paragraphs), 2)
        self.assertEqual(doc_ir.paragraphs[0].text, "Hello World")
        self.assertEqual(doc_ir.paragraphs[1].content[0].native_anchor.debug_path, "s1.p2.r1.tbl1")
        self.assertEqual(doc_ir.paragraphs[1].tables[0].row_count, 2)
        self.assertEqual(doc_ir.paragraphs[1].tables[0].col_count, 2)

    def test_style_embedding(self) -> None:
        doc_ir = build_doc_ir_from_mapping(self._sample_mapping(), style_map=self._sample_style_map())
        self.assertEqual(doc_ir.paragraphs[0].para_style.align, "center")
        self.assertTrue(doc_ir.paragraphs[0].runs[0].run_style.bold)
        self.assertEqual(doc_ir.paragraphs[1].tables[0].cells[0][0].cell_style.background, "#ffeeaa")

    def test_docir_subclass_from_mapping(self) -> None:
        class DocumentLM(DocIR):
            custom_field: int = 0

        doc = DocumentLM.from_mapping({"s1.p1.r1": "X"}, custom_field=7)
        self.assertIsInstance(doc, DocumentLM)
        self.assertEqual(doc.custom_field, 7)

    def test_docir_logging_defaults_and_file_output(self) -> None:
        logger = configure_logging()
        self.assertEqual(logger.level, logging.WARNING)
        self.assertEqual(get_logger().name, "document_processor")
        self.assertEqual(get_logger("helpers").name, "document_processor.helpers")
        self.assertTrue(any(isinstance(handler, logging.StreamHandler) for handler in logger.handlers))

        with tempfile.TemporaryDirectory() as tmp_dir:
            log_path = Path(tmp_dir) / "docir.log"
            try:
                configured_logger = DocIR.configure_logging(
                    level="INFO",
                    log_file=log_path,
                    console=False,
                )
                DocIR.from_mapping({"s1.p1.r1": "Logged"})
                for handler in configured_logger.handlers:
                    handler.flush()

                log_text = log_path.read_text(encoding="utf-8")
                self.assertIn("Building DocIR from mapping with 1 run(s)", log_text)
                self.assertIn("Built DocIR from mapping", log_text)
            finally:
                configure_logging()

    def test_content_is_source_of_truth(self) -> None:
        doc = DocIR(paragraphs=[
                ParagraphIR(content=[
                        RunIR(text="Hello"),
                        ImageIR(image_id="img1"),
                        TableIR(),
                    ],
                )
            ]
        )

        paragraph = doc.paragraphs[0]
        self.assertEqual(
            [type(node).__name__ for node in paragraph.content],
            ["RunIR", "ImageIR", "TableIR"],
        )
        self.assertEqual([run.text for run in paragraph.runs], ["Hello"])
        self.assertEqual([image.image_id for image in paragraph.images], ["img1"])
        self.assertEqual([table.native_anchor.debug_path for table in paragraph.tables], ["s1.p1.r1.tbl1"])

        content_annotation = ParagraphIR.model_fields["content"].annotation
        self.assertIn("RunIR", str(content_annotation))
        self.assertIn("ImageIR", str(content_annotation))
        self.assertIn("TableIR", str(content_annotation))

    def test_semantic_models_are_not_public_exports(self) -> None:
        import document_processor

        self.assertFalse(hasattr(document_processor, "SemanticBlock"))
        self.assertFalse(hasattr(document_processor, "SemanticDocument"))
        self.assertFalse(hasattr(document_processor, "SemanticBlockIR"))
        self.assertFalse(hasattr(document_processor, "SemanticIR"))
        self.assertNotIn("SemanticBlock", document_processor.__all__)
        self.assertNotIn("SemanticDocument", document_processor.__all__)
        self.assertNotIn("SemanticBlockIR", document_processor.__all__)
        self.assertNotIn("SemanticIR", document_processor.__all__)
        self.assertNotIn("SemanticBlock", document_models.__all__)
        self.assertNotIn("SemanticDocument", document_models.__all__)
        self.assertNotIn("SemanticBlockIR", document_models.__all__)
        self.assertNotIn("SemanticIR", document_models.__all__)

    def test_to_semantic_returns_model_with_markdown_tables(self) -> None:
        doc_ir = build_doc_ir_from_mapping(
            self._sample_mapping(),
            source_path="sample.pdf",
            source_doc_type="pdf",
            doc_id="doc_1",
        )
        doc_ir.paragraphs[0].page_number = 1
        doc_ir.paragraphs[0].bbox = BoundingBox(left_pt=10, bottom_pt=20, right_pt=110, top_pt=40)
        doc_ir.paragraphs[1].page_number = 2
        table = doc_ir.paragraphs[1].tables[0]
        table.bbox = BoundingBox(left_pt=12, bottom_pt=50, right_pt=210, top_pt=120)
        table.previous_table_id = "tbl_previous"
        table.next_table_id = "tbl_next"
        doc_ir.paragraphs.append(
            ParagraphIR(
                page_number=3,
                content=[
                    ImageIR(
                        image_id="img_1",
                        alt_text="Sample image",
                        bbox=BoundingBox(left_pt=20, bottom_pt=30, right_pt=80, top_pt=90),
                    )
                ],
            )
        )
        doc_ir.ensure_node_identity()

        semantic = doc_ir.to_semantic()

        self.assertEqual(type(semantic).__name__, "SemanticIR")
        self.assertEqual(type(semantic.blocks[0]).__name__, "SemanticBlockIR")
        self.assertEqual(semantic.doc_id, "doc_1")
        self.assertEqual(semantic.source_path, "sample.pdf")
        self.assertEqual(semantic.source_doc_type, "pdf")
        self.assertEqual([block.kind for block in semantic.blocks], ["paragraph", "table", "image"])
        self.assertEqual(semantic.blocks[0].text, "Hello World")
        self.assertEqual(semantic.blocks[0].page_number, 1)
        self.assertEqual(semantic.blocks[0].bbox.left_pt, 10)
        self.assertEqual(semantic.blocks[1].node_id, table.node_id)
        self.assertEqual(semantic.blocks[1].debug_path, table.native_anchor.debug_path)
        self.assertEqual(semantic.blocks[1].bbox.right_pt, 210)
        self.assertIn("| col1 | col2 |", semantic.blocks[1].text)
        self.assertIn("| A1 | B1 |", semantic.blocks[1].text)
        self.assertEqual(semantic.blocks[1].previous_table_id, "tbl_previous")
        self.assertEqual(semantic.blocks[1].next_table_id, "tbl_next")
        self.assertEqual(semantic.blocks[2].text, "Sample image")
        self.assertEqual(semantic.blocks[2].bbox.top_pt, 90)

        semantic_dict = semantic.model_dump(mode="json", exclude_none=True)
        self.assertIsInstance(semantic_dict, dict)
        self.assertEqual(semantic_dict["blocks"][0]["bbox"]["bottom_pt"], 20)
        self.assertNotIn("previous_table_id", semantic_dict["blocks"][0])
        self.assertEqual(semantic_dict["blocks"][1]["next_table_id"], "tbl_next")

        semantic_json = semantic.model_dump_json(exclude_none=True, indent=2)
        self.assertEqual(json.loads(semantic_json)["blocks"][1]["previous_table_id"], "tbl_previous")

        with self.assertRaises(TypeError):
            doc_ir.to_semantic(format="dict")

    def test_from_file_docx_path_and_file_object(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "sample.docx"

            doc = Document()
            doc.add_paragraph("Hello")
            doc.save(str(docx_path))

            from_path = DocIR.from_file(docx_path)
            with docx_path.open("rb") as handle:
                from_file_object = DocIR.from_file(handle)

        self.assertEqual(from_path.source_doc_type, "docx")
        self.assertEqual(from_path.source_path, str(docx_path))
        self.assertEqual(from_path.paragraphs[0].text, "Hello")
        self.assertEqual(from_file_object.source_doc_type, "docx")
        self.assertEqual(from_file_object.paragraphs[0].text, "Hello")

    def test_from_file_hwpx_bytes_and_file_object(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p><hp:run><hp:t>Hello HWPX</hp:t></hp:run></hp:p>
</hs:sec>
""",
            )
        hwpx_bytes = hwpx_bytes_io.getvalue()

        from_bytes = DocIR.from_file(hwpx_bytes, doc_type="hwpx")
        from_file_object = DocIR.from_file(BytesIO(hwpx_bytes), doc_type="hwpx")

        self.assertEqual(from_bytes.source_doc_type, "hwpx")
        self.assertEqual(from_bytes.paragraphs[0].text, "Hello HWPX")
        self.assertEqual(from_file_object.paragraphs[0].text, "Hello HWPX")

    def test_export_hwpx_structured_mapping_accepts_hwpx_document(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p><hp:run><hp:t>Hello from doc object</hp:t></hp:run></hp:p>
</hs:sec>
""",
            )

        with HwpxDocument.open(hwpx_bytes_io.getvalue()) as doc:
            mapping = export_hwpx_structured_mapping(doc)

        self.assertEqual(mapping, {"s1.p1.r1": "Hello from doc object"})

    def test_from_file_hwpx_reads_mixed_content_inside_hp_t(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:t>&lt;<hp:fwSpace />수요기업 협업 규모 및 분야<hp:fwSpace />&gt;</hp:t>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )

        doc = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        self.assertEqual(doc.paragraphs[0].text, "<수요기업 협업 규모 및 분야>")

    def test_from_file_hwpx_extracts_page_info_and_assigns_pages(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p pageBreak="0">
    <hp:run>
      <hp:secPr>
        <hp:pagePr width="59528" height="84188">
          <hp:margin left="5669" right="5669" top="2834" bottom="2834" />
        </hp:pagePr>
      </hp:secPr>
      <hp:t>Page 1-A</hp:t>
    </hp:run>
    <hp:linesegarray><hp:lineseg vertpos="0" /></hp:linesegarray>
  </hp:p>
  <hp:p pageBreak="0">
    <hp:run><hp:t>Page 1-B</hp:t></hp:run>
    <hp:linesegarray><hp:lineseg vertpos="2200" /></hp:linesegarray>
  </hp:p>
  <hp:p pageBreak="0">
    <hp:run><hp:t>Page 2-A</hp:t></hp:run>
    <hp:linesegarray><hp:lineseg vertpos="0" /></hp:linesegarray>
  </hp:p>
</hs:sec>
""",
            )

        doc = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        self.assertEqual([page.page_number for page in doc.pages], [1, 2])
        self.assertAlmostEqual(doc.pages[0].width_pt or 0.0, 595.28, places=2)
        self.assertAlmostEqual(doc.pages[0].height_pt or 0.0, 841.88, places=2)
        self.assertAlmostEqual(doc.pages[0].margin_left_pt or 0.0, 56.69, places=2)
        self.assertEqual([paragraph.page_number for paragraph in doc.paragraphs], [1, 1, 2])

    def test_from_file_docx_extracts_page_info_from_page_breaks(self) -> None:
        from docx import Document
        from docx.shared import Inches

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "pages.docx"
            doc = Document()
            section = doc.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            doc.add_paragraph("Page 1")
            doc.add_page_break()
            doc.add_paragraph("Page 2")
            doc.save(str(docx_path))

            parsed = DocIR.from_file(docx_path, skip_empty=True)

        self.assertEqual([page.page_number for page in parsed.pages], [1, 2])
        self.assertAlmostEqual(parsed.pages[0].width_pt or 0.0, 612.0, places=1)
        self.assertAlmostEqual(parsed.pages[0].height_pt or 0.0, 792.0, places=1)
        self.assertAlmostEqual(parsed.pages[0].margin_left_pt or 0.0, 72.0, places=1)
        self.assertEqual([paragraph.text for paragraph in parsed.paragraphs], ["Page 1", "Page 2"])
        self.assertEqual([paragraph.page_number for paragraph in parsed.paragraphs], [1, 2])

    def test_from_file_docx_extracts_section_column_layouts(self) -> None:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def set_section_columns(section, *, count: int, space_twips: int) -> None:
            cols = section._sectPr.find(qn("w:cols"))
            if cols is None:
                cols = OxmlElement("w:cols")
                section._sectPr.append(cols)
            cols.set(qn("w:num"), str(count))
            cols.set(qn("w:space"), str(space_twips))

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "columns.docx"
            doc = Document()
            doc.add_paragraph("One column title")
            three_col_section = doc.add_section(WD_SECTION.CONTINUOUS)
            set_section_columns(three_col_section, count=3, space_twips=720)
            doc.add_paragraph("Three column body")
            doc.save(str(docx_path))

            parsed = DocIR.from_file(docx_path, skip_empty=True)

        self.assertEqual([paragraph.text for paragraph in parsed.paragraphs], ["One column title", "Three column body"])
        self.assertIsNone(parsed.paragraphs[0].para_style)
        self.assertEqual(parsed.paragraphs[1].para_style.column_layout.count, 3)
        self.assertAlmostEqual(parsed.paragraphs[1].para_style.column_layout.gap_pt or 0.0, 36.0, places=2)

    def test_from_file_hwpx_extracts_paragraph_column_layouts(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:ctrl><hp:colPr colCount="1" sameSz="1" sameGap="0" /></hp:ctrl>
      <hp:t>One column title</hp:t>
    </hp:run>
  </hp:p>
  <hp:p>
    <hp:run>
      <hp:ctrl><hp:colPr colCount="3" sameSz="1" sameGap="300" /></hp:ctrl>
      <hp:t>Three column body start</hp:t>
    </hp:run>
  </hp:p>
  <hp:p>
    <hp:run><hp:t>Three column body continued</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )

        parsed = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        self.assertEqual([paragraph.text for paragraph in parsed.paragraphs], [
            "One column title",
            "Three column body start",
            "Three column body continued",
        ])
        self.assertIsNone(parsed.paragraphs[0].para_style)
        self.assertEqual([paragraph.para_style.column_layout.count for paragraph in parsed.paragraphs[1:]], [3, 3])
        self.assertAlmostEqual(parsed.paragraphs[1].para_style.column_layout.gap_pt or 0.0, 3.0, places=2)

    def test_from_file_docx_extracts_list_markers_from_numbering(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "lists.docx"
            doc = Document()
            doc.add_paragraph("First numbered item", style="List Number")
            doc.add_paragraph("Second numbered item", style="List Number")
            doc.add_paragraph("Bullet item", style="List Bullet")
            doc.save(str(docx_path))

            parsed = DocIR.from_file(docx_path, skip_empty=True)

        list_infos = [paragraph.para_style.list_info for paragraph in parsed.paragraphs]
        self.assertEqual([info.marker for info in list_infos], ["1.", "2.", "•"])
        self.assertEqual([info.level for info in list_infos], [0, 0, 0])
        self.assertEqual([info.marker_type for info in list_infos], ["decimal", "decimal", "bullet"])

        read_result = read_document(document=DocumentInput(doc_ir=parsed))
        self.assertEqual(
            [paragraph.display_text for paragraph in read_result.paragraphs],
            ["1. First numbered item", "2. Second numbered item", "• Bullet item"],
        )

        html = parsed.to_html()
        self.assertIn('class="document-list-marker"', html)
        self.assertIn(">1.</span>First numbered item", html)
        self.assertIn(">2.</span>Second numbered item", html)
        self.assertIn(">•</span>Bullet item", html)

    def test_from_file_docx_without_numbering_part_does_not_raise(self) -> None:
        docx_path = THIS_DIR / "doc_samples" / "new_test" / "element_side_by_side_test.docx"

        parsed = DocIR.from_file(docx_path, skip_empty=True)

        self.assertGreater(len(parsed.paragraphs), 0)

    def test_from_file_hwpx_extracts_number_and_bullet_headings(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:refList>
    <hh:numberings itemCnt="1">
      <hh:numbering id="1" start="1">
        <hh:paraHead start="1" level="1" numFormat="DIGIT">^1.</hh:paraHead>
      </hh:numbering>
    </hh:numberings>
    <hh:bullets itemCnt="1">
      <hh:bullet id="2" char="•" useImg="0">
        <hh:paraHead level="0" numFormat="DIGIT" />
      </hh:bullet>
    </hh:bullets>
    <hh:paraProperties itemCnt="2">
      <hh:paraPr id="1">
        <hh:heading type="NUMBER" idRef="1" level="0" />
      </hh:paraPr>
      <hh:paraPr id="2">
        <hh:heading type="BULLET" idRef="2" level="0" />
      </hh:paraPr>
    </hh:paraProperties>
  </hh:refList>
</hh:head>
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p paraPrIDRef="1"><hp:run><hp:t>First numbered item</hp:t></hp:run></hp:p>
  <hp:p paraPrIDRef="1"><hp:run><hp:t>Second numbered item</hp:t></hp:run></hp:p>
  <hp:p paraPrIDRef="2"><hp:run><hp:t>Bullet item</hp:t></hp:run></hp:p>
</hs:sec>
""",
            )

        parsed = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        list_infos = [paragraph.para_style.list_info for paragraph in parsed.paragraphs]
        self.assertEqual([info.marker for info in list_infos], ["1.", "2.", "•"])
        self.assertEqual([info.marker_type for info in list_infos], ["decimal", "decimal", "bullet"])

    def test_from_file_hwp_file_object_materializes_temp_path(self) -> None:
        fake_hwp = b"fake-hwp"

        with (
            patch("document_processor.core.document_ir_parser.convert_hwp_to_hwpx_bytes") as convert_hwp,
            patch("document_processor.core.style_extractor.extract_styles") as extract_styles,
        ):
            hwpx_bytes_io = BytesIO()
            with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
                zf.writestr(
                    "Contents/header.xml",
                    """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
                )
                zf.writestr(
                    "Contents/section0.xml",
                    """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p><hp:run><hp:t>Converted</hp:t></hp:run></hp:p>
</hs:sec>
""",
                )
            convert_hwp.return_value = hwpx_bytes_io.getvalue()
            extract_styles.return_value = StyleMap()

            doc = DocIR.from_file(BytesIO(fake_hwp), doc_type="hwp")

        self.assertEqual(doc.source_doc_type, "hwp")
        self.assertEqual(doc.paragraphs[0].text, "Converted")
        convert_source = convert_hwp.call_args.kwargs.get("hwp_path", convert_hwp.call_args.args[0])
        self.assertTrue(isinstance(convert_source, Path))
        self.assertEqual(convert_source.suffix, ".hwp")

    def test_hwpx_vertical_merge_uses_logical_column_ids(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>Main</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="0" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="4"/></hp:tc>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>관</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="1" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>A</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="2" rowAddr="0"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
        </hp:tr>
        <hp:tr>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>항</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="1" rowAddr="1"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
          <hp:tc><hp:subList><hp:p><hp:run><hp:t>B</hp:t></hp:run></hp:p></hp:subList><hp:cellAddr colAddr="2" rowAddr="1"/><hp:cellSpan colSpan="1" rowSpan="1"/></hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        mapping = export_hwpx_structured_mapping(hwpx_bytes_io.getvalue())

        self.assertIn("s1.p1.r1.tbl1.tr1.tc2.p1.r1", mapping)
        self.assertIn("s1.p1.r1.tbl1.tr2.tc2.p1.r1", mapping)
        self.assertIn("s1.p1.r1.tbl1.tr2.tc3.p1.r1", mapping)
        self.assertNotIn("s1.p1.r1.tbl1.tr2.tc1.p1.r1", mapping)

    def test_builder_supports_nested_tables_in_cell_paragraphs(self) -> None:
        mapping = {
            "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Outer",
            "s1.p1.r1.tbl1.tr1.tc1.p1.tbl1.tr1.tc1.p1.r1": "Inner",
        }

        doc = DocIR.from_mapping(mapping)
        outer_cell_paragraph = doc.paragraphs[0].tables[0].cells[0][0].paragraphs[0]

        self.assertEqual(outer_cell_paragraph.text, "Outer\nInner")
        self.assertEqual(len(outer_cell_paragraph.tables), 1)
        self.assertEqual(
            outer_cell_paragraph.tables[0].native_anchor.debug_path,
            "s1.p1.r1.tbl1.tr1.tc1.p1.tbl1",
        )
        self.assertEqual(
            outer_cell_paragraph.tables[0].cells[0][0].paragraphs[0].runs[0].text,
            "Inner",
        )

    def test_table_markdown_repeats_merged_cells(self) -> None:
        table = TableIR(
            cells=[
                [
                    TableCellIR(
                        cell_style=CellStyleInfo(rowspan=2, colspan=2),
                        paragraphs=[ParagraphIR(content=[RunIR(text="Merged")])],
                    ),
                    TableCellIR(paragraphs=[ParagraphIR(content=[RunIR(text="Right")])]),
                ],
                [TableCellIR(paragraphs=[ParagraphIR(content=[RunIR(text="Bottom")])])],
            ],
        )

        markdown = table.markdown

        self.assertFalse(hasattr(table.cells[0][0], "row_index"))
        self.assertEqual([len(row) for row in table.cells], [3, 3])
        self.assertIs(table.cells[0][1], table.cells[0][0])
        self.assertIs(table.cells[1][0], table.cells[0][0])
        self.assertIs(table.cells[1][1], table.cells[0][0])
        self.assertEqual(table.cells[0][2].paragraphs[0].content[0].text, "Right")
        self.assertEqual(table.cells[1][2].paragraphs[0].content[0].text, "Bottom")
        self.assertEqual(
            [
                (row_index, col_index, cell.paragraphs[0].content[0].text)
                for row_index, col_index, cell in table.iter_cell_positions()
            ],
            [(1, 1, "Merged"), (1, 3, "Right"), (2, 3, "Bottom")],
        )
        round_tripped = TableIR.model_validate(table.model_dump(mode="python"))
        self.assertEqual([len(row) for row in round_tripped.cells], [3, 3])
        self.assertIs(round_tripped.cells[0][1], round_tripped.cells[0][0])
        self.assertIs(round_tripped.cells[1][0], round_tripped.cells[0][0])
        self.assertIs(round_tripped.cells[1][1], round_tripped.cells[0][0])
        self.assertEqual(
            [
                (row_index, col_index, cell.paragraphs[0].content[0].text)
                for row_index, col_index, cell in round_tripped.iter_cell_positions()
            ],
            [(1, 1, "Merged"), (1, 3, "Right"), (2, 3, "Bottom")],
        )
        self.assertIn("| col1 | col2 | col3 |", markdown)
        self.assertIn("| Merged | Merged | Right |", markdown)
        self.assertIn("| Merged | Merged | Bottom |", markdown)

    def test_apply_style_map_expands_native_covered_slot_duplicates(self) -> None:
        def paragraph(text: str) -> ParagraphIR:
            paragraph_ir = ParagraphIR(content=[RunIR(text=text)])
            paragraph_ir.recompute_text()
            return paragraph_ir

        def cell(path: str, text: str) -> TableCellIR:
            table_cell = TableCellIR(
                native_anchor=NativeAnchor(
                    node_kind="cell",
                    debug_path=path,
                    structural_path=path,
                ),
                paragraphs=[paragraph(text)],
            )
            table_cell.recompute_text()
            return table_cell

        table = TableIR(
            row_count=2,
            col_count=2,
            native_anchor=NativeAnchor(
                node_kind="table",
                debug_path="s1.p1.r1.tbl1",
                structural_path="s1.p1.r1.tbl1",
            ),
            cells=[
                [cell("s1.p1.r1.tbl1.tr1.tc1", "Merged"), cell("s1.p1.r1.tbl1.tr1.tc2", "Right")],
                [cell("s1.p1.r1.tbl1.tr2.tc1", "Merged"), cell("s1.p1.r1.tbl1.tr2.tc2", "Bottom")],
            ],
        )
        doc = DocIR(paragraphs=[ParagraphIR(content=[table])])
        style_map = StyleMap(
            tables={"s1.p1.r1.tbl1": TableStyleInfo(row_count=2, col_count=2)},
            cells={"s1.p1.r1.tbl1.tr1.tc1": CellStyleInfo(rowspan=2)},
        )

        apply_style_map_to_doc_ir(doc, style_map)

        self.assertIs(table.cells[1][0], table.cells[0][0])
        self.assertEqual(
            [
                (row_index, col_index, cell.paragraphs[0].content[0].text)
                for row_index, col_index, cell in table.iter_cell_positions()
            ],
            [(1, 1, "Merged"), (1, 2, "Right"), (2, 2, "Bottom")],
        )
        self.assertIn("| Merged | Right |", table.markdown)
        self.assertIn("| Merged | Bottom |", table.markdown)

    def test_table_markdown_appends_nested_tables_by_reference(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Outer",
                "s1.p1.r1.tbl1.tr1.tc2.p1.tbl1.tr1.tc1.p1.r1": "Inner",
            }
        )
        outer = doc.paragraphs[0].tables[0]
        nested_path = outer.cells[0][1].paragraphs[0].tables[0].native_anchor.debug_path

        markdown = outer.markdown

        self.assertIn(f"| Outer | [tbl:{nested_path}] |", markdown)
        self.assertIn(f"[tbl:{nested_path}]", markdown)
        self.assertIn("| col1 |", markdown)
        self.assertIn("| Inner |", markdown)

    def test_docx_nested_tables_are_parsed(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "nested.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "Outer"
            nested = cell.add_table(rows=1, cols=1)
            nested.cell(0, 0).text = "Inner"
            doc.save(str(docx_path))

            parsed = DocIR.from_file(docx_path)

        outer_cell_paragraph = parsed.paragraphs[0].tables[0].cells[0][0].paragraphs[0]
        self.assertEqual(outer_cell_paragraph.runs[0].text, "Outer")
        self.assertEqual(len(outer_cell_paragraph.tables), 1)
        self.assertEqual(
            outer_cell_paragraph.tables[0].cells[0][0].paragraphs[0].runs[0].text,
            "Inner",
        )

    def test_hwpx_tables_inside_one_run_preserve_child_order(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:t>Before</hp:t>
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p><hp:run><hp:t>Cell</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
      <hp:t>After</hp:t>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )

        parsed = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")
        paragraph_ir = parsed.paragraphs[0]

        self.assertEqual(
            [type(node).__name__ for node in paragraph_ir.content],
            ["RunIR", "TableIR", "RunIR"],
        )
        self.assertEqual(
            [(run.native_anchor.debug_path, run.text) for run in paragraph_ir.runs],
            [("s1.p1.r1", "Before"), ("s1.p1.r3", "After")],
        )
        self.assertEqual(paragraph_ir.tables[0].native_anchor.debug_path, "s1.p1.r1.tbl1")
        self.assertEqual(paragraph_ir.tables[0].cells[0][0].paragraphs[0].runs[0].text, "Cell")

    def test_hwpx_nested_tables_are_parsed(self) -> None:
        hwpx_bytes_io = BytesIO()
        with zipfile.ZipFile(hwpx_bytes_io, "w") as zf:
            zf.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            zf.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p>
                <hp:run><hp:t>Outer</hp:t></hp:run>
                <hp:run>
                  <hp:tbl>
                    <hp:tr>
                      <hp:tc>
                        <hp:subList>
                          <hp:p><hp:run><hp:t>Inner</hp:t></hp:run></hp:p>
                        </hp:subList>
                        <hp:cellAddr colAddr="0" rowAddr="0"/>
                        <hp:cellSpan colSpan="1" rowSpan="1"/>
                      </hp:tc>
                    </hp:tr>
                  </hp:tbl>
                </hp:run>
              </hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )

        parsed = DocIR.from_file(hwpx_bytes_io.getvalue(), doc_type="hwpx")

        outer_cell_paragraph = parsed.paragraphs[0].tables[0].cells[0][0].paragraphs[0]
        self.assertEqual(outer_cell_paragraph.runs[0].text, "Outer")
        self.assertEqual(len(outer_cell_paragraph.tables), 1)
        self.assertEqual(
            outer_cell_paragraph.tables[0].cells[0][0].paragraphs[0].runs[0].text,
            "Inner",
        )


if __name__ == "__main__":
    unittest.main()
