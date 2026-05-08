from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path
import tempfile
import unittest
import zipfile

from document_processor import (
    DocumentInput,
    DocIR,
    StyleEdit,
    StructuralEdit,
    TextAnnotation,
    TextEdit,
    apply_document_edits,
    get_document_context,
    list_editable_targets,
    read_document,
    render_review_html,
)


class EditorApiTests(unittest.TestCase):
    @staticmethod
    def _text_hash(text: str) -> str:
        return hashlib.sha1(text.encode("utf-8")).hexdigest()

    @staticmethod
    def _build_sample_docx_bytes() -> bytes:
        from docx import Document

        docx = Document()
        paragraph = docx.add_paragraph()
        paragraph.add_run("Hello ")
        paragraph.add_run("World")
        docx.add_paragraph("Second paragraph")

        buffer = BytesIO()
        docx.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _build_sample_table_docx_bytes() -> bytes:
        from docx import Document

        docx = Document()
        table = docx.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Left"
        table.cell(0, 1).text = "Right"

        buffer = BytesIO()
        docx.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _build_sample_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run><hp:t>Hello </hp:t></hp:run>
    <hp:run><hp:t>World</hp:t></hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_sample_table_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p><hp:run><hp:t>Left</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
          <hp:tc>
            <hp:subList>
              <hp:p><hp:run><hp:t>Right</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="1" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_sample_styled_table_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">
  <hh:refList>
    <hh:borderFills itemCnt="1">
      <hh:borderFill id="0" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
        <hh:slash type="NONE" Crooked="0" isCounter="0"/>
        <hh:backSlash type="NONE" Crooked="0" isCounter="0"/>
        <hh:leftBorder type="SOLID" width="0.12 mm" color="#000000"/>
        <hh:rightBorder type="SOLID" width="0.12 mm" color="#000000"/>
        <hh:topBorder type="SOLID" width="0.12 mm" color="#000000"/>
        <hh:bottomBorder type="SOLID" width="0.12 mm" color="#000000"/>
        <hc:fillBrush><hc:winBrush faceColor="#DDEEFF" hatchColor="#999999" alpha="0"/></hc:fillBrush>
      </hh:borderFill>
    </hh:borderFills>
    <hh:charProperties itemCnt="1">
      <hh:charPr id="0" height="1000" textColor="#000000" shadeColor="none" useFontSpace="0" useKerning="0" symMark="NONE">
        <hh:fontRef hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>
        <hh:ratio hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/>
        <hh:spacing hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>
        <hh:relSz hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/>
        <hh:offset hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>
        <hh:underline type="NONE" shape="SOLID" color="#000000"/>
        <hh:strikeout shape="NONE" color="#000000"/>
      </hh:charPr>
    </hh:charProperties>
    <hh:paraProperties itemCnt="1">
      <hh:paraPr id="0" tabPrIDRef="0" condense="0" fontLineHeight="0" snapToGrid="1" suppressLineNumbers="0" checked="0">
        <hh:align horizontal="LEFT" vertical="BASELINE"/>
        <hh:margin>
          <hc:intent value="0" unit="HWPUNIT"/>
          <hc:left value="0" unit="HWPUNIT"/>
          <hc:right value="0" unit="HWPUNIT"/>
        </hh:margin>
      </hh:paraPr>
    </hh:paraProperties>
  </hh:refList>
</hh:head>
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p paraPrIDRef="0">
    <hp:run charPrIDRef="0"><hp:t>Hello</hp:t></hp:run>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc borderFillIDRef="0">
            <hp:subList vertAlign="TOP">
              <hp:p paraPrIDRef="0"><hp:run charPrIDRef="0"><hp:t>Cell</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
            <hp:cellSz width="5000" height="1000"/>
            <hp:cellMargin left="0" right="0" top="0" bottom="0"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_hwpx_table_after_intro_with_control_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p><hp:run><hp:t>Intro</hp:t></hp:run></hp:p>
  <hp:p>
    <hp:run>
      <hp:tbl>
        <hp:tr>
          <hp:tc>
            <hp:subList>
              <hp:p><hp:run><hp:t>Go </hp:t><hp:ctrl><hp:t>HIDDEN</hp:t></hp:ctrl><hp:t>Visible</hp:t></hp:run></hp:p>
            </hp:subList>
            <hp:cellAddr colAddr="0" rowAddr="0"/>
            <hp:cellSpan colSpan="1" rowSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
  </hp:p>
</hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    @staticmethod
    def _build_namespaced_hwpx_bytes() -> bytes:
        hwpx_bytes = BytesIO()
        with zipfile.ZipFile(hwpx_bytes, "w") as archive:
            archive.writestr(
                "Contents/header.xml",
                """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" />
""",
            )
            archive.writestr(
                "Contents/section0.xml",
                """<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><hs:sec xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core"><hc:pt0 x="0" y="0"/><hp:p><hp:run><hp:t>Hello </hp:t></hp:run><hp:run><hp:t>World</hp:t></hp:run></hp:p></hs:sec>
""",
            )
        return hwpx_bytes.getvalue()

    def test_get_document_context_accepts_bytes_backed_input(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        target_id = doc.paragraphs[0].runs[1].node_id

        result = get_document_context(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            target_ids=[target_id],
            before=0,
            after=1,
        )

        self.assertEqual(result.source_name, "sample.docx")
        self.assertEqual([paragraph.node_id for paragraph in result.paragraphs], [doc.paragraphs[0].node_id, doc.paragraphs[1].node_id])
        self.assertTrue(result.paragraphs[0].node_id.startswith("p_"))
        self.assertEqual(result.paragraphs[0].runs[1].text, "World")
        self.assertTrue(result.paragraphs[0].runs[1].node_id.startswith("r_"))
        self.assertEqual(result.paragraphs[0].text, "Hello World")
        self.assertEqual(result.paragraphs[0].text_hash, self._text_hash("Hello World"))
        self.assertEqual(result.paragraphs[0].runs[1].text_hash, self._text_hash("World"))
        self.assertEqual(
            [(run.text, run.start, run.end) for run in result.paragraphs[0].runs],
            [("Hello ", 0, 6), ("World", 6, 11)],
        )

    def test_document_input_accepts_path_source_path(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            source.write_bytes(self._build_sample_docx_bytes())

            result = read_document(
                document=DocumentInput(source_path=source),
                limit=1,
            )

        self.assertEqual(result.source_name, "sample.docx")
        self.assertEqual(result.paragraphs[0].text, "Hello World")

    def test_read_document_returns_bounded_stable_ids(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "First",
                "s1.p2.r1": "Second",
                "s1.p3.r1": "Third",
            },
            source_doc_type="docx",
        )

        result = read_document(
            document=DocumentInput(doc_ir=doc),
            start=1,
            limit=1,
        )

        self.assertEqual(result.total_paragraphs, 3)
        self.assertEqual(result.next_start, 2)
        self.assertEqual(result.paragraphs[0].text, "Second")
        self.assertEqual(result.paragraphs[0].node_id, doc.paragraphs[1].node_id)
        self.assertEqual(result.paragraphs[0].native_anchor.debug_path, "s1.p2")
        self.assertEqual(result.paragraphs[0].text_hash, self._text_hash("Second"))
        self.assertEqual(result.paragraphs[0].runs[0].start, 0)
        self.assertEqual(result.paragraphs[0].runs[0].end, len("Second"))

    def test_apply_document_edits_returns_output_bytes_for_text_edit_bytes_source(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                TextEdit(
                    target_kind="paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    expected_text_hash=self._text_hash("Hello World"),
                    new_text="Hello Legal World",
                    reason="Expand wording",
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.output_filename, "sample_edited.docx")
        self.assertIsNone(result.output_path)
        self.assertIsNotNone(result.output_bytes)
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual(DocIR.from_file(result.output_bytes).paragraphs[0].text, "Hello Legal World")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Legal World")

    def test_apply_document_edits_with_doc_ir_only_returns_updated_doc_ir_for_text_edit(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    expected_text_hash=self._text_hash("Hello World"),
                    new_text="Hello Contract World",
                    reason="Expand wording",
                )
            ],
        )

        self.assertTrue(result.ok)
        self.assertIsNone(result.output_path)
        self.assertIsNone(result.output_bytes)
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Contract World")

    def test_apply_document_edits_applies_flat_style_edit_to_doc_ir(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello",
            },
            source_doc_type="docx",
        )
        run_id = doc.paragraphs[0].runs[0].node_id

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                StyleEdit(
                    target_kind="run",
                    target_id=run_id,
                    bold=True,
                    color="#112233",
                    font_size_pt=14,
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.styles_applied, 1)
        updated_run = result.updated_doc_ir.paragraphs[0].runs[0]
        self.assertTrue(updated_run.run_style.bold)
        self.assertEqual(updated_run.run_style.color, "#112233")
        self.assertEqual(updated_run.run_style.size_pt, 14)

    def test_apply_document_edits_infers_style_target_kind(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello",
            },
            source_doc_type="docx",
        )
        run_id = doc.paragraphs[0].runs[0].node_id

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                StyleEdit(
                    client_edit_id="style-1",
                    target_id=run_id,
                    bold=True,
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertTrue(result.updated_doc_ir.paragraphs[0].runs[0].run_style.bold)
        self.assertEqual(result.edit_results[0].client_edit_id, "style-1")
        self.assertEqual(result.edit_results[0].target_kind, "run")
        self.assertEqual(result.edit_results[0].modified_target_ids, [run_id])

    def test_apply_document_edits_rejects_inferred_style_field_mismatch(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello",
            },
            source_doc_type="docx",
        )

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                StyleEdit(
                    client_edit_id="style-1",
                    target_id=doc.paragraphs[0].runs[0].node_id,
                    background="#FFFF00",
                )
            ],
            return_doc_ir=True,
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "invalid_style")
        self.assertEqual(result.edit_results[0].client_edit_id, "style-1")
        self.assertFalse(result.edit_results[0].ok)

    def test_apply_document_edits_broadcasts_doc_ir_cell_dimensions(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "A1",
                "s1.p1.r1.tbl1.tr1.tc2.p1.r1": "A2",
                "s1.p1.r1.tbl1.tr2.tc1.p1.r1": "B1",
                "s1.p1.r1.tbl1.tr2.tc2.p1.r1": "B2",
            },
            source_doc_type="docx",
        )
        table = doc.paragraphs[0].tables[0]
        target_cell = table.cells[0][0]

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                StyleEdit(
                    target_kind="cell",
                    target_id=target_cell.node_id,
                    background="#FFF2CC",
                    width_pt=100,
                    height_pt=50,
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        updated_table = result.updated_doc_ir.paragraphs[0].tables[0]
        cells = {(row_index, col_index): cell for row_index, col_index, cell in updated_table.iter_cell_positions()}

        self.assertEqual(cells[(1, 1)].cell_style.background, "#FFF2CC")
        self.assertIsNone(cells[(1, 2)].cell_style.background)
        self.assertIsNone(cells[(2, 1)].cell_style.background)
        self.assertIsNone(cells[(2, 2)].cell_style)

        self.assertEqual(cells[(1, 1)].cell_style.width_pt, 100)
        self.assertEqual(cells[(2, 1)].cell_style.width_pt, 100)
        self.assertIsNone(cells[(1, 2)].cell_style.width_pt)

        self.assertEqual(cells[(1, 1)].cell_style.height_pt, 50)
        self.assertEqual(cells[(1, 2)].cell_style.height_pt, 50)
        self.assertIsNone(cells[(2, 1)].cell_style.height_pt)

    def test_apply_document_edits_writes_docx_run_style(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        run_id = doc.paragraphs[0].runs[1].node_id

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                StyleEdit(
                    target_kind="run",
                    target_id=run_id,
                    bold=True,
                    color="#445566",
                    font_size_pt=16,
                )
            ],
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.styles_applied, 1)
        reparsed = DocIR.from_file(result.output_bytes, doc_type="docx")
        updated_run = reparsed.paragraphs[0].runs[1]
        self.assertTrue(updated_run.run_style.bold)
        self.assertEqual(updated_run.run_style.color, "#445566")
        self.assertAlmostEqual(updated_run.run_style.size_pt, 16.0)

    def test_style_edit_rejects_table_level_size_fields(self) -> None:
        with self.assertRaises(ValueError) as error:
            StyleEdit(
                target_kind="table",
                target_id="table-1",
                width_pt=360,
                height_pt=72,
            )

        self.assertIn("table style edits do not support fields", str(error.exception))

    def test_apply_document_edits_writes_docx_cell_geometry_and_background(self) -> None:
        from xml.etree import ElementTree as ET

        source_bytes = self._build_sample_table_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        cell = doc.paragraphs[0].tables[0].cells[0][0]

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                StyleEdit(
                    target_kind="cell",
                    target_id=cell.node_id,
                    width_pt=144,
                    height_pt=36,
                    background="#FFF2CC",
                )
            ],
        )

        self.assertTrue(result.ok, result.validation.issues)
        reparsed = DocIR.from_file(result.output_bytes, doc_type="docx")
        updated_cell = reparsed.paragraphs[0].tables[0].cells[0][0]
        self.assertAlmostEqual(updated_cell.cell_style.width_pt, 144.0)
        self.assertAlmostEqual(updated_cell.cell_style.height_pt, 36.0)
        self.assertEqual(updated_cell.cell_style.background, "#FFF2CC")

        with zipfile.ZipFile(BytesIO(result.output_bytes)) as archive:
            document_xml = archive.read("word/document.xml")
        root = ET.fromstring(document_xml)
        w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        self.assertEqual(root.find(f".//{w}tblLayout").get(f"{w}type"), "fixed")
        self.assertEqual(root.findall(f".//{w}tblGrid/{w}gridCol")[0].get(f"{w}w"), "2880")
        self.assertEqual(root.find(f".//{w}tcW").get(f"{w}w"), "2880")
        self.assertEqual(root.find(f".//{w}trHeight").get(f"{w}val"), "720")
        self.assertEqual(root.find(f".//{w}shd").get(f"{w}fill"), "FFF2CC")

    def test_apply_document_edits_writes_hwpx_run_paragraph_and_cell_style(self) -> None:
        from xml.etree import ElementTree as ET

        source_bytes = self._build_sample_styled_table_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        paragraph = doc.paragraphs[0]
        run = paragraph.runs[0]
        cell = paragraph.tables[0].cells[0][0]

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.hwpx",
            ),
            edits=[
                StyleEdit(
                    target_kind="run",
                    target_id=run.node_id,
                    bold=True,
                    color="#445566",
                    font_size_pt=16,
                ),
                StyleEdit(
                    target_kind="paragraph",
                    target_id=paragraph.node_id,
                    paragraph_align="center",
                    left_indent_pt=18,
                ),
                StyleEdit(
                    target_kind="cell",
                    target_id=cell.node_id,
                    background="#FFF2CC",
                    vertical_align="middle",
                    horizontal_align="center",
                    width_pt=120,
                    height_pt=50,
                    padding_left_pt=6,
                    padding_right_pt=6,
                    border_top="1pt single #445566",
                    border_right="1pt single #445566",
                    border_bottom="1pt single #445566",
                    border_left="1pt single #445566",
                ),
            ],
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertEqual(result.styles_applied, 3)
        with zipfile.ZipFile(BytesIO(result.output_bytes)) as archive:
            header_root = ET.fromstring(archive.read("Contents/header.xml"))
            section_root = ET.fromstring(archive.read("Contents/section0.xml"))
        hh = "{http://www.hancom.co.kr/hwpml/2011/head}"
        hp = "{http://www.hancom.co.kr/hwpml/2011/paragraph}"
        hc = "{http://www.hancom.co.kr/hwpml/2011/core}"
        updated_cell_el = section_root.find(f".//{hp}tc")
        border_fill_id = updated_cell_el.get("borderFillIDRef")
        border_fill = next(
            element
            for element in header_root.findall(f".//{hh}borderFill")
            if element.get("id") == border_fill_id
        )
        fill_brushes = [child for child in list(border_fill) if child.tag.rsplit("}", 1)[-1] == "fillBrush"]
        self.assertEqual(len(fill_brushes), 1)
        self.assertEqual(fill_brushes[0].find(f"{hc}winBrush").get("faceColor"), "#FFF2CC")
        cell_size = updated_cell_el.find(f"{hp}cellSz")
        self.assertEqual(cell_size.get("width"), "12000")
        self.assertEqual(cell_size.get("height"), "5000")
        sub_list = updated_cell_el.find(f"{hp}subList")
        self.assertEqual(sub_list.get("textWidth"), "10800")
        self.assertEqual(sub_list.get("textHeight"), "5000")

        reparsed = DocIR.from_file(result.output_bytes, doc_type="hwpx")
        updated_run = reparsed.paragraphs[0].runs[0]
        self.assertTrue(updated_run.run_style.bold)
        self.assertEqual(updated_run.run_style.color, "#445566")
        self.assertAlmostEqual(updated_run.run_style.size_pt, 16.0)

        updated_paragraph = reparsed.paragraphs[0]
        self.assertEqual(updated_paragraph.para_style.align, "center")
        self.assertAlmostEqual(updated_paragraph.para_style.left_indent_pt, 18.0)

        updated_cell = reparsed.paragraphs[0].tables[0].cells[0][0]
        self.assertEqual(updated_cell.cell_style.background, "#FFF2CC")
        self.assertIsNone(updated_cell.cell_style.horizontal_align)
        self.assertEqual(updated_cell.paragraphs[0].para_style.align, "center")
        self.assertEqual(updated_cell.cell_style.vertical_align, "center")
        self.assertAlmostEqual(updated_cell.cell_style.width_pt, 120.0)
        self.assertAlmostEqual(updated_cell.cell_style.height_pt, 50.0)
        self.assertAlmostEqual(updated_cell.cell_style.padding_left_pt, 6.0)
        self.assertAlmostEqual(updated_cell.cell_style.padding_right_pt, 6.0)
        self.assertEqual(updated_cell.cell_style.border_top, "2px solid #445566")

    def test_apply_document_edits_writes_hwpx_cell_border_px_width(self) -> None:
        from xml.etree import ElementTree as ET

        source_bytes = self._build_sample_styled_table_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        cell = doc.paragraphs[0].tables[0].cells[0][0]

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.hwpx",
            ),
            edits=[
                StyleEdit(
                    target_kind="cell",
                    target_id=cell.node_id,
                    border_top="3px solid #445566",
                ),
            ],
        )

        self.assertTrue(result.ok, result.validation.issues)
        with zipfile.ZipFile(BytesIO(result.output_bytes)) as archive:
            header_root = ET.fromstring(archive.read("Contents/header.xml"))
            section_root = ET.fromstring(archive.read("Contents/section0.xml"))

        hh = "{http://www.hancom.co.kr/hwpml/2011/head}"
        hp = "{http://www.hancom.co.kr/hwpml/2011/paragraph}"
        border_fill_id = section_root.find(f".//{hp}tc").get("borderFillIDRef")
        border_fill = next(
            element
            for element in header_root.findall(f".//{hh}borderFill")
            if element.get("id") == border_fill_id
        )
        top_border = border_fill.find(f"{hh}topBorder")
        self.assertEqual(top_border.get("width"), "0.7 mm")

        reparsed = DocIR.from_file(result.output_bytes, doc_type="hwpx")
        updated_cell = reparsed.paragraphs[0].tables[0].cells[0][0]
        self.assertEqual(updated_cell.cell_style.border_top, "3px solid #445566")

    def test_apply_document_edits_writes_hwpx_table_placement_style(self) -> None:
        source_bytes = self._build_sample_table_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        table = doc.paragraphs[0].tables[0]

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.hwpx",
            ),
            edits=[
                StyleEdit(
                    target_kind="table",
                    target_id=table.node_id,
                    placement_mode="floating",
                    wrap="square",
                    x_relative_to="page",
                    y_relative_to="paragraph",
                    x_offset_pt=10,
                    y_offset_pt=12,
                )
            ],
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.styles_applied, 1)
        with zipfile.ZipFile(BytesIO(result.output_bytes)) as archive:
            section_xml = archive.read("Contents/section0.xml").decode("utf-8")
        self.assertIn('textWrap="SQUARE"', section_xml)
        self.assertIn('treatAsChar="0"', section_xml)
        self.assertIn('horzOffset="1000"', section_xml)
        self.assertIn('vertOffset="1200"', section_xml)

    def test_apply_document_edits_accepts_stable_target_id_for_text_edit(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )
        target_id = doc.paragraphs[0].node_id

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="paragraph",
                    target_id=target_id,
                    expected_text_hash=self._text_hash("Hello World"),
                    new_text="Hello Stable World",
                )
            ],
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.modified_target_ids, [target_id])
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Stable World")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].node_id, target_id)

    def test_apply_document_edits_infers_text_target_kind(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello ",
                "s1.p1.r2": "World",
            },
            source_doc_type="docx",
        )
        target_id = doc.paragraphs[0].node_id

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    client_edit_id="edit-1",
                    target_id=target_id,
                    expected_text_hash=self._text_hash("Hello World"),
                    new_text="Hello Inferred World",
                )
            ],
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello Inferred World")
        self.assertEqual(len(result.edit_results), 1)
        self.assertEqual(result.edit_results[0].client_edit_id, "edit-1")
        self.assertEqual(result.edit_results[0].target_kind, "paragraph")
        self.assertEqual(result.edit_results[0].modified_target_ids, [target_id])

    def test_apply_document_edits_rejects_explicit_text_target_kind_mismatch(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"}, source_doc_type="docx")

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    client_edit_id="edit-1",
                    target_kind="paragraph",
                    target_id=doc.paragraphs[0].runs[0].node_id,
                    expected_text_hash=self._text_hash("Hello"),
                    new_text="Changed",
                )
            ],
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "target_kind_mismatch")
        self.assertEqual(result.edit_results[0].client_edit_id, "edit-1")
        self.assertFalse(result.edit_results[0].ok)

    def test_apply_document_edits_dry_run_returns_text_preview_without_native_output(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"}, source_doc_type="docx")

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="run",
                    target_id=doc.paragraphs[0].runs[0].node_id,
                    expected_text_hash=self._text_hash("Hello"),
                    new_text="Preview",
                )
            ],
            dry_run=True,
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.edits_applied, 0)
        self.assertIsNone(result.output_path)
        self.assertIsNone(result.output_bytes)
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Preview")
        self.assertEqual(doc.paragraphs[0].text, "Hello")
        self.assertEqual(result.edit_results[0].target_kind, "run")
        self.assertEqual(result.edit_results[0].modified_target_ids, [doc.paragraphs[0].runs[0].node_id])

    def test_apply_document_edits_rejects_missing_text_target_id(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "Hello",
                "s1.p2.r1": "Other",
            },
            source_doc_type="docx",
        )

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="paragraph",
                    target_id="missing",
                    expected_text_hash=self._text_hash("Hello"),
                    new_text="Changed",
                )
            ],
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "target_not_found")

    def test_apply_document_edits_rejects_text_hash_mismatch(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"}, source_doc_type="docx")

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_id=doc.paragraphs[0].node_id,
                    expected_text_hash=self._text_hash("Wrong"),
                    new_text="Changed",
                )
            ],
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "text_hash_mismatch")
        self.assertEqual(result.validation.issues[0].current_text_hash, self._text_hash("Hello"))
        self.assertEqual(result.edit_results[0].validation_issue.code, "text_hash_mismatch")

    def test_list_editable_targets_includes_cell_targets(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Left",
                "s1.p1.r1.tbl1.tr1.tc2.p1.r1": "Right",
            },
            source_doc_type="docx",
        )

        result = list_editable_targets(
            document=DocumentInput(doc_ir=doc),
            target_kinds=["cell"],
        )

        self.assertEqual(
            [(target.target_kind, target.native_anchor.debug_path, target.current_text) for target in result.targets],
            [
                ("cell", "s1.p1.r1.tbl1.tr1.tc1", "Left"),
                ("cell", "s1.p1.r1.tbl1.tr1.tc2", "Right"),
            ],
        )
        self.assertEqual([target.text_hash for target in result.targets], [self._text_hash("Left"), self._text_hash("Right")])
        table_id = doc.paragraphs[0].tables[0].node_id
        self.assertEqual(
            [
                (target.parent_table_id, target.row_index, target.column_index, target.rowspan, target.colspan)
                for target in result.targets
            ],
            [
                (table_id, 1, 1, 1, 1),
                (table_id, 1, 2, 1, 1),
            ],
        )

    def test_list_editable_targets_does_not_return_all_targets_for_missing_filter(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})

        result = list_editable_targets(
            document=DocumentInput(doc_ir=doc),
            target_ids=["missing"],
        )

        self.assertEqual(result.targets, [])
        self.assertEqual(result.missing_target_ids, ["missing"])

    def test_apply_document_edits_can_replace_doc_ir_cell_text(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1.tbl1.tr1.tc1.p1.r1": "Left",
                "s1.p1.r1.tbl1.tr1.tc2.p1.r1": "Right",
            },
            source_doc_type="docx",
        )

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                TextEdit(
                    target_kind="cell",
                    target_id=doc.paragraphs[0].tables[0].cells[0][0].node_id,
                    expected_text_hash=self._text_hash("Left"),
                    new_text="Changed",
                )
            ],
        )

        self.assertTrue(result.ok)
        cell = doc.paragraphs[0].tables[0].cells[0][0]
        run = cell.paragraphs[0].runs[0]
        self.assertEqual(result.modified_target_ids, [cell.node_id])
        self.assertEqual(result.modified_run_ids, [run.node_id])
        table = result.updated_doc_ir.paragraphs[0].tables[0]
        self.assertEqual(table.cells[0][0].text, "Changed")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Changed\nRight")

    def test_apply_document_edits_replaces_docx_cell_text(self) -> None:
        source_bytes = self._build_sample_table_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="table.docx",
            ),
            edits=[
                TextEdit(
                    target_kind="cell",
                    target_id=doc.paragraphs[0].tables[0].cells[0][0].node_id,
                    expected_text_hash=self._text_hash("Left"),
                    new_text="Changed",
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.output_filename, "table_edited.docx")
        self.assertEqual(DocIR.from_file(result.output_bytes).paragraphs[0].tables[0].cells[0][0].text, "Changed")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].tables[0].cells[0][0].text, "Changed")

    def test_apply_document_edits_replaces_hwpx_cell_text(self) -> None:
        source_bytes = self._build_sample_table_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="table.hwpx",
            ),
            edits=[
                TextEdit(
                    target_kind="cell",
                    target_id=doc.paragraphs[0].tables[0].cells[0][0].node_id,
                    expected_text_hash=self._text_hash("Left"),
                    new_text="Changed",
                )
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.output_filename, "table_edited.hwpx")
        self.assertEqual(DocIR.from_file(result.output_bytes, doc_type="hwpx").paragraphs[0].tables[0].cells[0][0].text, "Changed")
        self.assertEqual(result.updated_doc_ir.paragraphs[0].tables[0].cells[0][0].text, "Changed")

    def test_apply_document_edits_normalizes_hwpx_output_suffix_for_path_backed_writeback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            requested_output = Path(tmp_dir) / "sample_edited.docx"
            source.write_bytes(self._build_sample_hwpx_bytes())
            doc = DocIR.from_file(source, doc_type="hwpx")

            result = apply_document_edits(
                document=DocumentInput(source_path=str(source)),
                edits=[
                    TextEdit(
                        target_kind="run",
                        target_id=doc.paragraphs[0].runs[1].node_id,
                        expected_text_hash=self._text_hash("World"),
                        new_text="HWPX",
                        reason="Rename token",
                    )
                ],
                output_path=str(requested_output),
                return_doc_ir=True,
            )

            self.assertTrue(result.ok)
            self.assertEqual(Path(result.output_path).suffix, ".hwpx")
            self.assertEqual(Path(result.output_path).name, "sample_edited.hwpx")
            self.assertTrue(any("adjusted output path" in warning for warning in result.warnings))
            self.assertEqual(result.updated_doc_ir.paragraphs[0].text, "Hello HWPX")

    def test_apply_document_edits_rejects_invalid_output_filename_extensions(self) -> None:
        source_bytes = self._build_sample_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        edit = TextEdit(
            target_kind="run",
            target_id=doc.paragraphs[0].runs[1].node_id,
            expected_text_hash=self._text_hash("World"),
            new_text="HWPX",
            reason="Rename token",
        )

        mismatched = apply_document_edits(
            document=DocumentInput(source_bytes=source_bytes, source_name="sample.hwpx"),
            edits=[edit],
            output_filename="sample_edited.docx",
        )
        unsupported = apply_document_edits(
            document=DocumentInput(source_bytes=source_bytes, source_name="sample.hwpx"),
            edits=[edit],
            output_filename="sample_edited.txt",
        )

        self.assertFalse(mismatched.ok)
        self.assertIn("does not match", mismatched.validation.issues[0].message)
        self.assertFalse(unsupported.ok)
        self.assertIn("supported write-back extension", unsupported.validation.issues[0].message)

    def test_apply_document_edits_preserves_hwpx_namespace_prefixes_and_declaration(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.hwpx"
            output = Path(tmp_dir) / "sample_edited.hwpx"
            source.write_bytes(self._build_namespaced_hwpx_bytes())
            doc = DocIR.from_file(source, doc_type="hwpx")

            result = apply_document_edits(
                document=DocumentInput(source_path=str(source)),
                edits=[
                    TextEdit(
                        target_kind="run",
                        target_id=doc.paragraphs[0].runs[1].node_id,
                        expected_text_hash=self._text_hash("World"),
                        new_text="HWPX",
                        reason="Rename token",
                    )
                ],
                output_path=str(output),
            )

            self.assertTrue(result.ok)
            with zipfile.ZipFile(output) as archive:
                section_xml = archive.read("Contents/section0.xml")

            self.assertTrue(section_xml.startswith(b'<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>'))
            self.assertIn(b'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core"', section_xml)
            self.assertIn(b"<hc:pt0", section_xml)
            self.assertNotIn(b"xmlns:ns", section_xml)
            self.assertIn(b"HWPX", section_xml)

    def test_apply_document_edits_updates_doc_ir_structure_with_stable_ids(self) -> None:
        doc = DocIR.from_mapping(
            {
                "s1.p1.r1": "First",
                "s1.p2.r1": "Second",
                "s1.p3.r1.tbl1.tr1.tc1.p1.r1": "Left",
                "s1.p3.r1.tbl1.tr1.tc2.p1.r1": "Right",
            },
            source_doc_type="docx",
        )
        second_id = doc.paragraphs[1].node_id
        cell_id = doc.paragraphs[2].tables[0].cells[0][0].node_id

        result = apply_document_edits(
            document=DocumentInput(doc_ir=doc),
            edits=[
                StructuralEdit(
                    operation="insert_paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    position="after",
                    text="Inserted",
                ),
                StructuralEdit(
                    operation="insert_run",
                    target_id=second_id,
                    position="end",
                    text="!",
                ),
                StructuralEdit(
                    operation="set_cell_text",
                    target_id=cell_id,
                    expected_text_hash=self._text_hash("Left"),
                    text="Line one\nLine two",
                ),
                StructuralEdit(
                    operation="insert_table_row",
                    target_id=cell_id,
                    position="after",
                    values=["Bottom left", "Bottom right"],
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.operations_applied, 4)
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual([p.text for p in result.updated_doc_ir.paragraphs[:3]], ["First", "Inserted", "Second!"])
        self.assertEqual(result.updated_doc_ir.paragraphs[2].node_id, second_id)
        self.assertEqual(result.updated_doc_ir.paragraphs[2].native_anchor.structural_path, "s1.p3")
        table = result.updated_doc_ir.paragraphs[3].tables[0]
        self.assertEqual(table.row_count, 2)
        self.assertEqual(table.cells[0][0].text, "Line one\nLine two")
        self.assertTrue(result.created_target_ids)

    def test_apply_document_edits_writes_docx_paragraph_table_and_cells(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    position="before",
                    text="Preface",
                ),
                StructuralEdit(
                    operation="insert_table",
                    target_id=doc.paragraphs[1].node_id,
                    position="after",
                    rows=[["A", "B"], ["C", "D"]],
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertEqual(result.output_filename, "sample_edited.docx")
        self.assertIsNotNone(result.output_bytes)
        parsed = DocIR.from_file(result.output_bytes, doc_type="docx")
        self.assertEqual(parsed.paragraphs[0].text, "Preface")
        self.assertEqual(parsed.paragraphs[-1].tables[0].cells[1][1].text, "D")
        self.assertIsNotNone(result.updated_doc_ir)
        self.assertEqual(result.updated_doc_ir.paragraphs[1].node_id, doc.paragraphs[0].node_id)

    def test_apply_document_edits_writes_visible_docx_table_defaults(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_table",
                    target_id=doc.paragraphs[0].node_id,
                    position="after",
                    rows=[["A", "B"], ["C", "D"]],
                ),
            ],
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertIsNotNone(result.output_bytes)
        with zipfile.ZipFile(BytesIO(result.output_bytes)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        self.assertIn("<w:tblBorders>", document_xml)
        self.assertIn("<w:tblLayout w:type=\"fixed\"/>", document_xml)
        self.assertRegex(document_xml, r"<w:gridCol w:w=\"[1-9][0-9]*\"/>")
        self.assertRegex(document_xml, r"<w:tcW w:type=\"dxa\" w:w=\"[1-9][0-9]*\"/>")

        parsed = DocIR.from_file(result.output_bytes, doc_type="docx")
        table = parsed.paragraphs[1].tables[0]
        self.assertIsNotNone(table.table_style)
        self.assertGreater(table.table_style.width_pt or 0, 0)
        self.assertEqual(table.cells[0][0].cell_style.border_top, "1px solid #000000")
        self.assertGreater(table.cells[0][0].cell_style.width_pt or 0, 0)

    def test_apply_document_edits_mixes_structural_and_text_edits_in_order(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        second_id = doc.paragraphs[1].node_id

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    position="before",
                    text="Preface",
                ),
                TextEdit(
                    target_kind="paragraph",
                    target_id=second_id,
                    expected_text_hash=self._text_hash("Second paragraph"),
                    new_text="Updated second paragraph",
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        parsed = DocIR.from_file(result.output_bytes, doc_type="docx")
        self.assertEqual([paragraph.text for paragraph in parsed.paragraphs[:3]], ["Preface", "Hello World", "Updated second paragraph"])
        self.assertEqual(result.updated_doc_ir.paragraphs[2].node_id, second_id)
        self.assertEqual(result.updated_doc_ir.paragraphs[2].native_anchor.structural_path, "s1.p3")

    def test_apply_document_edits_preview_assigns_page_numbers_to_inserted_nodes(self) -> None:
        source_bytes = self._build_sample_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        anchor_page_number = doc.paragraphs[0].page_number

        self.assertIsNotNone(anchor_page_number)

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.docx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    position="after",
                    text="Inserted paragraph",
                ),
                StructuralEdit(
                    operation="insert_table",
                    target_id=doc.paragraphs[0].node_id,
                    position="after",
                    rows=[["A", "B"]],
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertIsNotNone(result.updated_doc_ir)
        updated = result.updated_doc_ir
        table_wrapper = updated.paragraphs[1]
        inserted_paragraph = updated.paragraphs[2]

        self.assertEqual(table_wrapper.page_number, anchor_page_number)
        self.assertEqual(inserted_paragraph.page_number, anchor_page_number)
        self.assertEqual(table_wrapper.tables[0].cells[0][0].paragraphs[0].page_number, anchor_page_number)

        review = render_review_html(
            document=DocumentInput(doc_ir=updated),
            annotations=[],
        )

        self.assertTrue(review.ok, review.validation.issues)
        self.assertIsNotNone(review.html)
        self.assertNotIn('<section class="document-unpaged">', review.html)
        self.assertIn("Inserted paragraph", review.html)

    def test_apply_document_edits_writes_docx_table_row_column_and_cell_text(self) -> None:
        source_bytes = self._build_sample_table_docx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="docx")
        table = doc.paragraphs[0].tables[0]
        left_cell_id = table.cells[0][0].node_id

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="table.docx",
            ),
            edits=[
                StructuralEdit(
                    operation="set_cell_text",
                    target_id=left_cell_id,
                    expected_text_hash=self._text_hash("Left"),
                    text="Changed",
                ),
                StructuralEdit(
                    operation="insert_table_row",
                    target_id=left_cell_id,
                    position="after",
                    values=["Bottom left", "Bottom right"],
                ),
                StructuralEdit(
                    operation="insert_table_column",
                    target_id=left_cell_id,
                    position="after",
                    values=["Middle top", "Middle bottom"],
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        parsed = DocIR.from_file(result.output_bytes, doc_type="docx")
        parsed_table = parsed.paragraphs[0].tables[0]
        self.assertEqual(parsed_table.row_count, 2)
        self.assertEqual(parsed_table.col_count, 3)
        self.assertEqual(parsed_table.cells[0][0].text, "Changed")
        self.assertEqual(parsed_table.cells[0][1].text, "Middle top")
        self.assertEqual(parsed_table.cells[-1][-1].text, "Bottom right")

    def test_apply_document_edits_writes_hwpx_paragraph_and_table_changes(self) -> None:
        source_bytes = self._build_sample_table_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        table = doc.paragraphs[0].tables[0]
        left_cell_id = table.cells[0][0].node_id

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="table.hwpx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_paragraph",
                    target_id=doc.paragraphs[0].node_id,
                    position="after",
                    text="After table",
                ),
                StructuralEdit(
                    operation="set_cell_text",
                    target_id=left_cell_id,
                    expected_text_hash=self._text_hash("Left"),
                    text="Changed",
                ),
                StructuralEdit(
                    operation="insert_table_row",
                    target_id=left_cell_id,
                    position="after",
                    values=["Bottom left", "Bottom right"],
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        parsed = DocIR.from_file(result.output_bytes, doc_type="hwpx")
        self.assertEqual(parsed.paragraphs[0].tables[0].row_count, 2)
        self.assertEqual(parsed.paragraphs[0].tables[0].cells[0][0].text, "Changed")
        self.assertEqual(parsed.paragraphs[1].text, "After table")

    def test_apply_document_edits_writes_inline_visible_hwpx_table_defaults(self) -> None:
        source_bytes = self._build_sample_hwpx_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="sample.hwpx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_table",
                    target_id=doc.paragraphs[0].node_id,
                    position="after",
                    rows=[["A", "B"], ["C", "D"]],
                ),
            ],
        )

        self.assertTrue(result.ok, result.validation.issues)
        self.assertIsNotNone(result.output_bytes)
        with zipfile.ZipFile(BytesIO(result.output_bytes)) as archive:
            header_xml = archive.read("Contents/header.xml").decode("utf-8")
            section_xml = archive.read("Contents/section0.xml").decode("utf-8")

        self.assertIn("type=\"SOLID\"", header_xml)
        self.assertIn("textWrap=\"TOP_AND_BOTTOM\"", section_xml)
        self.assertIn("treatAsChar=\"1\"", section_xml)
        self.assertRegex(section_xml, r"borderFillIDRef=\"[1-9][0-9]*\"")
        self.assertRegex(section_xml, r"<hp:cellSz width=\"[1-9][0-9]*\" height=\"[1-9][0-9]*\" ?/>")
        self.assertIn("<hp:cellMargin", section_xml)

        parsed = DocIR.from_file(result.output_bytes, doc_type="hwpx")
        table = parsed.paragraphs[1].tables[0]
        self.assertIsNotNone(table.table_style)
        self.assertGreater(table.table_style.width_pt or 0, 0)
        self.assertEqual(table.cells[0][0].cell_style.border_top, "1px solid #000000")
        self.assertGreater(table.cells[0][0].cell_style.width_pt or 0, 0)

    def test_apply_document_edits_writes_hwpx_cell_after_prior_inserts_and_ignores_control_text(self) -> None:
        source_bytes = self._build_hwpx_table_after_intro_with_control_bytes()
        doc = DocIR.from_file(source_bytes, doc_type="hwpx")
        intro_id = doc.paragraphs[0].node_id
        original_cell = doc.paragraphs[1].tables[0].cells[0][0]

        self.assertEqual(original_cell.text, "Go Visible")

        result = apply_document_edits(
            document=DocumentInput(
                source_bytes=source_bytes,
                source_name="table.hwpx",
            ),
            edits=[
                StructuralEdit(
                    operation="insert_table",
                    target_id=intro_id,
                    position="after",
                    rows=[["Inserted table"]],
                ),
                StructuralEdit(
                    operation="insert_paragraph",
                    target_id=intro_id,
                    position="after",
                    text="Inserted paragraph",
                ),
                StructuralEdit(
                    operation="set_cell_text",
                    target_id=original_cell.node_id,
                    expected_text_hash=self._text_hash("Go Visible"),
                    text="Changed",
                ),
            ],
            return_doc_ir=True,
        )

        self.assertTrue(result.ok, result.validation.issues)
        parsed = DocIR.from_file(result.output_bytes, doc_type="hwpx")
        self.assertEqual([paragraph.text for paragraph in parsed.paragraphs], ["Intro", "Inserted paragraph", "Inserted table", "Changed"])
        self.assertNotIn("HIDDEN", "".join(paragraph.text for paragraph in parsed.paragraphs))
        self.assertIsNotNone(result.updated_doc_ir)
        updated_cell = result.updated_doc_ir.paragraphs[3].tables[0].cells[0][0]
        self.assertEqual(updated_cell.node_id, original_cell.node_id)
        self.assertEqual(updated_cell.native_anchor.structural_path, "s1.p4.r1.tbl1.tr1.tc1")

    def test_render_review_html_accepts_doc_ir_input(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})

        result = render_review_html(
            document=DocumentInput(doc_ir=doc),
            annotations=[
                TextAnnotation(
                    target_kind="run",
                    target_id=doc.paragraphs[0].runs[0].node_id,
                    selected_text="Hello",
                    label="Greeting",
                )
            ],
        )

        self.assertTrue(result.ok)
        self.assertIn("<mark", result.html or "")
        self.assertEqual(result.resolved_annotations[0].selected_text, "Hello")

    def test_render_review_html_accepts_stable_target_id(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello"})
        run_id = doc.paragraphs[0].runs[0].node_id

        result = render_review_html(
            document=DocumentInput(doc_ir=doc),
            annotations=[
                TextAnnotation(
                    target_kind="run",
                    target_id=run_id,
                    selected_text="Hello",
                    label="Greeting",
                )
            ],
        )

        self.assertTrue(result.ok)
        self.assertEqual(result.resolved_annotations[0].target_id, run_id)

    def test_render_review_html_rejects_ambiguous_selected_text(self) -> None:
        doc = DocIR.from_mapping({"s1.p1.r1": "Hello Hello"})

        result = render_review_html(
            document=DocumentInput(doc_ir=doc),
            annotations=[
                TextAnnotation(
                    target_kind="run",
                    target_id=doc.paragraphs[0].runs[0].node_id,
                    selected_text="Hello",
                    label="Ambiguous",
                )
            ],
        )

        self.assertFalse(result.ok)
        self.assertEqual(result.validation.issues[0].code, "selected_text_ambiguous")


if __name__ == "__main__":
    unittest.main()
