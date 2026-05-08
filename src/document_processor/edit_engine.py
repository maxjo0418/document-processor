from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
import difflib
import hashlib
from io import BytesIO
from pathlib import Path
import re
from typing import BinaryIO, Callable
import tempfile
from xml.etree import ElementTree as ET
import zipfile

from pydantic import BaseModel, Field

from .api_types import AppliedEditResult, StyleEdit, StructuralEdit, TextEdit
from .core import convert_hwp_to_hwpx_bytes
from .io_utils import SourceDocType, TemporarySourcePath, coerce_source_to_supported_value, infer_doc_type
from .models import DocIR, ImageIR, NativeAnchor, NodeKind, ParagraphIR, RunIR, TableCellIR, TableIR, _anchored_node_id, _make_native_anchor
from .style_types import CellStyleInfo, ObjectPlacementInfo, ParaStyleInfo, RunStyleInfo, TableStyleInfo


class EditValidationError(ValueError):
    def __init__(
        self,
        message: str,
        *,
        code: str = "invalid_operation",
        target_kind: str | None = None,
        target_id: str | None = None,
        operation: str | None = None,
        expected_text_hash: str | None = None,
        current_text_hash: str | None = None,
        current_text: str | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.target_kind = target_kind
        self.target_id = target_id
        self.operation = operation
        self.expected_text_hash = expected_text_hash
        self.current_text_hash = current_text_hash
        self.current_text = current_text


class _EditEngineResult(BaseModel):
    model_config = {"arbitrary_types_allowed": True}

    source_doc_type: str | None = None
    output_path: str | None = None
    output_filename: str | None = None
    output_bytes: bytes | None = None
    updated_doc_ir: DocIR | None = None
    edits_applied: int = 0
    operations_applied: int = 0
    styles_applied: int = 0
    modified_target_ids: list[str] = Field(default_factory=list)
    created_target_ids: list[str] = Field(default_factory=list)
    removed_target_ids: list[str] = Field(default_factory=list)
    modified_run_ids: list[str] = Field(default_factory=list)
    edit_results: list[AppliedEditResult] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


def _text_hash(text: str | None) -> str | None:
    if text is None:
        return None
    return hashlib.sha1(text.encode("utf-8")).hexdigest()


class _EditableRunRef:
    def __init__(
        self,
        *,
        node_id: str,
        get_text: Callable[[], str],
        set_text: Callable[[str], None],
    ) -> None:
        self.node_id = node_id
        self._get_text = get_text
        self._set_text = set_text

    @property
    def text(self) -> str:
        return self._get_text()

    @text.setter
    def text(self, value: str) -> None:
        self._set_text(value)


class _EditableParagraphRef:
    def __init__(
        self,
        *,
        node_id: str,
        runs: list[_EditableRunRef],
        has_non_run_content: bool = False,
        recompute: Callable[[], None] | None = None,
    ) -> None:
        self.node_id = node_id
        self.runs = runs
        self.has_non_run_content = has_non_run_content
        self._recompute = recompute

    @property
    def text(self) -> str:
        return "".join(run.text for run in self.runs)

    def recompute(self) -> None:
        if self._recompute is not None:
            self._recompute()


class _EditableCellRef:
    def __init__(
        self,
        *,
        node_id: str,
        paragraphs: list[_EditableParagraphRef],
        recompute: Callable[[], None] | None = None,
    ) -> None:
        self.node_id = node_id
        self.paragraphs = paragraphs
        self._recompute = recompute

    @property
    def text(self) -> str:
        return "\n".join(paragraph.text for paragraph in self.paragraphs)

    def recompute(self) -> None:
        if self._recompute is not None:
            self._recompute()


class _EditableDocIndex:
    def __init__(
        self,
        *,
        paragraphs: dict[str, _EditableParagraphRef],
        runs: dict[str, _EditableRunRef],
        cells: dict[str, _EditableCellRef],
        run_to_paragraph: dict[str, _EditableParagraphRef],
    ) -> None:
        self.paragraphs = paragraphs
        self.runs = runs
        self.cells = cells
        self.run_to_paragraph = run_to_paragraph


class _RunSpan(BaseModel):
    start: int
    end: int
    full_start: int
    full_end: int
    run: _EditableRunRef

    model_config = {"arbitrary_types_allowed": True}


def _iter_doc_ir_paragraphs(paragraphs: list[ParagraphIR]):
    for paragraph in paragraphs:
        yield paragraph
        for table in paragraph.tables:
            yield from _iter_doc_ir_table_paragraphs(table)


def _iter_doc_ir_table_paragraphs(table: TableIR):
    for cell in table.iter_cells():
        for paragraph in cell.paragraphs:
            yield paragraph
            for nested_table in paragraph.tables:
                yield from _iter_doc_ir_table_paragraphs(nested_table)


def _build_doc_ir_index(doc: DocIR) -> _EditableDocIndex:
    doc.ensure_node_identity()
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    cells: dict[str, _EditableCellRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(
        paragraph: ParagraphIR,
        *,
        recompute_after: Callable[[], None] | None = None,
    ) -> _EditableParagraphRef:
        run_refs: list[_EditableRunRef] = []

        def recompute() -> None:
            paragraph.recompute_text()
            if recompute_after is not None:
                recompute_after()

        paragraph_ref = _EditableParagraphRef(
            node_id=paragraph.node_id,
            runs=run_refs,
            has_non_run_content=bool(paragraph.images or paragraph.tables),
            recompute=recompute,
        )
        paragraphs[paragraph.node_id] = paragraph_ref
        for run in paragraph.runs:
            run_ref = _EditableRunRef(
                node_id=run.node_id,
                get_text=lambda node=run: node.text,
                set_text=lambda value, node=run: setattr(node, "text", value),
            )
            run_refs.append(run_ref)
            runs[run.node_id] = run_ref
            run_to_paragraph[run.node_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table: TableIR, *, recompute_after: Callable[[], None] | None) -> None:
        for cell in table.iter_cells():
            cell_paragraph_refs: list[_EditableParagraphRef] = []

            def recompute_cell(node: TableCellIR = cell) -> None:
                node.recompute_text()
                if recompute_after is not None:
                    recompute_after()

            cell_ref = _EditableCellRef(
                node_id=cell.node_id,
                paragraphs=cell_paragraph_refs,
                recompute=recompute_cell,
            )
            cells[cell.node_id] = cell_ref

            for cell_paragraph in cell.paragraphs:
                paragraph_ref = register_paragraph(
                    cell_paragraph,
                    recompute_after=cell_ref.recompute,
                )
                cell_paragraph_refs.append(paragraph_ref)
                for nested_table in cell_paragraph.tables:
                    walk_table(nested_table, recompute_after=paragraph_ref.recompute)

    for paragraph in doc.paragraphs:
        paragraph_ref = register_paragraph(paragraph)
        for table in paragraph.tables:
            walk_table(table, recompute_after=paragraph_ref.recompute)

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, cells=cells, run_to_paragraph=run_to_paragraph)


def _iter_docx_blocks(doc):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    iter_inner_content = getattr(doc, "iter_inner_content", None)
    if callable(iter_inner_content):
        yield from iter_inner_content()
        return

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _iter_docx_blocks_from_element(parent, element):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _build_docx_index(doc) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    cells: dict[str, _EditableCellRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(paragraph, paragraph_path: str, *, has_non_run_content: bool = False) -> _EditableParagraphRef:
        run_refs: list[_EditableRunRef] = []
        paragraph_node_id = _anchored_node_id("paragraph", paragraph_path)
        paragraph_ref = _EditableParagraphRef(
            node_id=paragraph_node_id,
            runs=run_refs,
            has_non_run_content=has_non_run_content,
        )
        paragraphs[paragraph_node_id] = paragraph_ref
        for run_index, run in enumerate(paragraph.runs, start=1):
            run_path = f"{paragraph_path}.r{run_index}"
            run_node_id = _anchored_node_id("run", run_path)
            run_ref = _EditableRunRef(
                node_id=run_node_id,
                get_text=lambda node=run: node.text or "",
                set_text=lambda value, node=run: setattr(node, "text", value),
            )
            run_refs.append(run_ref)
            runs[run_node_id] = run_ref
            run_to_paragraph[run_node_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table, table_base: str) -> None:
        for tr_idx, row in enumerate(table.rows, start=1):
            for tc_idx, cell in enumerate(row.cells, start=1):
                cell_path = f"{table_base}.tr{tr_idx}.tc{tc_idx}"
                cell_node_id = _anchored_node_id("cell", cell_path)
                cell_paragraph_refs: list[_EditableParagraphRef] = []
                cells[cell_node_id] = _EditableCellRef(node_id=cell_node_id, paragraphs=cell_paragraph_refs)
                cp_idx = 0
                current_paragraph_path: str | None = None
                nested_table_counter_by_paragraph: dict[str, int] = {}

                for block in _iter_docx_blocks_from_element(cell, cell._tc):
                    if block.__class__.__name__ == "Paragraph":
                        cp_idx += 1
                        current_paragraph_path = f"{cell_path}.p{cp_idx}"
                        cell_paragraph_refs.append(register_paragraph(block, current_paragraph_path))
                        continue

                    if block.__class__.__name__ != "Table":
                        continue

                    if current_paragraph_path is None:
                        cp_idx += 1
                        current_paragraph_path = f"{cell_path}.p{cp_idx}"
                        paragraph_node_id = _anchored_node_id("paragraph", current_paragraph_path)
                        paragraph_ref = _EditableParagraphRef(
                            node_id=paragraph_node_id,
                            runs=[],
                            has_non_run_content=True,
                        )
                        paragraphs[paragraph_node_id] = paragraph_ref
                        cell_paragraph_refs.append(paragraph_ref)
                    else:
                        paragraphs[_anchored_node_id("paragraph", current_paragraph_path)].has_non_run_content = True

                    tbl_counter = nested_table_counter_by_paragraph.get(current_paragraph_path, 0) + 1
                    nested_table_counter_by_paragraph[current_paragraph_path] = tbl_counter
                    nested_table_base = f"{current_paragraph_path}.tbl{tbl_counter}"
                    walk_table(block, nested_table_base)

    p_idx = 0
    tbl_counter = 0
    for block in _iter_docx_blocks(doc):
        if block.__class__.__name__ == "Paragraph":
            p_idx += 1
            register_paragraph(block, f"s1.p{p_idx}")
            continue

        if block.__class__.__name__ != "Table":
            continue

        tbl_counter += 1
        p_idx += 1
        walk_table(block, f"s1.p{p_idx}.r1.tbl{tbl_counter}")

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, cells=cells, run_to_paragraph=run_to_paragraph)


_SECTION_NAME_RE = re.compile(r"^Contents/section(\d+)\.xml$")
_HEADER_NAME = "Contents/header.xml"
_HH_NS = "http://www.hancom.co.kr/hwpml/2011/head"
_HP_NS = "http://www.hancom.co.kr/hwpml/2011/paragraph"
_HC_NS = "http://www.hancom.co.kr/hwpml/2011/core"
_HH = f"{{{_HH_NS}}}"
_HP = f"{{{_HP_NS}}}"
_HC = f"{{{_HC_NS}}}"
_XML_PREFIX_AND_ROOT_RE = re.compile(
    rb"^(?P<prefix>\s*(?:<\?xml[^>]*\?>\s*)?)(?P<root_open><[^!?][^>]*>)",
    re.DOTALL,
)
_XML_ROOT_NAME_RE = re.compile(rb"<(?P<name>[^\s>/]+)")
_DOCX_DEFAULT_TABLE_WIDTH_TWIPS = 8640
_DOCX_MIN_COLUMN_WIDTH_TWIPS = 1440
_DOCX_DEFAULT_CELL_MARGIN_TWIPS = 120
_DOCX_DEFAULT_BORDER_SIZE = "6"
_HWPX_DEFAULT_TABLE_WIDTH = 36000
_HWPX_MIN_CELL_WIDTH = 6000
_HWPX_DEFAULT_CELL_HEIGHT = 1800
_HWPX_DEFAULT_CELL_MARGIN_X = 510
_HWPX_DEFAULT_CELL_MARGIN_Y = 141
_EMU_PER_PT = 12700.0
_TWIPS_PER_PT = 20.0
_HWPUNIT_PER_PT = 100.0
_CSS_PX_PER_MM = 3.78
_HWPX_LINE_WIDTHS_MM = (
    (0.1, "0.1 mm"),
    (0.12, "0.12 mm"),
    (0.15, "0.15 mm"),
    (0.2, "0.2 mm"),
    (0.25, "0.25 mm"),
    (0.3, "0.3 mm"),
    (0.4, "0.4 mm"),
    (0.5, "0.5 mm"),
    (0.6, "0.6 mm"),
    (0.7, "0.7 mm"),
    (1.0, "1.0 mm"),
    (1.5, "1.5 mm"),
    (2.0, "2.0 mm"),
    (3.0, "3.0 mm"),
    (4.0, "4.0 mm"),
    (5.0, "5.0 mm"),
)


def _pt_to_twips(value: float | int | None) -> int | None:
    if value is None:
        return None
    return int(round(float(value) * _TWIPS_PER_PT))


def _pt_to_emu(value: float | int | None) -> int | None:
    if value is None:
        return None
    return int(round(float(value) * _EMU_PER_PT))


def _pt_to_hwpunit(value: float | int | None) -> int | None:
    if value is None:
        return None
    return int(round(float(value) * _HWPUNIT_PER_PT))


def _hwpx_line_width_value(width_mm: float) -> str:
    _distance, width = min(
        (abs(width_mm - candidate_mm), candidate_width)
        for candidate_mm, candidate_width in _HWPX_LINE_WIDTHS_MM
    )
    return width


@dataclass
class _EditableHwpxSection:
    name: str
    root: ET.Element
    xml_prefix: bytes
    original_root_open: bytes
    namespaces: list[tuple[str, str]]


@dataclass
class _EditableHwpxHeader:
    name: str
    root: ET.Element
    xml_prefix: bytes
    original_root_open: bytes
    namespaces: list[tuple[str, str]]


def _run_text(run_el: ET.Element) -> str:
    return "".join("".join(node.itertext()) for node in run_el.findall(f"{_HP}t"))


def _hwpx_paragraph_visible_text(paragraph_el: ET.Element) -> str:
    return "".join(_run_text(run_el) for run_el in paragraph_el.findall(f"{_HP}run"))


def _hwpx_cell_visible_text(cell_el: ET.Element) -> str:
    return "\n".join(_hwpx_paragraph_visible_text(paragraph_el) for paragraph_el in _iter_cell_paragraphs(cell_el))


def _set_hwpx_run_text(run_el: ET.Element, new_text: str) -> None:
    for node in list(run_el):
        if node.tag == f"{_HP}t":
            run_el.remove(node)
    text_el = ET.SubElement(run_el, f"{_HP}t")
    text_el.text = new_text


def _iter_section_paragraphs(section_root: ET.Element) -> list[ET.Element]:
    return section_root.findall(f"{_HP}p")


def _iter_paragraph_tables(paragraph_el: ET.Element) -> list[ET.Element]:
    return paragraph_el.findall(f"{_HP}run/{_HP}tbl")


def _iter_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    direct = cell_el.findall(f"{_HP}subList/{_HP}p")
    if direct:
        return direct
    return cell_el.findall(f".//{_HP}p")


def _safe_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _logical_table_cells(row_el: ET.Element) -> list[tuple[int, ET.Element]]:
    logical_cells: list[tuple[int, ET.Element]] = []
    fallback_col = 1
    for cell_el in row_el.findall(f"{_HP}tc"):
        cell_addr = cell_el.find(f"{_HP}cellAddr")
        col_addr = _safe_int(cell_addr.get("colAddr")) if cell_addr is not None else None
        logical_col = (col_addr + 1) if col_addr is not None else fallback_col
        logical_cells.append((logical_col, cell_el))

        cell_span = cell_el.find(f"{_HP}cellSpan")
        colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
        fallback_col = max(fallback_col, logical_col + max(colspan or 1, 1))
    return logical_cells


def _collect_xml_namespaces(xml_bytes: bytes) -> list[tuple[str, str]]:
    namespaces: list[tuple[str, str]] = []
    for _event, item in ET.iterparse(BytesIO(xml_bytes), events=("start-ns",)):
        if item not in namespaces:
            namespaces.append(item)
    return namespaces


def _split_xml_prefix_and_root_open(xml_bytes: bytes) -> tuple[bytes, bytes] | None:
    match = _XML_PREFIX_AND_ROOT_RE.match(xml_bytes)
    if match is None:
        return None
    return match.group("prefix"), match.group("root_open")


def _root_open_start_tag(root_open: bytes) -> bytes:
    stripped = root_open.rstrip()
    if not stripped.endswith(b"/>"):
        return root_open
    trailing = root_open[len(stripped) :]
    return stripped[:-2].rstrip() + b">" + trailing


def _root_close_tag(root_open: bytes) -> bytes:
    match = _XML_ROOT_NAME_RE.search(root_open)
    if match is None:
        return b""
    return b"</" + match.group("name") + b">"


def _serialize_hwpx_xml_part(part: _EditableHwpxSection | _EditableHwpxHeader) -> bytes:
    for prefix, uri in part.namespaces:
        ET.register_namespace(prefix, uri)

    serialized = ET.tostring(part.root, encoding="utf-8", xml_declaration=False)
    generated_parts = _split_xml_prefix_and_root_open(serialized)
    if generated_parts is None:
        return serialized if not part.xml_prefix else part.xml_prefix + serialized

    _generated_prefix, generated_root_open = generated_parts
    generated_body = serialized[len(generated_root_open) :]
    original_start = _root_open_start_tag(part.original_root_open)
    if not generated_body and generated_root_open.rstrip().endswith(b"/>"):
        return part.xml_prefix + original_start + _root_close_tag(part.original_root_open)
    return part.xml_prefix + original_start + generated_body


def _serialize_hwpx_section(section: _EditableHwpxSection) -> bytes:
    return _serialize_hwpx_xml_part(section)


class _EditableHwpxArchive:
    def __init__(
        self,
        *,
        source_path: Path,
        source_bytes: bytes,
        section_entries: list[_EditableHwpxSection],
        header_entry: _EditableHwpxHeader | None = None,
    ) -> None:
        self.source_path = source_path
        self.source_bytes = source_bytes
        self.section_entries = section_entries
        self.header_entry = header_entry

    @staticmethod
    def _load_section_entries(source_bytes: bytes) -> list[_EditableHwpxSection]:
        with zipfile.ZipFile(BytesIO(source_bytes)) as archive:
            return sorted(
                [
                    _EditableHwpxSection(
                        name=name,
                        root=ET.fromstring(section_bytes := archive.read(name)),
                        xml_prefix=(split_parts[0] if (split_parts := _split_xml_prefix_and_root_open(section_bytes)) else b""),
                        original_root_open=(split_parts[1] if split_parts else b""),
                        namespaces=_collect_xml_namespaces(section_bytes),
                    )
                    for name in archive.namelist()
                    if _SECTION_NAME_RE.match(name)
                ],
                key=lambda item: int(_SECTION_NAME_RE.match(item.name).group(1)),
            )

    @staticmethod
    def _load_header_entry(source_bytes: bytes) -> _EditableHwpxHeader | None:
        with zipfile.ZipFile(BytesIO(source_bytes)) as archive:
            try:
                header_bytes = archive.read(_HEADER_NAME)
            except KeyError:
                return None
        split_parts = _split_xml_prefix_and_root_open(header_bytes)
        return _EditableHwpxHeader(
            name=_HEADER_NAME,
            root=ET.fromstring(header_bytes),
            xml_prefix=(split_parts[0] if split_parts else b""),
            original_root_open=(split_parts[1] if split_parts else b""),
            namespaces=_collect_xml_namespaces(header_bytes),
        )

    @classmethod
    def open(cls, source_path: str | Path) -> "_EditableHwpxArchive":
        path = Path(source_path)
        source_bytes = path.read_bytes()
        section_entries = cls._load_section_entries(source_bytes)
        header_entry = cls._load_header_entry(source_bytes)
        return cls(source_path=path, source_bytes=source_bytes, section_entries=section_entries, header_entry=header_entry)

    @classmethod
    def from_bytes(
        cls,
        source_bytes: bytes,
        *,
        source_path: str | Path | None = None,
    ) -> "_EditableHwpxArchive":
        path = Path(source_path) if source_path is not None else Path("converted.hwpx")
        section_entries = cls._load_section_entries(source_bytes)
        header_entry = cls._load_header_entry(source_bytes)
        return cls(source_path=path, source_bytes=source_bytes, section_entries=section_entries, header_entry=header_entry)

    def write_to(self, output_path: str | Path) -> None:
        section_bytes = {
            section.name: _serialize_hwpx_section(section)
            for section in self.section_entries
        }
        replacement_bytes = dict(section_bytes)
        if self.header_entry is not None:
            replacement_bytes[self.header_entry.name] = _serialize_hwpx_xml_part(self.header_entry)

        output = Path(output_path)
        with zipfile.ZipFile(BytesIO(self.source_bytes), "r") as source_archive:
            with zipfile.ZipFile(output, "w") as target_archive:
                for info in source_archive.infolist():
                    data = replacement_bytes.get(info.filename, source_archive.read(info.filename))
                    target_archive.writestr(info, data)


def _build_hwpx_index(archive: _EditableHwpxArchive) -> _EditableDocIndex:
    paragraphs: dict[str, _EditableParagraphRef] = {}
    runs: dict[str, _EditableRunRef] = {}
    cells: dict[str, _EditableCellRef] = {}
    run_to_paragraph: dict[str, _EditableParagraphRef] = {}

    def register_paragraph(paragraph_el: ET.Element, paragraph_path: str) -> _EditableParagraphRef:
        run_elements = paragraph_el.findall(f"{_HP}run")
        paragraph_node_id = _anchored_node_id("paragraph", paragraph_path)
        paragraph_ref = _EditableParagraphRef(
            node_id=paragraph_node_id,
            runs=[],
            has_non_run_content=bool(_iter_paragraph_tables(paragraph_el)),
        )
        paragraphs[paragraph_node_id] = paragraph_ref
        if not run_elements:
            return paragraph_ref

        for run_index, run_el in enumerate(run_elements, start=1):
            run_path = f"{paragraph_path}.r{run_index}"
            run_node_id = _anchored_node_id("run", run_path)
            run_ref = _EditableRunRef(
                node_id=run_node_id,
                get_text=lambda node=run_el: _run_text(node),
                set_text=lambda value, node=run_el: _set_hwpx_run_text(node, value),
            )
            paragraph_ref.runs.append(run_ref)
            runs[run_node_id] = run_ref
            run_to_paragraph[run_node_id] = paragraph_ref
        return paragraph_ref

    def walk_table(table_el: ET.Element, table_base: str) -> None:
        for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
            for tc_idx, cell_el in _logical_table_cells(row_el):
                cell_path = f"{table_base}.tr{tr_idx}.tc{tc_idx}"
                cell_node_id = _anchored_node_id("cell", cell_path)
                cell_paragraph_refs: list[_EditableParagraphRef] = []
                cells[cell_node_id] = _EditableCellRef(node_id=cell_node_id, paragraphs=cell_paragraph_refs)
                cell_paragraphs = _iter_cell_paragraphs(cell_el)
                if not cell_paragraphs:
                    paragraph_path = f"{cell_path}.p1"
                    paragraph_node_id = _anchored_node_id("paragraph", paragraph_path)
                    paragraph_ref = _EditableParagraphRef(
                        node_id=paragraph_node_id,
                        runs=[],
                        has_non_run_content=False,
                    )
                    paragraphs[paragraph_node_id] = paragraph_ref
                    cell_paragraph_refs.append(paragraph_ref)
                    continue

                for cp_idx, paragraph_el in enumerate(cell_paragraphs, start=1):
                    paragraph_path = f"{cell_path}.p{cp_idx}"
                    paragraph_ref = register_paragraph(paragraph_el, paragraph_path)
                    cell_paragraph_refs.append(paragraph_ref)
                    for nested_index, nested_table in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                        paragraph_ref.has_non_run_content = True
                        walk_table(nested_table, f"{paragraph_path}.tbl{nested_index}")

    for section_index, section in enumerate(archive.section_entries, start=1):
        for paragraph_index, paragraph_el in enumerate(_iter_section_paragraphs(section.root), start=1):
            paragraph_id = f"s{section_index}.p{paragraph_index}"
            paragraph_ref = register_paragraph(paragraph_el, paragraph_id)
            for table_index, table_el in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                paragraph_ref.has_non_run_content = True
                walk_table(table_el, f"{paragraph_id}.r1.tbl{table_index}")

    return _EditableDocIndex(paragraphs=paragraphs, runs=runs, cells=cells, run_to_paragraph=run_to_paragraph)


def _build_run_spans(paragraph: _EditableParagraphRef) -> list[_RunSpan]:
    spans: list[_RunSpan] = []
    cursor = 0
    for run in paragraph.runs:
        length = len(run.text)
        spans.append(
            _RunSpan(
                start=cursor,
                end=cursor + length,
                full_start=cursor,
                full_end=cursor + length,
                run=run,
            )
        )
        cursor += length
    return spans


def _clip_run_spans(spans: list[_RunSpan], i1: int, i2: int) -> list[_RunSpan]:
    is_insert = i1 == i2
    clipped: list[_RunSpan] = []
    for span in spans:
        if is_insert:
            if span.end < i1 or span.start > i2:
                continue
        else:
            if span.end <= i1 or span.start >= i2:
                continue
        clipped.append(
            _RunSpan(
                start=span.start if is_insert else max(span.start, i1),
                end=span.end if is_insert else min(span.end, i2),
                full_start=span.full_start,
                full_end=span.full_end,
                run=span.run,
            )
        )
    if is_insert and len(clipped) > 1:
        preceding = [span for span in clipped if span.end == i1]
        if preceding:
            clipped = preceding[:1]
    return clipped


def _append_unique(values: list[str], value: str) -> None:
    if value not in values:
        values.append(value)


def _apply_to_run(run: _EditableRunRef, new_text: str, local_start: int, local_end: int) -> None:
    current = run.text
    run.text = current[:local_start] + new_text + current[local_end:]


def _apply_to_run_with_offsets(
    run: _EditableRunRef,
    new_text: str,
    local_start: int,
    local_end: int,
    offset_deltas: dict[str, int],
) -> None:
    delta = offset_deltas.get(run.node_id, 0)
    actual_start = local_start + delta
    actual_end = local_end + delta
    _apply_to_run(run, new_text, actual_start, actual_end)
    offset_deltas[run.node_id] = delta + len(new_text) - (local_end - local_start)


def _apply_multi_run(
    orig_sub: str,
    new_sub: str,
    spans: list[_RunSpan],
    result: _EditEngineResult,
    offset_deltas: dict[str, int],
    *,
    base_offset: int,
    depth: int = 0,
) -> None:
    if depth == 0:
        matcher = difflib.SequenceMatcher(None, orig_sub, new_sub, autojunk=False)
        for tag, a1, a2, b1, b2 in matcher.get_opcodes():
            if tag == "equal":
                continue
            abs_a1 = base_offset + a1
            abs_a2 = base_offset + a2
            sub_spans = _clip_run_spans(spans, abs_a1, abs_a2)
            if not sub_spans:
                continue
            if len(sub_spans) == 1:
                span = sub_spans[0]
                local_start = abs_a1 - span.full_start
                local_end = abs_a2 - span.full_start
                _apply_to_run_with_offsets(span.run, new_sub[b1:b2], local_start, local_end, offset_deltas)
                _append_unique(result.modified_run_ids, span.run.node_id)
                continue
            _apply_multi_run(
                orig_sub[a1:a2],
                new_sub[b1:b2],
                sub_spans,
                result,
                offset_deltas,
                base_offset=abs_a1,
                depth=1,
            )
        return

    first = spans[0]
    _apply_to_run_with_offsets(
        first.run,
        new_sub,
        first.start - first.full_start,
        first.end - first.full_start,
        offset_deltas,
    )
    _append_unique(result.modified_run_ids, first.run.node_id)
    for span in spans[1:]:
        _apply_to_run_with_offsets(
            span.run,
            "",
            span.start - span.full_start,
            span.end - span.full_start,
            offset_deltas,
        )
        _append_unique(result.modified_run_ids, span.run.node_id)
    result.warnings.append(
        "Multi-run fallback used for "
        f"{[span.run.node_id for span in spans]}: all replacement text assigned to {first.run.node_id}"
    )


def _validate_paragraph_edit(index: _EditableDocIndex, edit: TextEdit) -> _EditableParagraphRef:
    paragraph = index.paragraphs.get(edit.target_id)
    if paragraph is None:
        raise EditValidationError(f"Paragraph does not exist: {edit.target_id}")
    if paragraph.has_non_run_content:
        raise EditValidationError(
            f"Paragraph edit targets unsupported mixed content (tables/images): {edit.target_id}"
        )
    current_hash = _text_hash(paragraph.text)
    if current_hash != edit.expected_text_hash:
        raise EditValidationError(
            f"Paragraph text hash mismatch for {edit.target_id}.",
            code="text_hash_mismatch",
            target_kind="paragraph",
            target_id=edit.target_id,
            expected_text_hash=edit.expected_text_hash,
            current_text_hash=current_hash,
            current_text=paragraph.text,
        )
    return paragraph


def _validate_run_edit(index: _EditableDocIndex, edit: TextEdit) -> _EditableRunRef:
    run = index.runs.get(edit.target_id)
    if run is None:
        raise EditValidationError(f"Run does not exist: {edit.target_id}")
    current_hash = _text_hash(run.text)
    if current_hash != edit.expected_text_hash:
        raise EditValidationError(
            f"Run text hash mismatch for {edit.target_id}.",
            code="text_hash_mismatch",
            target_kind="run",
            target_id=edit.target_id,
            expected_text_hash=edit.expected_text_hash,
            current_text_hash=current_hash,
            current_text=run.text,
        )
    return run


def _validate_cell_edit(index: _EditableDocIndex, edit: TextEdit) -> _EditableCellRef:
    cell = index.cells.get(edit.target_id)
    if cell is None:
        raise EditValidationError(f"Cell does not exist: {edit.target_id}")
    if any(paragraph.has_non_run_content for paragraph in cell.paragraphs):
        raise EditValidationError(
            f"Cell edit targets unsupported mixed content (nested tables/images): {edit.target_id}"
        )
    if not cell.paragraphs or any(not paragraph.runs for paragraph in cell.paragraphs):
        raise EditValidationError(f"Cell does not contain editable text runs: {edit.target_id}")
    current_hash = _text_hash(cell.text)
    if current_hash != edit.expected_text_hash:
        raise EditValidationError(
            f"Cell text hash mismatch for {edit.target_id}.",
            code="text_hash_mismatch",
            target_kind="cell",
            target_id=edit.target_id,
            expected_text_hash=edit.expected_text_hash,
            current_text_hash=current_hash,
            current_text=cell.text,
        )
    expected_paragraphs = len(cell.paragraphs)
    new_paragraphs = len(edit.new_text.split("\n"))
    if new_paragraphs != expected_paragraphs:
        raise EditValidationError(
            f"Cell text replacement for {edit.target_id} must preserve paragraph count: "
            f"expected {expected_paragraphs} line(s), got {new_paragraphs}."
        )
    return cell


def _replace_paragraph_text(
    paragraph: _EditableParagraphRef,
    new_text: str,
    result: _EditEngineResult,
) -> None:
    spans = _build_run_spans(paragraph)
    original = paragraph.text
    if len(spans) == 1:
        run = spans[0].run
        run.text = new_text
        _append_unique(result.modified_run_ids, run.node_id)
        paragraph.recompute()
        return

    offset_deltas: dict[str, int] = {}
    matcher = difflib.SequenceMatcher(None, original, new_text, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        affected = _clip_run_spans(spans, i1, i2)
        if not affected:
            continue
        if len(affected) == 1:
            span = affected[0]
            local_start = i1 - span.full_start
            local_end = i2 - span.full_start
            _apply_to_run_with_offsets(span.run, new_text[j1:j2], local_start, local_end, offset_deltas)
            _append_unique(result.modified_run_ids, span.run.node_id)
            continue
        _apply_multi_run(
            original[i1:i2],
            new_text[j1:j2],
            affected,
            result,
            offset_deltas,
            base_offset=i1,
        )
    paragraph.recompute()


def _apply_single_edit(index: _EditableDocIndex, edit: TextEdit, result: _EditEngineResult) -> None:
    if edit.target_kind == "run":
        run = _validate_run_edit(index, edit)
        run.text = edit.new_text
        paragraph = index.run_to_paragraph.get(edit.target_id)
        if paragraph is not None:
            paragraph.recompute()
        _append_unique(result.modified_target_ids, edit.target_id)
        _append_unique(result.modified_run_ids, edit.target_id)
        result.edits_applied += 1
        return

    if edit.target_kind == "paragraph":
        paragraph = _validate_paragraph_edit(index, edit)
        _replace_paragraph_text(paragraph, edit.new_text, result)
        _append_unique(result.modified_target_ids, edit.target_id)
        result.edits_applied += 1
        return

    if edit.target_kind != "cell":
        raise EditValidationError(f"Unsupported edit target kind: {edit.target_kind!r}")

    cell = _validate_cell_edit(index, edit)
    for paragraph, new_paragraph_text in zip(cell.paragraphs, edit.new_text.split("\n"), strict=True):
        _replace_paragraph_text(paragraph, new_paragraph_text, result)
    cell.recompute()
    _append_unique(result.modified_target_ids, edit.target_id)
    result.edits_applied += 1


def _apply_text_edits_to_doc_ir(doc: DocIR, edits: list[TextEdit]) -> _EditEngineResult:
    updated = doc.model_copy(deep=True)
    index = _build_doc_ir_index(updated)
    result = _EditEngineResult(source_doc_type=updated.source_doc_type)
    for edit in edits:
        _apply_single_edit(index, edit, result)
    result.updated_doc_ir = updated
    return result


@dataclass
class _DocIrParagraphLocation:
    node: ParagraphIR
    container: list[ParagraphIR]
    index: int
    parent_cell: TableCellIR | None = None


@dataclass
class _DocIrRunLocation:
    node: RunIR
    paragraph: ParagraphIR
    content_index: int


@dataclass
class _DocIrTableLocation:
    node: TableIR
    paragraph: ParagraphIR
    content_index: int


@dataclass
class _DocIrImageLocation:
    node: ImageIR
    paragraph: ParagraphIR
    content_index: int


@dataclass
class _DocIrCellLocation:
    node: TableCellIR
    table: TableIR


class _StructuralDocIrIndex:
    def __init__(self) -> None:
        self.paragraphs: dict[str, _DocIrParagraphLocation] = {}
        self.runs: dict[str, _DocIrRunLocation] = {}
        self.tables: dict[str, _DocIrTableLocation] = {}
        self.images: dict[str, _DocIrImageLocation] = {}
        self.cells: dict[str, _DocIrCellLocation] = {}


def _build_structural_doc_ir_index(doc: DocIR) -> _StructuralDocIrIndex:
    doc.ensure_node_identity()
    index = _StructuralDocIrIndex()

    def walk_paragraphs(container: list[ParagraphIR], *, parent_cell: TableCellIR | None = None) -> None:
        for paragraph_index, paragraph in enumerate(container):
            index.paragraphs[paragraph.node_id] = _DocIrParagraphLocation(
                node=paragraph,
                container=container,
                index=paragraph_index,
                parent_cell=parent_cell,
            )
            for content_index, item in enumerate(paragraph.content):
                if isinstance(item, RunIR):
                    index.runs[item.node_id] = _DocIrRunLocation(
                        node=item,
                        paragraph=paragraph,
                        content_index=content_index,
                    )
                elif isinstance(item, ImageIR):
                    index.images[item.node_id] = _DocIrImageLocation(
                        node=item,
                        paragraph=paragraph,
                        content_index=content_index,
                    )
                elif isinstance(item, TableIR):
                    index.tables[item.node_id] = _DocIrTableLocation(
                        node=item,
                        paragraph=paragraph,
                        content_index=content_index,
                    )
                    walk_table(item)

    def walk_table(table: TableIR) -> None:
        for cell in table.iter_cells():
            index.cells[cell.node_id] = _DocIrCellLocation(node=cell, table=table)
            walk_paragraphs(cell.paragraphs, parent_cell=cell)

    walk_paragraphs(doc.paragraphs)
    return index


_RUN_STYLE_FIELD_MAP = {
    "bold": "bold",
    "italic": "italic",
    "underline": "underline",
    "strikethrough": "strikethrough",
    "superscript": "superscript",
    "subscript": "subscript",
    "color": "color",
    "highlight": "highlight",
    "font_size_pt": "size_pt",
}
_PARA_STYLE_FIELD_MAP = {
    "paragraph_align": "align",
    "left_indent_pt": "left_indent_pt",
    "right_indent_pt": "right_indent_pt",
    "first_line_indent_pt": "first_line_indent_pt",
    "hanging_indent_pt": "hanging_indent_pt",
}
_CELL_STYLE_FIELDS = {
    "background",
    "vertical_align",
    "horizontal_align",
    "width_pt",
    "height_pt",
    "padding_top_pt",
    "padding_right_pt",
    "padding_bottom_pt",
    "padding_left_pt",
    "border_top",
    "border_right",
    "border_bottom",
    "border_left",
}
_CELL_DIMENSION_FIELDS = {"width_pt", "height_pt"}
_CELL_DIRECT_STYLE_FIELDS = _CELL_STYLE_FIELDS - _CELL_DIMENSION_FIELDS
_PLACEMENT_FIELD_MAP = {
    "placement_mode": "mode",
    "wrap": "wrap",
    "text_flow": "text_flow",
    "x_relative_to": "x_relative_to",
    "y_relative_to": "y_relative_to",
    "x_align": "x_align",
    "y_align": "y_align",
    "x_offset_pt": "x_offset_pt",
    "y_offset_pt": "y_offset_pt",
    "margin_top_pt": "margin_top_pt",
    "margin_right_pt": "margin_right_pt",
    "margin_bottom_pt": "margin_bottom_pt",
    "margin_left_pt": "margin_left_pt",
    "allow_overlap": "allow_overlap",
    "flow_with_text": "flow_with_text",
    "z_order": "z_order",
}
_STYLE_EDIT_FIELDS = set(_RUN_STYLE_FIELD_MAP) | set(_PARA_STYLE_FIELD_MAP) | _CELL_STYLE_FIELDS | {"width_pt", "height_pt"} | set(_PLACEMENT_FIELD_MAP)


def _style_edit_supplied_fields(edit: StyleEdit) -> set[str]:
    return {field for field in _STYLE_EDIT_FIELDS if getattr(edit, field) is not None} | set(edit.clear_fields)


def _set_or_clear_field(obj: object, attr: str, value: object, *, clear: bool = False, clear_value: object = None) -> None:
    setattr(obj, attr, clear_value if clear else value)


def _apply_placement_edit(placement: ObjectPlacementInfo | None, edit: StyleEdit) -> ObjectPlacementInfo | None:
    if placement is None:
        placement = ObjectPlacementInfo()
    for field_name, attr in _PLACEMENT_FIELD_MAP.items():
        if field_name in edit.clear_fields:
            _set_or_clear_field(placement, attr, None, clear=True)
            continue
        value = getattr(edit, field_name)
        if value is not None:
            setattr(placement, attr, value)
    if not placement.model_dump(exclude_none=True):
        return None
    return placement


def _style_target_kind_for_doc_ir_index(index: _StructuralDocIrIndex, target_id: str) -> str | None:
    if target_id in index.paragraphs:
        return "paragraph"
    if target_id in index.runs:
        return "run"
    if target_id in index.cells:
        return "cell"
    if target_id in index.tables:
        return "table"
    if target_id in index.images:
        return "image"
    return None


def _cell_style_span(cell: TableCellIR, attr: str) -> int:
    style = cell.cell_style
    value = getattr(style, attr, 1) if style is not None else 1
    return max(value or 1, 1)


def _cell_covers_column(cell: TableCellIR, cell_col_index: int, target_col_index: int) -> bool:
    return cell_col_index <= target_col_index < cell_col_index + _cell_style_span(cell, "colspan")


def _cell_covers_row(cell: TableCellIR, cell_row_index: int, target_row_index: int) -> bool:
    return cell_row_index <= target_row_index < cell_row_index + _cell_style_span(cell, "rowspan")


def _apply_cell_style_fields(cell: TableCellIR, edit: StyleEdit, fields: set[str]) -> None:
    style = cell.cell_style or CellStyleInfo()
    for field_name in fields:
        if field_name in edit.clear_fields:
            _set_or_clear_field(style, field_name, None, clear=True)
            continue
        value = getattr(edit, field_name)
        if value is not None:
            setattr(style, field_name, value)
    cell.cell_style = style


def _apply_style_edit_to_doc_ir_index(index: _StructuralDocIrIndex, edit: StyleEdit, result: _EditEngineResult) -> None:
    actual_kind = _style_target_kind_for_doc_ir_index(index, edit.target_id)
    if actual_kind is None:
        raise EditValidationError(
            f"Style target does not exist: {edit.target_id}.",
            code="target_not_found",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )
    if actual_kind != edit.target_kind:
        raise EditValidationError(
            f"{edit.target_id} is a {actual_kind} target, not a {edit.target_kind} target.",
            code="target_kind_mismatch",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )

    if edit.target_kind == "run":
        run = index.runs[edit.target_id].node
        style = run.run_style or RunStyleInfo()
        for field_name, attr in _RUN_STYLE_FIELD_MAP.items():
            if field_name in edit.clear_fields:
                clear_value = False if attr in {"bold", "italic", "underline", "strikethrough", "superscript", "subscript"} else None
                _set_or_clear_field(style, attr, None, clear=True, clear_value=clear_value)
                continue
            value = getattr(edit, field_name)
            if value is not None:
                setattr(style, attr, value)
        run.run_style = style
    elif edit.target_kind == "paragraph":
        paragraph = index.paragraphs[edit.target_id].node
        style = paragraph.para_style or ParaStyleInfo()
        for field_name, attr in _PARA_STYLE_FIELD_MAP.items():
            if field_name in edit.clear_fields:
                _set_or_clear_field(style, attr, None, clear=True)
                continue
            value = getattr(edit, field_name)
            if value is not None:
                setattr(style, attr, value)
        paragraph.para_style = style
    elif edit.target_kind == "cell":
        location = index.cells[edit.target_id]
        cell = location.node
        _apply_cell_style_fields(cell, edit, _CELL_DIRECT_STYLE_FIELDS)
        cell_position = location.table.cell_position(cell)

        if cell_position is not None and "width_pt" in _style_edit_supplied_fields(edit):
            _row_index, target_col_index = cell_position
            for _table_row_index, table_col_index, table_cell in location.table.iter_cell_positions():
                if _cell_covers_column(table_cell, table_col_index, target_col_index):
                    _apply_cell_style_fields(table_cell, edit, {"width_pt"})
                    _append_unique(result.modified_target_ids, table_cell.node_id)

        if cell_position is not None and "height_pt" in _style_edit_supplied_fields(edit):
            target_row_index, _col_index = cell_position
            for table_row_index, _table_col_index, table_cell in location.table.iter_cell_positions():
                if _cell_covers_row(table_cell, table_row_index, target_row_index):
                    _apply_cell_style_fields(table_cell, edit, {"height_pt"})
                    _append_unique(result.modified_target_ids, table_cell.node_id)
    elif edit.target_kind == "table":
        table = index.tables[edit.target_id].node
        style = table.table_style or TableStyleInfo(row_count=table.row_count, col_count=table.col_count)
        style.placement = _apply_placement_edit(style.placement, edit)
        table.table_style = style
    elif edit.target_kind == "image":
        image = index.images[edit.target_id].node
        if "width_pt" in edit.clear_fields:
            image.display_width_pt = None
        elif edit.width_pt is not None:
            image.display_width_pt = edit.width_pt
        if "height_pt" in edit.clear_fields:
            image.display_height_pt = None
        elif edit.height_pt is not None:
            image.display_height_pt = edit.height_pt
        image.placement = _apply_placement_edit(image.placement, edit)
    else:
        raise EditValidationError(
            f"Unsupported style target kind: {edit.target_kind!r}.",
            code="invalid_style",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )

    _append_unique(result.modified_target_ids, edit.target_id)
    result.styles_applied += 1


def _apply_style_edits_to_doc_ir(doc: DocIR, edits: list[StyleEdit]) -> _EditEngineResult:
    updated = doc.model_copy(deep=True)
    index = _build_structural_doc_ir_index(updated)
    result = _EditEngineResult(source_doc_type=updated.source_doc_type)
    for edit in edits:
        _apply_style_edit_to_doc_ir_index(index, edit, result)
    result.updated_doc_ir = updated
    return result


def _text_digest(text: str | None) -> str:
    return hashlib.sha1((text or "").encode("utf-8")).hexdigest()[:12]


def _new_inserted_node_id(
    kind: NodeKind,
    seed: str,
    existing_ids: set[str],
) -> str:
    counter = 1
    while True:
        candidate = _anchored_node_id(kind, f"inserted.{seed}.{counter}")
        if candidate not in existing_ids:
            existing_ids.add(candidate)
            return candidate
        counter += 1


def _all_doc_ir_node_ids(doc: DocIR) -> set[str]:
    ids: set[str] = set()

    def walk_paragraph(paragraph: ParagraphIR) -> None:
        if paragraph.node_id:
            ids.add(paragraph.node_id)
        for item in paragraph.content:
            if isinstance(item, RunIR):
                if item.node_id:
                    ids.add(item.node_id)
            elif isinstance(item, ImageIR):
                if item.node_id:
                    ids.add(item.node_id)
            elif isinstance(item, TableIR):
                walk_table(item)

    def walk_table(table: TableIR) -> None:
        if table.node_id:
            ids.add(table.node_id)
        for cell in table.iter_cells():
            if cell.node_id:
                ids.add(cell.node_id)
            for paragraph in cell.paragraphs:
                walk_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        walk_paragraph(paragraph)
    return ids


def _inserted_anchor(kind: NodeKind, seed: str, *, source_doc_type: str | None, text: str | None = None) -> NativeAnchor:
    return _make_native_anchor(
        kind,
        f"inserted.{seed}",
        source_doc_type=source_doc_type,
        text=text,
    )


def _default_table_width_pt(col_count: int) -> float:
    if col_count <= 0:
        return 0.0
    return max(72.0 * col_count, _DOCX_DEFAULT_TABLE_WIDTH_TWIPS / 20.0)


def _default_cell_width_pt(col_count: int) -> float:
    if col_count <= 0:
        return 72.0
    return max(72.0, _default_table_width_pt(col_count) / col_count)


def _default_table_style(row_count: int, col_count: int) -> TableStyleInfo:
    return TableStyleInfo(
        row_count=row_count,
        col_count=col_count,
        width_pt=_default_table_width_pt(col_count),
        height_pt=max(row_count, 0) * 18.0,
    )


def _default_cell_style(*, row_count: int, col_count: int) -> CellStyleInfo:
    return CellStyleInfo(
        width_pt=_default_cell_width_pt(col_count),
        height_pt=18.0,
        padding_top_pt=3.0,
        padding_right_pt=6.0,
        padding_bottom_pt=3.0,
        padding_left_pt=6.0,
        border_top="1px solid #000000",
        border_bottom="1px solid #000000",
        border_left="1px solid #000000",
        border_right="1px solid #000000",
    )


def _make_inserted_run(
    *,
    text: str,
    seed: str,
    source_doc_type: str | None,
    existing_ids: set[str],
) -> RunIR:
    return RunIR(
        node_id=_new_inserted_node_id("run", f"{seed}.run.{_text_digest(text)}", existing_ids),
        text=text,
        native_anchor=_inserted_anchor("run", seed, source_doc_type=source_doc_type, text=text),
    )


def _make_inserted_paragraph(
    *,
    text: str,
    seed: str,
    source_doc_type: str | None,
    existing_ids: set[str],
    page_number: int | None = None,
) -> ParagraphIR:
    paragraph = ParagraphIR(
        node_id=_new_inserted_node_id("paragraph", f"{seed}.paragraph.{_text_digest(text)}", existing_ids),
        page_number=page_number,
        native_anchor=_inserted_anchor("paragraph", seed, source_doc_type=source_doc_type, text=text),
    )
    paragraph.content.append(
        _make_inserted_run(
            text=text,
            seed=f"{seed}.p.run1",
            source_doc_type=source_doc_type,
            existing_ids=existing_ids,
        )
    )
    paragraph.recompute_text()
    return paragraph


def _make_inserted_cell(
    *,
    row_index: int,
    col_index: int,
    text: str,
    seed: str,
    source_doc_type: str | None,
    existing_ids: set[str],
    row_count: int = 1,
    col_count: int = 1,
    cell_style: CellStyleInfo | None = None,
    page_number: int | None = None,
) -> TableCellIR:
    cell = TableCellIR(
        node_id=_new_inserted_node_id("cell", f"{seed}.cell.r{row_index}.c{col_index}.{_text_digest(text)}", existing_ids),
        cell_style=cell_style.model_copy(deep=True) if cell_style is not None else _default_cell_style(row_count=row_count, col_count=col_count),
        native_anchor=_inserted_anchor("cell", seed, source_doc_type=source_doc_type, text=text),
    )
    cell.paragraphs.append(
        _make_inserted_paragraph(
            text=text,
            seed=f"{seed}.cell.r{row_index}.c{col_index}.p1",
            source_doc_type=source_doc_type,
            existing_ids=existing_ids,
            page_number=page_number,
        )
    )
    cell.recompute_text()
    return cell


def _make_inserted_table(
    *,
    rows: list[list[str]],
    seed: str,
    source_doc_type: str | None,
    existing_ids: set[str],
    page_number: int | None = None,
) -> TableIR:
    row_count = len(rows)
    col_count = len(rows[0]) if rows else 0
    table = TableIR(
        node_id=_new_inserted_node_id("table", f"{seed}.table.{row_count}x{col_count}", existing_ids),
        row_count=row_count,
        col_count=col_count,
        table_style=_default_table_style(row_count, col_count),
        native_anchor=_inserted_anchor("table", seed, source_doc_type=source_doc_type),
    )
    for row_index, row in enumerate(rows, start=1):
        cell_row: list[TableCellIR] = []
        for col_index, text in enumerate(row, start=1):
            cell_row.append(
                _make_inserted_cell(
                    row_index=row_index,
                    col_index=col_index,
                    text=text,
                    seed=f"{seed}.table.r{row_index}.c{col_index}",
                    source_doc_type=source_doc_type,
                    existing_ids=existing_ids,
                    row_count=row_count,
                    col_count=col_count,
                    page_number=page_number,
                )
            )
        table.cells.append(cell_row)
    return table


def _normalize_table_rows(rows: list[list[str]] | None) -> list[list[str]]:
    normalized = rows if rows is not None else [[""]]
    if not normalized or not normalized[0]:
        raise EditValidationError("insert_table requires at least one row and one column.", code="invalid_table_shape")
    width = len(normalized[0])
    if any(len(row) != width for row in normalized):
        raise EditValidationError("insert_table rows must be rectangular.", code="invalid_table_shape")
    return [[str(value) for value in row] for row in normalized]


def _assign_page_number_to_paragraph(paragraph: ParagraphIR, page_number: int | None) -> None:
    paragraph.page_number = page_number
    for node in paragraph.content:
        if isinstance(node, TableIR):
            for cell in node.iter_cells():
                for cell_paragraph in cell.paragraphs:
                    _assign_page_number_to_paragraph(cell_paragraph, page_number)


def _infer_cell_page_number(cell: TableCellIR) -> int | None:
    for paragraph in cell.paragraphs:
        if paragraph.page_number is not None:
            return paragraph.page_number
    return None


def _infer_table_page_number(index: _StructuralDocIrIndex, table: TableIR) -> int | None:
    table_location = index.tables.get(table.node_id)
    if table_location is not None and table_location.paragraph.page_number is not None:
        return table_location.paragraph.page_number
    for cell in table.iter_cells():
        if (page_number := _infer_cell_page_number(cell)) is not None:
            return page_number
    return None


def _collect_doc_ir_node_ids(node) -> list[str]:
    ids: list[str] = []

    def add(value: str | None) -> None:
        if value is not None:
            _append_unique(ids, value)

    def walk_paragraph(paragraph: ParagraphIR) -> None:
        add(paragraph.node_id)
        for item in paragraph.content:
            if isinstance(item, RunIR):
                add(item.node_id)
            elif isinstance(item, ImageIR):
                add(item.node_id)
            elif isinstance(item, TableIR):
                walk_table(item)

    def walk_table(table: TableIR) -> None:
        add(table.node_id)
        for cell in table.iter_cells():
            add(cell.node_id)
            for paragraph in cell.paragraphs:
                walk_paragraph(paragraph)

    if isinstance(node, ParagraphIR):
        walk_paragraph(node)
    elif isinstance(node, RunIR):
        add(node.node_id)
    elif isinstance(node, TableIR):
        walk_table(node)
    elif isinstance(node, TableCellIR):
        add(node.node_id)
        for paragraph in node.paragraphs:
            walk_paragraph(paragraph)
    return ids


def _replace_cell_paragraphs(
    cell: TableCellIR,
    text: str,
    *,
    seed: str,
    source_doc_type: str | None,
    existing_ids: set[str],
    result: _EditEngineResult,
    page_number: int | None = None,
) -> None:
    for paragraph in cell.paragraphs:
        for node_id in _collect_doc_ir_node_ids(paragraph):
            _append_unique(result.removed_target_ids, node_id)

    lines = text.split("\n") if text != "" else [""]
    cell.paragraphs = [
        _make_inserted_paragraph(
            text=line,
            seed=f"{seed}.p{index}",
            source_doc_type=source_doc_type,
            existing_ids=existing_ids,
            page_number=page_number,
        )
        for index, line in enumerate(lines, start=1)
    ]
    for paragraph in cell.paragraphs:
        for node_id in _collect_doc_ir_node_ids(paragraph):
            _append_unique(result.created_target_ids, node_id)
    cell.recompute_text()


def _recompute_table_shape(table: TableIR) -> None:
    table.expand_merged_cells()
    row_count = len(table.cells)
    col_count = max((len(row) for row in table.cells), default=0)
    for row_index, col_index, cell in table.iter_cell_positions():
        row_count = max(row_count, row_index + _cell_style_span(cell, "rowspan") - 1)
        col_count = max(col_count, col_index + _cell_style_span(cell, "colspan") - 1)
    table.row_count = row_count
    table.col_count = col_count
    if table.table_style is not None:
        table.table_style.row_count = table.row_count
        table.table_style.col_count = table.col_count


def _table_row_count(table: TableIR) -> int:
    return table.row_count or len(table.cells)


def _table_col_count(table: TableIR) -> int:
    return table.col_count or max((len(row) for row in table.cells), default=0)


def _table_cell_at(table: TableIR, *, row_index: int, col_index: int) -> TableCellIR | None:
    for cell_row_index, cell_col_index, cell in table.iter_cell_positions():
        if cell_row_index == row_index and cell_col_index == col_index:
            return cell
    return None


def _cloned_or_default_cell_style(
    template: TableCellIR | None,
    *,
    row_count: int,
    col_count: int,
) -> CellStyleInfo:
    if template is not None and template.cell_style is not None:
        return template.cell_style.model_copy(deep=True)
    return _default_cell_style(row_count=row_count, col_count=col_count)


def _resolve_table_axis(
    index: _StructuralDocIrIndex,
    operation: StructuralEdit,
    *,
    axis: str,
) -> tuple[TableIR, int]:
    cell_location = index.cells.get(operation.target_id)
    if cell_location is not None:
        position = cell_location.table.cell_position(cell_location.node)
        if position is None:
            raise EditValidationError(
                f"{operation.operation} target cell is not in its parent table: {operation.target_id}.",
                code="target_not_found",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        return cell_location.table, position[0] if axis == "row" else position[1]

    table_location = index.tables.get(operation.target_id)
    if table_location is None:
        raise EditValidationError(
            f"{operation.operation} target must be a table or cell: {operation.target_id}.",
            code="target_kind_mismatch",
            target_id=operation.target_id,
            operation=operation.operation,
        )

    axis_index = operation.row_index if axis == "row" else operation.column_index
    if axis_index is None:
        raise EditValidationError(
            f"{operation.operation} with a table target requires {axis}_index.",
            code="index_out_of_bounds",
            target_kind="table",
            target_id=operation.target_id,
            operation=operation.operation,
        )
    return table_location.node, axis_index


def _validate_expected_text_hash(expected_text_hash: str | None, current_text: str, operation: StructuralEdit, *, target_kind: str) -> None:
    current_text_hash = _text_hash(current_text)
    if expected_text_hash is not None and current_text_hash != expected_text_hash:
        raise EditValidationError(
            f"Text hash mismatch for {operation.target_id}.",
            code="text_hash_mismatch",
            target_kind=target_kind,
            target_id=operation.target_id,
            operation=operation.operation,
            expected_text_hash=expected_text_hash,
            current_text_hash=current_text_hash,
            current_text=current_text,
        )


def _ensure_doc_ir_has_content(doc: DocIR, existing_ids: set[str], result: _EditEngineResult) -> None:
    if doc.paragraphs:
        return
    default_page_number = doc.pages[0].page_number if doc.pages else None
    paragraph = _make_inserted_paragraph(
        text="",
        seed="document.empty.p1",
        source_doc_type=doc.source_doc_type,
        existing_ids=existing_ids,
        page_number=default_page_number,
    )
    doc.paragraphs.append(paragraph)
    for node_id in _collect_doc_ir_node_ids(paragraph):
        _append_unique(result.created_target_ids, node_id)


def _ensure_cell_has_paragraph(cell: TableCellIR, doc: DocIR, existing_ids: set[str], result: _EditEngineResult) -> None:
    if cell.paragraphs:
        return
    page_number = _infer_cell_page_number(cell)
    paragraph = _make_inserted_paragraph(
        text="",
        seed=f"{cell.node_id}.empty.p1",
        source_doc_type=doc.source_doc_type,
        existing_ids=existing_ids,
        page_number=page_number,
    )
    cell.paragraphs.append(paragraph)
    cell.recompute_text()
    for node_id in _collect_doc_ir_node_ids(paragraph):
        _append_unique(result.created_target_ids, node_id)


def _apply_structural_doc_ir_operation(
    doc: DocIR,
    operation: StructuralEdit,
    result: _EditEngineResult,
    *,
    existing_ids: set[str],
    sequence: int,
) -> None:
    index = _build_structural_doc_ir_index(doc)
    seed = f"op{sequence}.{operation.operation}.{operation.target_id}"

    if operation.operation == "insert_paragraph":
        text = operation.text or ""
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is not None:
            if operation.position not in {"before", "after"}:
                raise EditValidationError(
                    "insert_paragraph with a paragraph target requires position before or after.",
                    code="invalid_position",
                    target_kind="paragraph",
                    target_id=operation.target_id,
                    operation=operation.operation,
                )
            inserted = _make_inserted_paragraph(
                text=text,
                seed=seed,
                source_doc_type=doc.source_doc_type,
                existing_ids=existing_ids,
                page_number=paragraph_location.node.page_number,
            )
            offset = 0 if operation.position == "before" else 1
            paragraph_location.container.insert(paragraph_location.index + offset, inserted)
        elif (cell_location := index.cells.get(operation.target_id)) is not None:
            if operation.position not in {"start", "end", "before", "after"}:
                raise EditValidationError("Invalid insert_paragraph cell position.", code="invalid_position", operation=operation.operation)
            inserted = _make_inserted_paragraph(
                text=text,
                seed=seed,
                source_doc_type=doc.source_doc_type,
                existing_ids=existing_ids,
                page_number=_infer_cell_page_number(cell_location.node),
            )
            insert_index = 0 if operation.position in {"start", "before"} else len(cell_location.node.paragraphs)
            cell_location.node.paragraphs.insert(insert_index, inserted)
            cell_location.node.recompute_text()
        else:
            raise EditValidationError(
                f"insert_paragraph target must be a paragraph or cell: {operation.target_id}.",
                code="target_not_found",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        for node_id in _collect_doc_ir_node_ids(inserted):
            _append_unique(result.created_target_ids, node_id)
        _append_unique(result.modified_target_ids, operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_paragraph":
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is None:
            raise EditValidationError(
                f"Paragraph does not exist: {operation.target_id}.",
                code="target_not_found",
                target_kind="paragraph",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        _validate_expected_text_hash(operation.expected_text_hash, paragraph_location.node.text, operation, target_kind="paragraph")
        for node_id in _collect_doc_ir_node_ids(paragraph_location.node):
            _append_unique(result.removed_target_ids, node_id)
        del paragraph_location.container[paragraph_location.index]
        if paragraph_location.parent_cell is not None:
            _ensure_cell_has_paragraph(paragraph_location.parent_cell, doc, existing_ids, result)
            paragraph_location.parent_cell.recompute_text()
        else:
            _ensure_doc_ir_has_content(doc, existing_ids, result)
        result.operations_applied += 1
        return

    if operation.operation == "insert_run":
        text = operation.text or ""
        if (run_location := index.runs.get(operation.target_id)) is not None:
            if operation.position not in {"before", "after"}:
                raise EditValidationError(
                    "insert_run with a run target requires position before or after.",
                    code="invalid_position",
                    target_kind="run",
                    target_id=operation.target_id,
                    operation=operation.operation,
                )
            inserted = _make_inserted_run(
                text=text,
                seed=seed,
                source_doc_type=doc.source_doc_type,
                existing_ids=existing_ids,
            )
            offset = 0 if operation.position == "before" else 1
            run_location.paragraph.content.insert(run_location.content_index + offset, inserted)
            run_location.paragraph.recompute_text()
        elif (paragraph_location := index.paragraphs.get(operation.target_id)) is not None:
            if operation.position not in {"start", "end", "before", "after"}:
                raise EditValidationError("Invalid insert_run paragraph position.", code="invalid_position", operation=operation.operation)
            inserted = _make_inserted_run(
                text=text,
                seed=seed,
                source_doc_type=doc.source_doc_type,
                existing_ids=existing_ids,
            )
            insert_index = 0 if operation.position in {"start", "before"} else len(paragraph_location.node.content)
            paragraph_location.node.content.insert(insert_index, inserted)
            paragraph_location.node.recompute_text()
        else:
            raise EditValidationError(
                f"insert_run target must be a run or paragraph: {operation.target_id}.",
                code="target_not_found",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        _append_unique(result.created_target_ids, inserted.node_id)
        _append_unique(result.modified_target_ids, operation.target_id)
        _append_unique(result.modified_run_ids, inserted.node_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_run":
        run_location = index.runs.get(operation.target_id)
        if run_location is None:
            raise EditValidationError(
                f"Run does not exist: {operation.target_id}.",
                code="target_not_found",
                target_kind="run",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        _validate_expected_text_hash(operation.expected_text_hash, run_location.node.text, operation, target_kind="run")
        run_location.paragraph.content.pop(run_location.content_index)
        run_location.paragraph.recompute_text()
        _append_unique(result.removed_target_ids, operation.target_id)
        _append_unique(result.modified_target_ids, run_location.paragraph.node_id)
        result.operations_applied += 1
        return

    if operation.operation == "insert_table":
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is None:
            raise EditValidationError(
                f"insert_table target must be a paragraph: {operation.target_id}.",
                code="target_not_found",
                target_kind="paragraph",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        if operation.position not in {"before", "after"}:
            raise EditValidationError("insert_table requires position before or after.", code="invalid_position", operation=operation.operation)
        table = _make_inserted_table(
            rows=_normalize_table_rows(operation.rows),
            seed=seed,
            source_doc_type=doc.source_doc_type,
            existing_ids=existing_ids,
            page_number=paragraph_location.node.page_number,
        )
        table_paragraph = ParagraphIR(
            node_id=_new_inserted_node_id("paragraph", f"{seed}.table.paragraph", existing_ids),
            page_number=paragraph_location.node.page_number,
            native_anchor=_inserted_anchor("paragraph", seed, source_doc_type=doc.source_doc_type),
            content=[table],
        )
        table_paragraph.recompute_text()
        _assign_page_number_to_paragraph(table_paragraph, paragraph_location.node.page_number)
        offset = 0 if operation.position == "before" else 1
        paragraph_location.container.insert(paragraph_location.index + offset, table_paragraph)
        for node_id in _collect_doc_ir_node_ids(table_paragraph):
            _append_unique(result.created_target_ids, node_id)
        _append_unique(result.modified_target_ids, operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_table":
        table_location = index.tables.get(operation.target_id)
        if table_location is None:
            raise EditValidationError(
                f"Table does not exist: {operation.target_id}.",
                code="target_not_found",
                target_kind="table",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        for node_id in _collect_doc_ir_node_ids(table_location.node):
            _append_unique(result.removed_target_ids, node_id)
        table_location.paragraph.content.pop(table_location.content_index)
        table_location.paragraph.recompute_text()
        _append_unique(result.modified_target_ids, table_location.paragraph.node_id)
        result.operations_applied += 1
        return

    if operation.operation == "set_cell_text":
        cell_location = index.cells.get(operation.target_id)
        if cell_location is None:
            raise EditValidationError(
                f"Cell does not exist: {operation.target_id}.",
                code="target_not_found",
                target_kind="cell",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        _validate_expected_text_hash(operation.expected_text_hash, cell_location.node.text, operation, target_kind="cell")
        _replace_cell_paragraphs(
            cell_location.node,
            operation.text or "",
            seed=seed,
            source_doc_type=doc.source_doc_type,
            existing_ids=existing_ids,
            result=result,
            page_number=_infer_cell_page_number(cell_location.node),
        )
        _append_unique(result.modified_target_ids, operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation in {"insert_table_row", "remove_table_row"}:
        table, row_index = _resolve_table_axis(index, operation, axis="row")
        row_count = _table_row_count(table)
        col_count = _table_col_count(table)
        table_page_number = _infer_table_page_number(index, table)
        if row_index < 1 or row_index > row_count:
            raise EditValidationError("Table row index is out of bounds.", code="index_out_of_bounds", operation=operation.operation)
        if operation.operation == "remove_table_row":
            if row_count <= 1:
                raise EditValidationError("Cannot remove the only table row.", code="invalid_table_shape", operation=operation.operation)
            removed_cells = [cell for cell_row_index, _col_index, cell in table.iter_cell_positions() if cell_row_index == row_index]
            removed_cell_ids = {id(cell) for cell in removed_cells}
            for cell in removed_cells:
                for node_id in _collect_doc_ir_node_ids(cell):
                    _append_unique(result.removed_target_ids, node_id)
            if removed_cell_ids:
                table.cells = [
                    [cell for cell in row if id(cell) not in removed_cell_ids]
                    for row in table.cells
                ]
            if row_index - 1 < len(table.cells):
                del table.cells[row_index - 1]
        else:
            values = operation.values or ["" for _ in range(col_count)]
            if len(values) != col_count:
                raise EditValidationError("Inserted row values must match table column count.", code="invalid_table_shape", operation=operation.operation)
            insert_at = row_index if operation.position in {"after", "end"} else row_index - 1
            template_row_index = row_index
            template_styles = {
                col_index: _cloned_or_default_cell_style(
                    _table_cell_at(table, row_index=template_row_index, col_index=col_index),
                    row_count=row_count + 1,
                    col_count=col_count,
                )
                for col_index in range(1, col_count + 1)
            }
            new_row: list[TableCellIR] = []
            for col_index, text in enumerate(values, start=1):
                new_cell = _make_inserted_cell(
                    row_index=insert_at + 1,
                    col_index=col_index,
                    text=text,
                    seed=f"{seed}.row{insert_at + 1}.c{col_index}",
                    source_doc_type=doc.source_doc_type,
                    existing_ids=existing_ids,
                    row_count=row_count + 1,
                    col_count=col_count,
                    cell_style=template_styles[col_index],
                    page_number=table_page_number,
                )
                new_row.append(new_cell)
                for node_id in _collect_doc_ir_node_ids(new_cell):
                    _append_unique(result.created_target_ids, node_id)
            table.cells.insert(insert_at, new_row)
        _recompute_table_shape(table)
        _append_unique(result.modified_target_ids, table.node_id)
        result.operations_applied += 1
        return

    if operation.operation in {"insert_table_column", "remove_table_column"}:
        table, column_index = _resolve_table_axis(index, operation, axis="column")
        row_count = _table_row_count(table)
        col_count = _table_col_count(table)
        table_page_number = _infer_table_page_number(index, table)
        if column_index < 1 or column_index > col_count:
            raise EditValidationError("Table column index is out of bounds.", code="index_out_of_bounds", operation=operation.operation)
        if operation.operation == "remove_table_column":
            if col_count <= 1:
                raise EditValidationError("Cannot remove the only table column.", code="invalid_table_shape", operation=operation.operation)
            removed_cells = [cell for _row_index, col_index, cell in table.iter_cell_positions() if col_index == column_index]
            removed_cell_ids = {id(cell) for cell in removed_cells}
            for cell in removed_cells:
                for node_id in _collect_doc_ir_node_ids(cell):
                    _append_unique(result.removed_target_ids, node_id)
            table.cells = [
                [
                    cell
                    for cell in row
                    if id(cell) not in removed_cell_ids
                ]
                for row in table.cells
            ]
        else:
            values = operation.values or ["" for _ in range(row_count)]
            if len(values) != row_count:
                raise EditValidationError("Inserted column values must match table row count.", code="invalid_table_shape", operation=operation.operation)
            insert_at = column_index if operation.position in {"after", "end"} else column_index - 1
            template_column_index = column_index
            template_styles = {
                row_index: _cloned_or_default_cell_style(
                    _table_cell_at(table, row_index=row_index, col_index=template_column_index),
                    row_count=row_count,
                    col_count=col_count + 1,
                )
                for row_index in range(1, row_count + 1)
            }
            for row_index, text in enumerate(values, start=1):
                new_cell = _make_inserted_cell(
                    row_index=row_index,
                    col_index=insert_at + 1,
                    text=text,
                    seed=f"{seed}.column{insert_at + 1}.r{row_index}",
                    source_doc_type=doc.source_doc_type,
                    existing_ids=existing_ids,
                    row_count=row_count,
                    col_count=col_count + 1,
                    cell_style=template_styles[row_index],
                    page_number=table_page_number,
                )
                table.append_cell(new_cell, row_index=row_index, col_index=insert_at + 1)
                for node_id in _collect_doc_ir_node_ids(new_cell):
                    _append_unique(result.created_target_ids, node_id)
        _recompute_table_shape(table)
        _append_unique(result.modified_target_ids, table.node_id)
        result.operations_applied += 1
        return

    raise EditValidationError(
        f"Unsupported structural operation: {operation.operation!r}.",
        code="invalid_operation",
        operation=operation.operation,
    )


def _refresh_anchor(
    node,
    kind: NodeKind,
    structural_path: str,
    *,
    source_doc_type: str | None,
    parent_debug_path: str | None,
    text: str | None = None,
) -> None:
    if node.native_anchor is None:
        node.native_anchor = _make_native_anchor(
            kind,
            structural_path,
            source_doc_type=source_doc_type,
            parent_debug_path=parent_debug_path,
            text=text,
        )
    else:
        node.native_anchor.node_kind = kind
        node.native_anchor.debug_path = structural_path
        node.native_anchor.structural_path = structural_path
        node.native_anchor.parent_debug_path = parent_debug_path
        node.native_anchor.source_doc_type = source_doc_type
        node.native_anchor.text_hash = hashlib.sha1((text or "").encode("utf-8")).hexdigest() if text is not None else None
    if node.node_id is None:
        node.node_id = _anchored_node_id(kind, structural_path)


def _refresh_doc_ir_native_paths(doc: DocIR) -> None:
    docx_top_level_table_counter = 0

    def walk_paragraph(
        paragraph: ParagraphIR,
        paragraph_path: str,
        *,
        parent_debug_path: str | None,
        top_level: bool,
    ) -> None:
        nonlocal docx_top_level_table_counter
        paragraph.recompute_text()
        _refresh_anchor(
            paragraph,
            "paragraph",
            paragraph_path,
            source_doc_type=doc.source_doc_type,
            parent_debug_path=parent_debug_path,
            text=paragraph.text,
        )
        run_index = 0
        image_index = 0
        table_index = 0
        for item in paragraph.content:
            if isinstance(item, RunIR):
                run_index += 1
                _refresh_anchor(
                    item,
                    "run",
                    f"{paragraph_path}.r{run_index}",
                    source_doc_type=doc.source_doc_type,
                    parent_debug_path=paragraph_path,
                    text=item.text,
                )
            elif isinstance(item, ImageIR):
                image_index += 1
                _refresh_anchor(
                    item,
                    "image",
                    f"{paragraph_path}.img{image_index}",
                    source_doc_type=doc.source_doc_type,
                    parent_debug_path=paragraph_path,
                )
            elif isinstance(item, TableIR):
                table_index += 1
                if top_level:
                    if doc.source_doc_type == "docx":
                        docx_top_level_table_counter += 1
                        native_table_index = docx_top_level_table_counter
                    else:
                        native_table_index = table_index
                    table_path = f"{paragraph_path}.r1.tbl{native_table_index}"
                else:
                    table_path = f"{paragraph_path}.tbl{table_index}"
                walk_table(item, table_path, parent_debug_path=paragraph_path)

    def walk_table(table: TableIR, table_path: str, *, parent_debug_path: str | None) -> None:
        _recompute_table_shape(table)
        _refresh_anchor(
            table,
            "table",
            table_path,
            source_doc_type=doc.source_doc_type,
            parent_debug_path=parent_debug_path,
        )
        for row_index, col_index, cell in table.iter_cell_positions():
            cell.recompute_text()
            cell_path = f"{table_path}.tr{row_index}.tc{col_index}"
            _refresh_anchor(
                cell,
                "cell",
                cell_path,
                source_doc_type=doc.source_doc_type,
                parent_debug_path=table_path,
                text=cell.text,
            )
            for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                walk_paragraph(
                    paragraph,
                    f"{cell_path}.p{paragraph_index}",
                    parent_debug_path=cell_path,
                    top_level=False,
                )

    for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
        walk_paragraph(paragraph, f"s1.p{paragraph_index}", parent_debug_path=None, top_level=True)


def _apply_structural_edits_to_doc_ir(doc: DocIR, operations: list[StructuralEdit]) -> _EditEngineResult:
    updated = doc.model_copy(deep=True)
    updated.ensure_node_identity()
    existing_ids = _all_doc_ir_node_ids(updated)
    result = _EditEngineResult(source_doc_type=updated.source_doc_type)
    for sequence, operation in enumerate(operations, start=1):
        _apply_structural_doc_ir_operation(
            updated,
            operation,
            result,
            existing_ids=existing_ids,
            sequence=sequence,
        )
    _refresh_doc_ir_native_paths(updated)
    result.updated_doc_ir = updated
    return result


def _default_output_path(source_path: Path, *, output_suffix: str | None = None) -> Path:
    suffix = output_suffix if output_suffix is not None else source_path.suffix
    return source_path.with_name(f"{source_path.stem}_edited{suffix}")


def _expected_writeback_suffix(source_doc_type: str | None) -> str | None:
    if source_doc_type == "docx":
        return ".docx"
    if source_doc_type in {"hwpx", "hwp"}:
        return ".hwpx"
    return None


def _normalize_output_path_for_source_doc_type(
    target_path: Path,
    *,
    source_doc_type: str | None,
    result: _EditEngineResult,
) -> Path:
    expected_suffix = _expected_writeback_suffix(source_doc_type)
    if expected_suffix is None or target_path.suffix.lower() == expected_suffix:
        return target_path

    adjusted_target_path = target_path.with_suffix(expected_suffix)
    if source_doc_type == "hwp":
        result.warnings.append(
            f"HWP sources are written back as HWPX; adjusted output path to {adjusted_target_path}."
        )
    else:
        result.warnings.append(
            f"{str(source_doc_type).upper()} write-back keeps the native {expected_suffix} format; "
            f"adjusted output path to {adjusted_target_path}."
        )
    return adjusted_target_path


def _same_path(left: Path, right: Path) -> bool:
    return left.expanduser().resolve(strict=False) == right.expanduser().resolve(strict=False)


def _default_output_filename(
    *,
    source_name: str | None,
    source_doc_type: str | None,
) -> str:
    source_doc_type = source_doc_type or "docx"
    if source_name:
        source = Path(source_name)
    else:
        suffix = f".{source_doc_type}" if source_doc_type != "hwp" else ".hwp"
        source = Path(f"document{suffix}")

    suffix = ".hwpx" if source_doc_type == "hwp" else (source.suffix or f".{source_doc_type}")
    stem = source.stem if source.suffix else source.name
    return f"{stem}_edited{suffix}"


def _source_suffix_for_doc_type(doc_type: SourceDocType | str) -> str:
    return {
        "docx": ".docx",
        "hwpx": ".hwpx",
        "hwp": ".hwp",
        "pdf": ".pdf",
    }.get(str(doc_type), ".bin")


def _resolve_bytes_doc_type(
    source_bytes: bytes,
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
) -> str:
    if doc_type != "auto":
        return doc_type
    if source_name:
        try:
            return infer_doc_type(Path(source_name), "auto")
        except ValueError:
            pass
    return infer_doc_type(source_bytes, "auto")


@dataclass
class _DocxParagraphLocation:
    paragraph: object
    path: str


@dataclass
class _DocxRunLocation:
    run: object
    paragraph: object
    path: str


@dataclass
class _DocxTableLocation:
    table: object
    path: str


@dataclass
class _DocxImageLocation:
    drawing_el: object
    path: str


@dataclass
class _DocxCellLocation:
    cell: object
    table: object
    row_index: int
    col_index: int
    path: str


class _DocxStructuralIndex:
    def __init__(self) -> None:
        self.paragraphs: dict[str, _DocxParagraphLocation] = {}
        self.runs: dict[str, _DocxRunLocation] = {}
        self.tables: dict[str, _DocxTableLocation] = {}
        self.images: dict[str, _DocxImageLocation] = {}
        self.cells: dict[str, _DocxCellLocation] = {}


def _build_docx_structural_index(doc) -> _DocxStructuralIndex:
    index = _DocxStructuralIndex()

    def register_paragraph(paragraph, paragraph_path: str) -> None:
        paragraph_id = _anchored_node_id("paragraph", paragraph_path)
        index.paragraphs[paragraph_id] = _DocxParagraphLocation(paragraph=paragraph, path=paragraph_path)
        image_index = 0
        for run_index, run in enumerate(paragraph.runs, start=1):
            run_path = f"{paragraph_path}.r{run_index}"
            index.runs[_anchored_node_id("run", run_path)] = _DocxRunLocation(
                run=run,
                paragraph=paragraph,
                path=run_path,
            )
            for element in run._r.iter():
                tag = getattr(element, "tag", "")
                if not isinstance(tag, str) or not tag.endswith("}blip"):
                    continue
                drawing_parent = element
                while drawing_parent is not None and not (
                    isinstance(getattr(drawing_parent, "tag", None), str)
                    and drawing_parent.tag.rsplit("}", 1)[-1] in {"inline", "anchor"}
                ):
                    drawing_parent = drawing_parent.getparent()
                if drawing_parent is None:
                    continue
                image_index += 1
                image_path = f"{paragraph_path}.img{image_index}"
                index.images[_anchored_node_id("image", image_path)] = _DocxImageLocation(
                    drawing_el=drawing_parent,
                    path=image_path,
                )

    def walk_table(table, table_path: str) -> None:
        index.tables[_anchored_node_id("table", table_path)] = _DocxTableLocation(table=table, path=table_path)
        for row_index, row in enumerate(table.rows, start=1):
            for col_index, cell in enumerate(row.cells, start=1):
                cell_path = f"{table_path}.tr{row_index}.tc{col_index}"
                index.cells[_anchored_node_id("cell", cell_path)] = _DocxCellLocation(
                    cell=cell,
                    table=table,
                    row_index=row_index,
                    col_index=col_index,
                    path=cell_path,
                )
                cp_idx = 0
                current_paragraph_path: str | None = None
                nested_table_counter_by_paragraph: dict[str, int] = {}
                for block in _iter_docx_blocks_from_element(cell, cell._tc):
                    if block.__class__.__name__ == "Paragraph":
                        cp_idx += 1
                        current_paragraph_path = f"{cell_path}.p{cp_idx}"
                        register_paragraph(block, current_paragraph_path)
                        continue
                    if block.__class__.__name__ != "Table":
                        continue
                    if current_paragraph_path is None:
                        cp_idx += 1
                        current_paragraph_path = f"{cell_path}.p{cp_idx}"
                    table_counter = nested_table_counter_by_paragraph.get(current_paragraph_path, 0) + 1
                    nested_table_counter_by_paragraph[current_paragraph_path] = table_counter
                    walk_table(block, f"{current_paragraph_path}.tbl{table_counter}")

    p_idx = 0
    table_counter = 0
    for block in _iter_docx_blocks(doc):
        if block.__class__.__name__ == "Paragraph":
            p_idx += 1
            register_paragraph(block, f"s1.p{p_idx}")
            continue
        if block.__class__.__name__ != "Table":
            continue
        table_counter += 1
        p_idx += 1
        walk_table(block, f"s1.p{p_idx}.r1.tbl{table_counter}")

    return index


def _docx_text_run_el(text: str):
    from docx.oxml import OxmlElement

    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def _docx_paragraph_el(text: str):
    from docx.oxml import OxmlElement

    paragraph = OxmlElement("w:p")
    if text != "":
        paragraph.append(_docx_text_run_el(text))
    return paragraph


def _docx_default_column_widths_twips(col_count: int) -> list[int]:
    if col_count <= 0:
        return []
    width = max(_DOCX_MIN_COLUMN_WIDTH_TWIPS, _DOCX_DEFAULT_TABLE_WIDTH_TWIPS // col_count)
    return [width for _ in range(col_count)]


def _docx_border_el(name: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    border = OxmlElement(f"w:{name}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), _DOCX_DEFAULT_BORDER_SIZE)
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")
    return border


def _docx_table_pr_el(table_width_twips: int):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table_pr = OxmlElement("w:tblPr")
    table_w = OxmlElement("w:tblW")
    table_w.set(qn("w:type"), "dxa")
    table_w.set(qn("w:w"), str(table_width_twips))
    table_pr.append(table_w)

    table_layout = OxmlElement("w:tblLayout")
    table_layout.set(qn("w:type"), "fixed")
    table_pr.append(table_layout)

    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_docx_border_el(name))
    table_pr.append(borders)

    cell_margin = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), str(_DOCX_DEFAULT_CELL_MARGIN_TWIPS))
        margin.set(qn("w:type"), "dxa")
        cell_margin.append(margin)
    table_pr.append(cell_margin)
    return table_pr


def _docx_tc_pr_width(tc_pr, width_twips: int | None) -> None:
    if width_twips is None:
        return

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(width_twips))


def _docx_cell_el(text: str, *, width_twips: int | None = None, tc_pr_template=None):
    from docx.oxml import OxmlElement

    cell = OxmlElement("w:tc")
    cell_pr = deepcopy(tc_pr_template) if tc_pr_template is not None else OxmlElement("w:tcPr")
    _docx_tc_pr_width(cell_pr, width_twips)
    cell.append(cell_pr)
    cell.append(_docx_paragraph_el(text))
    return cell


def _docx_row_el(values: list[str], *, col_widths: list[int] | None = None, template_row=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    row = OxmlElement("w:tr")
    if template_row is not None:
        template_row_pr = template_row.find(qn("w:trPr"))
        if template_row_pr is not None:
            row.append(deepcopy(template_row_pr))

    template_cells = _docx_row_cells(template_row) if template_row is not None else []
    col_widths = col_widths or _docx_default_column_widths_twips(len(values))
    for index, value in enumerate(values):
        template_tc_pr = None
        if index < len(template_cells):
            template_tc_pr = template_cells[index].find(qn("w:tcPr"))
        row.append(_docx_cell_el(value, width_twips=col_widths[index] if index < len(col_widths) else None, tc_pr_template=template_tc_pr))
    return row


def _docx_table_el(rows: list[list[str]]):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = OxmlElement("w:tbl")
    col_widths = _docx_default_column_widths_twips(len(rows[0]))
    table_pr = _docx_table_pr_el(sum(col_widths))
    table_grid = OxmlElement("w:tblGrid")
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        table_grid.append(grid_col)
    table.append(table_pr)
    table.append(table_grid)
    for row in rows:
        table.append(_docx_row_el(row, col_widths=col_widths))
    return table


def _docx_table_rows(table_el) -> list:
    from docx.oxml.ns import qn

    return list(table_el.findall(qn("w:tr")))


def _docx_row_cells(row_el) -> list:
    from docx.oxml.ns import qn

    return list(row_el.findall(qn("w:tc")))


def _docx_table_col_count(table_el) -> int:
    rows = _docx_table_rows(table_el)
    if not rows:
        return 0
    return len(_docx_row_cells(rows[0]))


def _docx_adjust_table_grid(table_el, *, column_index: int, operation: str, position: str = "after", width_twips: int | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    grid = table_el.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table_el.insert(1 if len(table_el) else 0, grid)
    grid_cols = list(grid.findall(qn("w:gridCol")))
    if operation == "insert":
        new_col = OxmlElement("w:gridCol")
        if width_twips is None and 1 <= column_index <= len(grid_cols):
            width_twips = _safe_int(grid_cols[column_index - 1].get(qn("w:w")))
        new_col.set(qn("w:w"), str(width_twips or _DOCX_MIN_COLUMN_WIDTH_TWIPS))
        insert_index = column_index if position in {"after", "end"} else column_index - 1
        grid.insert(max(0, min(insert_index, len(grid_cols))), new_col)
        return
    if 1 <= column_index <= len(grid_cols):
        grid.remove(grid_cols[column_index - 1])


def _ensure_docx_container_has_paragraph(parent_el) -> None:
    from docx.oxml.ns import qn

    if parent_el.findall(qn("w:p")):
        return
    parent_el.append(_docx_paragraph_el(""))


def _docx_insert_paragraph_in_cell(cell, paragraph_el, *, position: str) -> None:
    from docx.oxml.ns import qn

    tc = cell._tc
    if position in {"start", "before"}:
        children = list(tc)
        insert_index = 1 if children and children[0].tag == qn("w:tcPr") else 0
        tc.insert(insert_index, paragraph_el)
    else:
        tc.append(paragraph_el)


def _docx_set_cell_text(cell, text: str) -> None:
    cell._tc.clear_content()
    lines = text.split("\n") if text != "" else [""]
    for line in lines:
        cell.add_paragraph(line)


def _resolve_docx_table_axis(
    index: _DocxStructuralIndex,
    operation: StructuralEdit,
    *,
    axis: str,
):
    if cell_location := index.cells.get(operation.target_id):
        return cell_location.table, cell_location.row_index if axis == "row" else cell_location.col_index
    if table_location := index.tables.get(operation.target_id):
        axis_index = operation.row_index if axis == "row" else operation.column_index
        if axis_index is None:
            raise EditValidationError(
                f"{operation.operation} with a table target requires {axis}_index.",
                code="index_out_of_bounds",
                target_kind="table",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        return table_location.table, axis_index
    raise EditValidationError(
        f"{operation.operation} target must be a table or cell: {operation.target_id}.",
        code="target_not_found",
        target_id=operation.target_id,
        operation=operation.operation,
    )


def _docx_get_or_add_child(parent, tag: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _docx_set_on_off(parent, tag: str, value: bool | None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    existing = parent.find(qn(tag))
    if value is None:
        if existing is not None:
            parent.remove(existing)
        return
    if existing is None:
        existing = OxmlElement(tag)
        parent.append(existing)
    existing.set(qn("w:val"), "1" if value else "0")


def _docx_hex_color(value: str) -> str:
    return value.strip().lstrip("#").upper()


def _apply_docx_run_style(run, edit: StyleEdit, result: _EditEngineResult) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if "bold" in edit.clear_fields:
        run.bold = None
    elif edit.bold is not None:
        run.bold = edit.bold
    if "italic" in edit.clear_fields:
        run.italic = None
    elif edit.italic is not None:
        run.italic = edit.italic
    if "underline" in edit.clear_fields:
        run.underline = None
    elif edit.underline is not None:
        run.underline = edit.underline

    font = run.font
    for field_name, attr in (("strikethrough", "strike"), ("superscript", "superscript"), ("subscript", "subscript")):
        if field_name in edit.clear_fields:
            setattr(font, attr, None)
        elif (value := getattr(edit, field_name)) is not None:
            setattr(font, attr, value)
    if edit.superscript:
        font.subscript = False
    if edit.subscript:
        font.superscript = False

    if "color" in edit.clear_fields:
        r_pr = run._r.get_or_add_rPr()
        color_el = r_pr.find(qn("w:color"))
        if color_el is not None:
            r_pr.remove(color_el)
    elif edit.color is not None:
        font.color.rgb = RGBColor.from_string(_docx_hex_color(edit.color))

    if "font_size_pt" in edit.clear_fields:
        font.size = None
    elif edit.font_size_pt is not None:
        font.size = Pt(edit.font_size_pt)

    if "highlight" in edit.clear_fields:
        r_pr = run._r.get_or_add_rPr()
        highlight_el = r_pr.find(qn("w:highlight"))
        if highlight_el is not None:
            r_pr.remove(highlight_el)
    elif edit.highlight is not None:
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), str(edit.highlight).lower().lstrip("#"))
        r_pr = run._r.get_or_add_rPr()
        old = r_pr.find(qn("w:highlight"))
        if old is not None:
            r_pr.remove(old)
        r_pr.append(highlight)


def _apply_docx_paragraph_style(paragraph, edit: StyleEdit) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if "paragraph_align" in edit.clear_fields:
        paragraph.alignment = None
    elif edit.paragraph_align is not None:
        paragraph.alignment = align_map[edit.paragraph_align]

    fmt = paragraph.paragraph_format
    for field_name, attr in (
        ("left_indent_pt", "left_indent"),
        ("right_indent_pt", "right_indent"),
    ):
        if field_name in edit.clear_fields:
            setattr(fmt, attr, None)
        elif (value := getattr(edit, field_name)) is not None:
            setattr(fmt, attr, Pt(value))

    if "first_line_indent_pt" in edit.clear_fields or "hanging_indent_pt" in edit.clear_fields:
        fmt.first_line_indent = None
    if edit.first_line_indent_pt is not None:
        fmt.first_line_indent = Pt(edit.first_line_indent_pt)
    if edit.hanging_indent_pt is not None:
        fmt.first_line_indent = Pt(-edit.hanging_indent_pt)


def _docx_get_or_add_tbl_pr(table_el):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table_el.insert(0, tbl_pr)
    return tbl_pr


def _docx_get_or_add_tc_pr(cell_el):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell_el.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        cell_el.insert(0, tc_pr)
    return tc_pr


def _docx_set_width_el(parent, tag: str, value_pt: float | None, *, clear: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    existing = parent.find(qn(tag))
    if clear:
        if existing is not None:
            parent.remove(existing)
        return
    width = _pt_to_twips(value_pt)
    if width is None:
        return
    if existing is None:
        existing = OxmlElement(tag)
        parent.insert(0, existing)
    existing.set(qn("w:type"), "dxa")
    existing.set(qn("w:w"), str(width))


def _docx_set_width_twips_el(parent, tag: str, width_twips: int | None, *, clear: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    existing = parent.find(qn(tag))
    if clear:
        if existing is not None:
            parent.remove(existing)
        return
    if width_twips is None:
        return
    if existing is None:
        existing = OxmlElement(tag)
        parent.insert(0, existing)
    existing.set(qn("w:type"), "dxa")
    existing.set(qn("w:w"), str(width_twips))


def _docx_grid_span(tc_el) -> int:
    from docx.oxml.ns import qn

    tc_pr = tc_el.find(qn("w:tcPr"))
    grid_span = tc_pr.find(qn("w:gridSpan")) if tc_pr is not None else None
    return max(_safe_int(grid_span.get(qn("w:val"))) if grid_span is not None else 1, 1)


def _docx_row_logical_cells(row_el):
    logical_col = 1
    for cell_el in _docx_row_cells(row_el):
        span = _docx_grid_span(cell_el)
        yield logical_col, span, cell_el
        logical_col += span


def _docx_table_logical_col_count(table_el) -> int:
    from docx.oxml.ns import qn

    grid = table_el.find(qn("w:tblGrid"))
    grid_col_count = len(grid.findall(qn("w:gridCol"))) if grid is not None else 0
    row_col_count = 0
    for row_el in _docx_table_rows(table_el):
        for logical_col, span, _cell_el in _docx_row_logical_cells(row_el):
            row_col_count = max(row_col_count, logical_col + span - 1)
    return max(grid_col_count, row_col_count)


def _docx_set_fixed_table_layout(table_el) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = _docx_get_or_add_tbl_pr(table_el)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _docx_ensure_table_grid(table_el, col_count: int):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    grid = table_el.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        insert_at = 1 if table_el.find(qn("w:tblPr")) is not None else 0
        table_el.insert(insert_at, grid)
    grid_cols = list(grid.findall(qn("w:gridCol")))
    while len(grid_cols) < col_count:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(_DOCX_MIN_COLUMN_WIDTH_TWIPS))
        grid.append(grid_col)
        grid_cols.append(grid_col)
    return grid, grid_cols


def _docx_grid_widths(table_el, col_count: int) -> list[int]:
    from docx.oxml.ns import qn

    _grid, grid_cols = _docx_ensure_table_grid(table_el, col_count)
    widths: list[int] = []
    for grid_col in grid_cols[:col_count]:
        widths.append(_safe_int(grid_col.get(qn("w:w"))) or _DOCX_MIN_COLUMN_WIDTH_TWIPS)
    return widths


def _docx_set_column_width(table_el, *, col_index: int, width_twips: int) -> list[int]:
    from docx.oxml.ns import qn

    col_count = max(_docx_table_logical_col_count(table_el), col_index)
    grid, grid_cols = _docx_ensure_table_grid(table_el, col_count)
    widths = _docx_grid_widths(table_el, col_count)
    widths[col_index - 1] = width_twips
    for index, grid_col in enumerate(grid_cols[:col_count]):
        grid_col.set(qn("w:w"), str(widths[index]))

    tbl_pr = _docx_get_or_add_tbl_pr(table_el)
    _docx_set_width_twips_el(tbl_pr, "w:tblW", sum(widths))
    _docx_set_fixed_table_layout(table_el)
    return widths


def _docx_apply_column_width_to_cells(table_el, *, col_index: int, grid_widths: list[int]) -> None:
    for row_el in _docx_table_rows(table_el):
        for logical_col, span, cell_el in _docx_row_logical_cells(row_el):
            if not (logical_col <= col_index < logical_col + span):
                continue
            tc_pr = _docx_get_or_add_tc_pr(cell_el)
            start = logical_col - 1
            end = min(start + span, len(grid_widths))
            width_twips = sum(grid_widths[start:end]) if end > start else grid_widths[col_index - 1]
            _docx_set_width_twips_el(tc_pr, "w:tcW", width_twips)


def _docx_apply_cell_width(cell_location: _DocxCellLocation, edit: StyleEdit) -> None:
    if "width_pt" not in _style_edit_supplied_fields(edit):
        return

    table_el = cell_location.table._tbl
    if "width_pt" in edit.clear_fields:
        tc_pr = _docx_get_or_add_tc_pr(cell_location.cell._tc)
        _docx_set_width_twips_el(tc_pr, "w:tcW", None, clear=True)
        return

    width_twips = _pt_to_twips(edit.width_pt)
    if width_twips is None:
        return
    grid_widths = _docx_set_column_width(table_el, col_index=cell_location.col_index, width_twips=width_twips)
    _docx_apply_column_width_to_cells(table_el, col_index=cell_location.col_index, grid_widths=grid_widths)


def _docx_set_cell_padding(tc_pr, edit: StyleEdit) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    margin_fields = {
        "padding_top_pt": "top",
        "padding_right_pt": "right",
        "padding_bottom_pt": "bottom",
        "padding_left_pt": "left",
    }
    if not any(getattr(edit, field) is not None or field in edit.clear_fields for field in margin_fields):
        return
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for field_name, side in margin_fields.items():
        side_el = tc_mar.find(qn(f"w:{side}"))
        if field_name in edit.clear_fields:
            if side_el is not None:
                tc_mar.remove(side_el)
            continue
        value = getattr(edit, field_name)
        if value is None:
            continue
        if side_el is None:
            side_el = OxmlElement(f"w:{side}")
            tc_mar.append(side_el)
        side_el.set(qn("w:type"), "dxa")
        side_el.set(qn("w:w"), str(_pt_to_twips(value) or 0))


def _docx_border_attrs(border_value: str) -> dict[str, str]:
    color = "000000"
    match = re.search(r"#([0-9A-Fa-f]{6})", border_value)
    if match:
        color = match.group(1).upper()
    width_match = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*pt", border_value)
    size = str(max(2, int(round(float(width_match.group(1)) * 8)))) if width_match else _DOCX_DEFAULT_BORDER_SIZE
    lowered = border_value.lower()
    if "dashed" in lowered:
        val = "dashed"
    elif "dotted" in lowered:
        val = "dotted"
    elif "none" in lowered:
        val = "nil"
    else:
        val = "single"
    return {"val": val, "sz": size, "color": color}


def _docx_set_cell_borders(tc_pr, edit: StyleEdit) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    border_fields = {
        "border_top": "top",
        "border_right": "right",
        "border_bottom": "bottom",
        "border_left": "left",
    }
    if not any(getattr(edit, field) is not None or field in edit.clear_fields for field in border_fields):
        return
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for field_name, side in border_fields.items():
        side_el = borders.find(qn(f"w:{side}"))
        if field_name in edit.clear_fields:
            if side_el is not None:
                borders.remove(side_el)
            continue
        value = getattr(edit, field_name)
        if value is None:
            continue
        if side_el is None:
            side_el = OxmlElement(f"w:{side}")
            borders.append(side_el)
        attrs = _docx_border_attrs(value)
        side_el.set(qn("w:val"), attrs["val"])
        side_el.set(qn("w:sz"), attrs["sz"])
        side_el.set(qn("w:space"), "0")
        side_el.set(qn("w:color"), attrs["color"])


def _apply_docx_cell_style(cell_location: _DocxCellLocation, edit: StyleEdit) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cell = cell_location.cell
    tc_pr = _docx_get_or_add_tc_pr(cell._tc)
    _docx_apply_cell_width(cell_location, edit)

    if "background" in edit.clear_fields:
        shd = tc_pr.find(qn("w:shd"))
        if shd is not None:
            tc_pr.remove(shd)
    elif edit.background is not None:
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), _docx_hex_color(edit.background))

    if "vertical_align" in edit.clear_fields:
        v_align = tc_pr.find(qn("w:vAlign"))
        if v_align is not None:
            tc_pr.remove(v_align)
    elif edit.vertical_align is not None:
        v_align = tc_pr.find(qn("w:vAlign"))
        if v_align is None:
            v_align = OxmlElement("w:vAlign")
            tc_pr.append(v_align)
        v_align.set(qn("w:val"), {"top": "top", "middle": "center", "bottom": "bottom"}[edit.vertical_align])

    if "horizontal_align" in edit.clear_fields:
        for paragraph in cell.paragraphs:
            paragraph.alignment = None
    elif edit.horizontal_align is not None:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        for paragraph in cell.paragraphs:
            paragraph.alignment = align_map[edit.horizontal_align]

    if "height_pt" in edit.clear_fields or edit.height_pt is not None:
        row = cell._tc.getparent()
        tr_pr = row.find(qn("w:trPr"))
        if tr_pr is None:
            tr_pr = OxmlElement("w:trPr")
            row.insert(0, tr_pr)
        tr_height = tr_pr.find(qn("w:trHeight"))
        if "height_pt" in edit.clear_fields:
            if tr_height is not None:
                tr_pr.remove(tr_height)
        elif edit.height_pt is not None:
            if tr_height is None:
                tr_height = OxmlElement("w:trHeight")
                tr_pr.append(tr_height)
            tr_height.set(qn("w:val"), str(_pt_to_twips(edit.height_pt) or 0))
            tr_height.set(qn("w:hRule"), "atLeast")

    _docx_set_cell_padding(tc_pr, edit)
    _docx_set_cell_borders(tc_pr, edit)


def _docx_set_table_placement(table_el, edit: StyleEdit) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    placement_fields = set(_PLACEMENT_FIELD_MAP)
    if not any(getattr(edit, field) is not None or field in edit.clear_fields for field in placement_fields):
        return

    tbl_pr = _docx_get_or_add_tbl_pr(table_el)
    tblp = tbl_pr.find(qn("w:tblpPr"))
    if edit.placement_mode == "inline":
        if tblp is not None:
            tbl_pr.remove(tblp)
        return
    if tblp is None:
        tblp = OxmlElement("w:tblpPr")
        tbl_pr.insert(0, tblp)

    if edit.x_relative_to is not None:
        tblp.set(qn("w:horzAnchor"), {"page": "page", "margin": "margin", "column": "text", "paragraph": "text", "character": "text"}[edit.x_relative_to])
    if edit.y_relative_to is not None:
        tblp.set(qn("w:vertAnchor"), {"page": "page", "margin": "margin", "paragraph": "text", "line": "text"}[edit.y_relative_to])
    if edit.x_align is not None:
        tblp.set(qn("w:tblpXSpec"), edit.x_align)
    elif edit.x_offset_pt is not None:
        tblp.set(qn("w:tblpX"), str(_pt_to_twips(edit.x_offset_pt) or 0))
    if edit.y_align is not None:
        tblp.set(qn("w:tblpYSpec"), edit.y_align)
    elif edit.y_offset_pt is not None:
        tblp.set(qn("w:tblpY"), str(_pt_to_twips(edit.y_offset_pt) or 0))
    for field_name, attr in (
        ("margin_top_pt", "topFromText"),
        ("margin_right_pt", "rightFromText"),
        ("margin_bottom_pt", "bottomFromText"),
        ("margin_left_pt", "leftFromText"),
    ):
        if field_name in edit.clear_fields:
            tblp.attrib.pop(qn(f"w:{attr}"), None)
        elif (value := getattr(edit, field_name)) is not None:
            tblp.set(qn(f"w:{attr}"), str(_pt_to_twips(value) or 0))


def _apply_docx_table_style(table_location: _DocxTableLocation, edit: StyleEdit) -> None:
    table_el = table_location.table._tbl
    _docx_set_table_placement(table_el, edit)


def _local_name(element) -> str:
    tag = getattr(element, "tag", "")
    return tag.rsplit("}", 1)[-1] if isinstance(tag, str) else ""


def _docx_child_by_local_name(parent, name: str):
    for child in list(parent):
        if _local_name(child) == name:
            return child
    return None


def _docx_set_drawing_extent(drawing_el, edit: StyleEdit) -> None:
    width = None if "width_pt" in edit.clear_fields else _pt_to_emu(edit.width_pt)
    height = None if "height_pt" in edit.clear_fields else _pt_to_emu(edit.height_pt)
    if width is None and height is None:
        return
    for element in drawing_el.iter():
        if _local_name(element) in {"extent", "ext"} and element.get("cx") is not None and element.get("cy") is not None:
            if width is not None:
                element.set("cx", str(width))
            if height is not None:
                element.set("cy", str(height))


def _docx_new_position_el(axis: str, *, relative_to: str | None, align: str | None, offset_pt: float | None):
    from docx.oxml import OxmlElement

    element = OxmlElement(f"wp:position{axis}")
    element.set("relativeFrom", relative_to or ("column" if axis == "H" else "paragraph"))
    if offset_pt is not None:
        child = OxmlElement("wp:posOffset")
        child.text = str(_pt_to_emu(offset_pt) or 0)
    else:
        child = OxmlElement("wp:align")
        child.text = align or ("left" if axis == "H" else "top")
    element.append(child)
    return element


def _docx_new_wrap_el(edit: StyleEdit):
    from docx.oxml import OxmlElement

    wrap = edit.wrap or "square"
    wrap_tag = {
        "none": "wrapNone",
        "square": "wrapSquare",
        "tight": "wrapTight",
        "through": "wrapThrough",
        "top_bottom": "wrapTopAndBottom",
        "behind_text": "wrapNone",
        "in_front_of_text": "wrapNone",
    }[wrap]
    element = OxmlElement(f"wp:{wrap_tag}")
    if wrap_tag in {"wrapSquare", "wrapTight", "wrapThrough"}:
        element.set("wrapText", {"both_sides": "bothSides", "left": "left", "right": "right", "largest": "largest"}.get(edit.text_flow or "both_sides", "bothSides"))
    for field_name, attr in (
        ("margin_top_pt", "distT"),
        ("margin_right_pt", "distR"),
        ("margin_bottom_pt", "distB"),
        ("margin_left_pt", "distL"),
    ):
        value = getattr(edit, field_name)
        if value is not None and wrap_tag in {"wrapSquare", "wrapTight", "wrapThrough", "wrapTopAndBottom"}:
            element.set(attr, str(_pt_to_emu(value) or 0))
    if wrap_tag in {"wrapTight", "wrapThrough"}:
        polygon = OxmlElement("wp:wrapPolygon")
        polygon.set("edited", "0")
        start = OxmlElement("wp:start")
        start.set("x", "0")
        start.set("y", "0")
        polygon.append(start)
        for x, y in (("0", "21600"), ("21600", "21600"), ("21600", "0")):
            point = OxmlElement("wp:lineTo")
            point.set("x", x)
            point.set("y", y)
            polygon.append(point)
        element.append(polygon)
    return element


def _docx_convert_inline_to_anchor(inline_el, edit: StyleEdit):
    from docx.oxml import OxmlElement

    anchor = OxmlElement("wp:anchor")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", str(edit.z_order if edit.z_order is not None else 0))
    anchor.set("behindDoc", "1" if edit.wrap == "behind_text" else "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1" if edit.allow_overlap else "0")
    for field_name, attr in (
        ("margin_top_pt", "distT"),
        ("margin_right_pt", "distR"),
        ("margin_bottom_pt", "distB"),
        ("margin_left_pt", "distL"),
    ):
        value = getattr(edit, field_name)
        anchor.set(attr, str(_pt_to_emu(value) or 0))

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)
    anchor.append(_docx_new_position_el("H", relative_to=edit.x_relative_to, align=edit.x_align, offset_pt=edit.x_offset_pt))
    anchor.append(_docx_new_position_el("V", relative_to=edit.y_relative_to, align=edit.y_align, offset_pt=edit.y_offset_pt))

    for child_name in ("extent", "effectExtent"):
        child = _docx_child_by_local_name(inline_el, child_name)
        if child is not None:
            anchor.append(deepcopy(child))
    anchor.append(_docx_new_wrap_el(edit))
    for child_name in ("docPr", "cNvGraphicFramePr", "graphic"):
        child = _docx_child_by_local_name(inline_el, child_name)
        if child is not None:
            anchor.append(deepcopy(child))

    parent = inline_el.getparent()
    parent.replace(inline_el, anchor)
    return anchor


def _docx_convert_anchor_to_inline(anchor_el):
    from docx.oxml import OxmlElement

    inline = OxmlElement("wp:inline")
    for child_name in ("extent", "effectExtent", "docPr", "cNvGraphicFramePr", "graphic"):
        child = _docx_child_by_local_name(anchor_el, child_name)
        if child is not None:
            inline.append(deepcopy(child))
    parent = anchor_el.getparent()
    parent.replace(anchor_el, inline)
    return inline


def _docx_update_anchor_position(anchor_el, edit: StyleEdit) -> None:
    for old in list(anchor_el):
        if _local_name(old) in {"positionH", "positionV", "wrapNone", "wrapSquare", "wrapTight", "wrapThrough", "wrapTopAndBottom"}:
            anchor_el.remove(old)
    simple_pos = _docx_child_by_local_name(anchor_el, "simplePos")
    insert_at = list(anchor_el).index(simple_pos) + 1 if simple_pos is not None else 0
    anchor_el.insert(insert_at, _docx_new_position_el("H", relative_to=edit.x_relative_to, align=edit.x_align, offset_pt=edit.x_offset_pt))
    anchor_el.insert(insert_at + 1, _docx_new_position_el("V", relative_to=edit.y_relative_to, align=edit.y_align, offset_pt=edit.y_offset_pt))
    effect_extent = _docx_child_by_local_name(anchor_el, "effectExtent")
    extent = _docx_child_by_local_name(anchor_el, "extent")
    wrap_anchor = effect_extent or extent
    wrap_index = list(anchor_el).index(wrap_anchor) + 1 if wrap_anchor is not None else insert_at + 2
    anchor_el.insert(wrap_index, _docx_new_wrap_el(edit))
    if edit.z_order is not None:
        anchor_el.set("relativeHeight", str(edit.z_order))
    if edit.allow_overlap is not None:
        anchor_el.set("allowOverlap", "1" if edit.allow_overlap else "0")
    if edit.wrap in {"behind_text", "in_front_of_text"}:
        anchor_el.set("behindDoc", "1" if edit.wrap == "behind_text" else "0")


def _apply_docx_image_style(image_location: _DocxImageLocation, edit: StyleEdit) -> None:
    drawing_el = image_location.drawing_el
    if edit.placement_mode == "floating" and _local_name(drawing_el) == "inline":
        drawing_el = _docx_convert_inline_to_anchor(drawing_el, edit)
        image_location.drawing_el = drawing_el
    elif edit.placement_mode == "inline" and _local_name(drawing_el) == "anchor":
        drawing_el = _docx_convert_anchor_to_inline(drawing_el)
        image_location.drawing_el = drawing_el

    _docx_set_drawing_extent(drawing_el, edit)
    if _local_name(drawing_el) == "anchor" and any(
        getattr(edit, field) is not None
        for field in _PLACEMENT_FIELD_MAP
        if field not in {"placement_mode"}
    ):
        _docx_update_anchor_position(drawing_el, edit)


def _apply_docx_style_edit(doc, edit: StyleEdit, result: _EditEngineResult) -> None:
    index = _build_docx_structural_index(doc)
    if edit.target_kind == "run" and (location := index.runs.get(edit.target_id)):
        _apply_docx_run_style(location.run, edit, result)
    elif edit.target_kind == "paragraph" and (location := index.paragraphs.get(edit.target_id)):
        _apply_docx_paragraph_style(location.paragraph, edit)
    elif edit.target_kind == "cell" and (location := index.cells.get(edit.target_id)):
        _apply_docx_cell_style(location, edit)
    elif edit.target_kind == "table" and (location := index.tables.get(edit.target_id)):
        _apply_docx_table_style(location, edit)
    elif edit.target_kind == "image" and (location := index.images.get(edit.target_id)):
        _apply_docx_image_style(location, edit)
    else:
        raise EditValidationError(
            f"{edit.target_kind.capitalize()} style target does not exist: {edit.target_id}.",
            code="target_not_found",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )
    _append_unique(result.modified_target_ids, edit.target_id)
    result.styles_applied += 1


def _apply_docx_structural_operation(doc, operation: StructuralEdit, result: _EditEngineResult) -> None:
    from docx.oxml.ns import qn

    index = _build_docx_structural_index(doc)

    if operation.operation == "insert_paragraph":
        paragraph_el = _docx_paragraph_el(operation.text or "")
        if paragraph_location := index.paragraphs.get(operation.target_id):
            if operation.position == "before":
                paragraph_location.paragraph._p.addprevious(paragraph_el)
            elif operation.position == "after":
                paragraph_location.paragraph._p.addnext(paragraph_el)
            else:
                raise EditValidationError("insert_paragraph with a paragraph target requires before/after.", code="invalid_position")
        elif cell_location := index.cells.get(operation.target_id):
            _docx_insert_paragraph_in_cell(cell_location.cell, paragraph_el, position=operation.position)
        else:
            raise EditValidationError("insert_paragraph target must be a paragraph or cell.", code="target_not_found", target_id=operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_paragraph":
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is None:
            raise EditValidationError("Paragraph does not exist.", code="target_not_found", target_kind="paragraph", target_id=operation.target_id)
        current_text = paragraph_location.paragraph.text or ""
        current_text_hash = _text_hash(current_text)
        if operation.expected_text_hash is not None and current_text_hash != operation.expected_text_hash:
            raise EditValidationError(
                "Paragraph text hash mismatch.",
                code="text_hash_mismatch",
                target_kind="paragraph",
                target_id=operation.target_id,
                operation=operation.operation,
                current_text=current_text,
                expected_text_hash=operation.expected_text_hash,
                current_text_hash=current_text_hash,
            )
        parent = paragraph_location.paragraph._p.getparent()
        parent.remove(paragraph_location.paragraph._p)
        _ensure_docx_container_has_paragraph(parent)
        result.operations_applied += 1
        return

    if operation.operation == "insert_run":
        run_el = _docx_text_run_el(operation.text or "")
        if run_location := index.runs.get(operation.target_id):
            if operation.position == "before":
                run_location.run._r.addprevious(run_el)
            elif operation.position == "after":
                run_location.run._r.addnext(run_el)
            else:
                raise EditValidationError("insert_run with a run target requires before/after.", code="invalid_position")
        elif paragraph_location := index.paragraphs.get(operation.target_id):
            paragraph_el = paragraph_location.paragraph._p
            if operation.position in {"start", "before"}:
                insert_index = 1 if len(paragraph_el) and paragraph_el[0].tag.endswith("}pPr") else 0
                paragraph_el.insert(insert_index, run_el)
            elif operation.position in {"end", "after"}:
                paragraph_el.append(run_el)
            else:
                raise EditValidationError("Invalid insert_run position.", code="invalid_position")
        else:
            raise EditValidationError("insert_run target must be a run or paragraph.", code="target_not_found", target_id=operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_run":
        run_location = index.runs.get(operation.target_id)
        if run_location is None:
            raise EditValidationError("Run does not exist.", code="target_not_found", target_kind="run", target_id=operation.target_id)
        current_text = run_location.run.text or ""
        current_text_hash = _text_hash(current_text)
        if operation.expected_text_hash is not None and current_text_hash != operation.expected_text_hash:
            raise EditValidationError(
                "Run text hash mismatch.",
                code="text_hash_mismatch",
                target_kind="run",
                target_id=operation.target_id,
                operation=operation.operation,
                current_text=current_text,
                expected_text_hash=operation.expected_text_hash,
                current_text_hash=current_text_hash,
            )
        run_location.run._r.getparent().remove(run_location.run._r)
        result.operations_applied += 1
        return

    if operation.operation == "insert_table":
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is None:
            raise EditValidationError("insert_table target must be a paragraph.", code="target_not_found", target_kind="paragraph", target_id=operation.target_id)
        table_el = _docx_table_el(_normalize_table_rows(operation.rows))
        if operation.position == "before":
            paragraph_location.paragraph._p.addprevious(table_el)
        elif operation.position == "after":
            paragraph_location.paragraph._p.addnext(table_el)
        else:
            raise EditValidationError("insert_table requires before/after.", code="invalid_position")
        result.operations_applied += 1
        return

    if operation.operation == "remove_table":
        table_location = index.tables.get(operation.target_id)
        if table_location is None:
            raise EditValidationError("Table does not exist.", code="target_not_found", target_kind="table", target_id=operation.target_id)
        table_location.table._tbl.getparent().remove(table_location.table._tbl)
        result.operations_applied += 1
        return

    if operation.operation == "set_cell_text":
        cell_location = index.cells.get(operation.target_id)
        if cell_location is None:
            raise EditValidationError("Cell does not exist.", code="target_not_found", target_kind="cell", target_id=operation.target_id)
        current_text = cell_location.cell.text or ""
        current_text_hash = _text_hash(current_text)
        if operation.expected_text_hash is not None and current_text_hash != operation.expected_text_hash:
            raise EditValidationError(
                "Cell text hash mismatch.",
                code="text_hash_mismatch",
                target_kind="cell",
                target_id=operation.target_id,
                operation=operation.operation,
                current_text=current_text,
                expected_text_hash=operation.expected_text_hash,
                current_text_hash=current_text_hash,
            )
        _docx_set_cell_text(cell_location.cell, operation.text or "")
        result.operations_applied += 1
        return

    if operation.operation in {"insert_table_row", "remove_table_row"}:
        table, row_index = _resolve_docx_table_axis(index, operation, axis="row")
        rows = _docx_table_rows(table._tbl)
        if row_index < 1 or row_index > len(rows):
            raise EditValidationError("Table row index is out of bounds.", code="index_out_of_bounds")
        if operation.operation == "remove_table_row":
            if len(rows) <= 1:
                raise EditValidationError("Cannot remove the only table row.", code="invalid_table_shape")
            table._tbl.remove(rows[row_index - 1])
        else:
            col_count = _docx_table_col_count(table._tbl)
            values = operation.values or ["" for _ in range(col_count)]
            if len(values) != col_count:
                raise EditValidationError("Inserted row values must match table column count.", code="invalid_table_shape")
            template_row = rows[row_index - 1]
            template_widths: list[int | None] = []
            for cell in _docx_row_cells(template_row):
                tc_pr = cell.find(qn("w:tcPr"))
                tc_width = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
                template_widths.append(_safe_int(tc_width.get(qn("w:w"))) if tc_width is not None else None)
            default_widths = _docx_default_column_widths_twips(col_count)
            col_widths = [
                (template_widths[index] if index < len(template_widths) else None) or default_widths[index]
                for index in range(col_count)
            ]
            new_row = _docx_row_el(values, col_widths=col_widths, template_row=template_row)
            if operation.position in {"after", "end"}:
                rows[row_index - 1].addnext(new_row)
            else:
                rows[row_index - 1].addprevious(new_row)
        result.operations_applied += 1
        return

    if operation.operation in {"insert_table_column", "remove_table_column"}:
        table, column_index = _resolve_docx_table_axis(index, operation, axis="column")
        rows = _docx_table_rows(table._tbl)
        col_count = _docx_table_col_count(table._tbl)
        if column_index < 1 or column_index > col_count:
            raise EditValidationError("Table column index is out of bounds.", code="index_out_of_bounds")
        if operation.operation == "remove_table_column":
            if col_count <= 1:
                raise EditValidationError("Cannot remove the only table column.", code="invalid_table_shape")
            for row in rows:
                cells = _docx_row_cells(row)
                if len(cells) >= column_index:
                    row.remove(cells[column_index - 1])
            _docx_adjust_table_grid(table._tbl, column_index=column_index, operation="remove")
        else:
            values = operation.values or ["" for _ in rows]
            if len(values) != len(rows):
                raise EditValidationError("Inserted column values must match table row count.", code="invalid_table_shape")
            grid = table._tbl.find(qn("w:tblGrid"))
            grid_cols = list(grid.findall(qn("w:gridCol"))) if grid is not None else []
            inserted_width = None
            if 1 <= column_index <= len(grid_cols):
                inserted_width = _safe_int(grid_cols[column_index - 1].get(qn("w:w")))
            for row_index, row in enumerate(rows):
                cells = _docx_row_cells(row)
                template_tc_pr = cells[column_index - 1].find(qn("w:tcPr")) if len(cells) >= column_index else None
                new_cell = _docx_cell_el(values[row_index], width_twips=inserted_width, tc_pr_template=template_tc_pr)
                if operation.position in {"after", "end"}:
                    cells[column_index - 1].addnext(new_cell)
                else:
                    cells[column_index - 1].addprevious(new_cell)
            _docx_adjust_table_grid(
                table._tbl,
                column_index=column_index,
                operation="insert",
                position=operation.position,
                width_twips=inserted_width,
            )
        result.operations_applied += 1
        return

    raise EditValidationError(f"Unsupported structural operation: {operation.operation!r}.", code="invalid_operation")


@dataclass
class _HwpxParagraphLocation:
    element: ET.Element
    parent: ET.Element
    path: str


@dataclass
class _HwpxRunLocation:
    element: ET.Element
    parent: ET.Element
    path: str


@dataclass
class _HwpxTableLocation:
    element: ET.Element
    parent: ET.Element
    path: str


@dataclass
class _HwpxImageLocation:
    element: ET.Element
    parent: ET.Element
    path: str


@dataclass
class _HwpxCellLocation:
    element: ET.Element
    table: ET.Element
    row_index: int
    col_index: int
    path: str


class _HwpxStructuralIndex:
    def __init__(self) -> None:
        self.paragraphs: dict[str, _HwpxParagraphLocation] = {}
        self.runs: dict[str, _HwpxRunLocation] = {}
        self.tables: dict[str, _HwpxTableLocation] = {}
        self.images: dict[str, _HwpxImageLocation] = {}
        self.cells: dict[str, _HwpxCellLocation] = {}


def _build_hwpx_structural_index(archive: _EditableHwpxArchive) -> _HwpxStructuralIndex:
    index = _HwpxStructuralIndex()

    def register_paragraph(paragraph_el: ET.Element, parent_el: ET.Element, paragraph_path: str) -> None:
        index.paragraphs[_anchored_node_id("paragraph", paragraph_path)] = _HwpxParagraphLocation(
            element=paragraph_el,
            parent=parent_el,
            path=paragraph_path,
        )
        image_index = 0
        for run_index, run_el in enumerate(paragraph_el.findall(f"{_HP}run"), start=1):
            run_path = f"{paragraph_path}.r{run_index}"
            index.runs[_anchored_node_id("run", run_path)] = _HwpxRunLocation(
                element=run_el,
                parent=paragraph_el,
                path=run_path,
            )
            for pic_el in run_el.findall(f"{_HP}pic"):
                image_index += 1
                image_path = f"{paragraph_path}.img{image_index}"
                index.images[_anchored_node_id("image", image_path)] = _HwpxImageLocation(
                    element=pic_el,
                    parent=run_el,
                    path=image_path,
                )

    def walk_table(table_el: ET.Element, parent_el: ET.Element, table_path: str) -> None:
        index.tables[_anchored_node_id("table", table_path)] = _HwpxTableLocation(
            element=table_el,
            parent=parent_el,
            path=table_path,
        )
        for row_index, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
            for col_index, cell_el in _logical_table_cells(row_el):
                cell_path = f"{table_path}.tr{row_index}.tc{col_index}"
                index.cells[_anchored_node_id("cell", cell_path)] = _HwpxCellLocation(
                    element=cell_el,
                    table=table_el,
                    row_index=row_index,
                    col_index=col_index,
                    path=cell_path,
                )
                sub_list = cell_el.find(f"{_HP}subList")
                if sub_list is None:
                    continue
                for paragraph_index, paragraph_el in enumerate([child for child in list(sub_list) if child.tag == f"{_HP}p"], start=1):
                    paragraph_path = f"{cell_path}.p{paragraph_index}"
                    register_paragraph(paragraph_el, sub_list, paragraph_path)
                    for table_index, nested_table in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                        walk_table(nested_table, paragraph_el, f"{paragraph_path}.tbl{table_index}")

    for section_index, section in enumerate(archive.section_entries, start=1):
        for paragraph_index, paragraph_el in enumerate(_iter_section_paragraphs(section.root), start=1):
            paragraph_path = f"s{section_index}.p{paragraph_index}"
            register_paragraph(paragraph_el, section.root, paragraph_path)
            for table_index, table_el in enumerate(_iter_paragraph_tables(paragraph_el), start=1):
                walk_table(table_el, paragraph_el, f"{paragraph_path}.r1.tbl{table_index}")

    return index


def _hwpx_find_or_create_child(parent: ET.Element, name: str) -> ET.Element:
    child = parent.find(f"{_HP}{name}")
    if child is None:
        child = ET.Element(f"{_HP}{name}")
        if name in {"sz", "pos", "outMargin"}:
            order = {"sz": 0, "pos": 1, "outMargin": 2}
            insert_at = len(parent)
            for index, existing in enumerate(list(parent)):
                existing_name = existing.tag.rsplit("}", 1)[-1]
                existing_order = order.get(existing_name)
                if existing_order is None or existing_order > order[name]:
                    insert_at = index
                    break
            parent.insert(insert_at, child)
        elif name in {"subList", "cellAddr", "cellSpan", "cellSz", "cellMargin"}:
            order = {"subList": 0, "cellAddr": 1, "cellSpan": 2, "cellSz": 3, "cellMargin": 4}
            insert_at = len(parent)
            for index, existing in enumerate(list(parent)):
                existing_name = existing.tag.rsplit("}", 1)[-1]
                existing_order = order.get(existing_name)
                if existing_order is None or existing_order > order[name]:
                    insert_at = index
                    break
            parent.insert(insert_at, child)
        else:
            parent.append(child)
    return child


def _hwpx_find_or_create_hh_child(parent: ET.Element, name: str) -> ET.Element:
    child = parent.find(f"{_HH}{name}")
    if child is None:
        child = _hh_el(name)
        parent.append(child)
    return child


def _hwpx_find_or_create_hc_child(parent: ET.Element, name: str) -> ET.Element:
    child = parent.find(f"{_HC}{name}")
    if child is None:
        child = _hc_el(name)
        parent.append(child)
    return child


def _hwpx_header_or_raise(archive: _EditableHwpxArchive, edit: StyleEdit) -> _EditableHwpxHeader:
    if archive.header_entry is None:
        raise EditValidationError(
            "HWPX native style write-back requires Contents/header.xml.",
            code="invalid_style",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )
    return archive.header_entry


def _ensure_hwpx_header_style_container(
    header: _EditableHwpxHeader,
    container_name: str,
) -> ET.Element:
    ref_list = _ensure_hwpx_ref_list(header)
    container = ref_list.find(f"{_HH}{container_name}")
    if container is not None:
        return container
    container = _hh_el(container_name, {"itemCnt": "0"})
    ref_list.append(container)
    return container


def _hwpx_next_header_ref_id(container: ET.Element, child_name: str) -> str:
    existing_ids = [
        item_id
        for child in container.findall(f"{_HH}{child_name}")
        if (item_id := _safe_int(child.get("id"))) is not None
    ]
    return str((max(existing_ids) + 1) if existing_ids else 0)


def _hwpx_header_child_by_id(container: ET.Element, child_name: str, child_id: str | None) -> ET.Element | None:
    if child_id is None:
        return None
    for child in container.findall(f"{_HH}{child_name}"):
        if child.get("id") == child_id:
            return child
    return None


def _hwpx_default_char_pr(char_pr_id: str) -> ET.Element:
    char_pr = _hh_el(
        "charPr",
        {
            "id": char_pr_id,
            "height": "1000",
            "textColor": "#000000",
            "shadeColor": "none",
            "useFontSpace": "0",
            "useKerning": "0",
            "symMark": "NONE",
        },
    )
    char_pr.append(_hh_el("fontRef", {"hangul": "0", "latin": "0", "hanja": "0", "japanese": "0", "other": "0", "symbol": "0", "user": "0"}))
    char_pr.append(_hh_el("ratio", {"hangul": "100", "latin": "100", "hanja": "100", "japanese": "100", "other": "100", "symbol": "100", "user": "100"}))
    char_pr.append(_hh_el("spacing", {"hangul": "0", "latin": "0", "hanja": "0", "japanese": "0", "other": "0", "symbol": "0", "user": "0"}))
    char_pr.append(_hh_el("relSz", {"hangul": "100", "latin": "100", "hanja": "100", "japanese": "100", "other": "100", "symbol": "100", "user": "100"}))
    char_pr.append(_hh_el("offset", {"hangul": "0", "latin": "0", "hanja": "0", "japanese": "0", "other": "0", "symbol": "0", "user": "0"}))
    char_pr.append(_hh_el("underline", {"type": "NONE", "shape": "SOLID", "color": "#000000"}))
    char_pr.append(_hh_el("strikeout", {"shape": "NONE", "color": "#000000"}))
    char_pr.append(_hh_el("outline", {"type": "NONE"}))
    char_pr.append(_hh_el("shadow", {"type": "NONE", "color": "#C0C0C0", "offsetX": "10", "offsetY": "10"}))
    return char_pr


def _hwpx_default_para_pr(para_pr_id: str) -> ET.Element:
    para_pr = _hh_el(
        "paraPr",
        {
            "id": para_pr_id,
            "tabPrIDRef": "0",
            "condense": "0",
            "fontLineHeight": "0",
            "snapToGrid": "1",
            "suppressLineNumbers": "0",
            "checked": "0",
        },
    )
    para_pr.append(_hh_el("align", {"horizontal": "LEFT", "vertical": "BASELINE"}))
    margin = _hh_el("margin")
    for name in ("intent", "left", "right", "prev", "next"):
        margin.append(_hc_el(name, {"value": "0", "unit": "HWPUNIT"}))
    para_pr.append(margin)
    return para_pr


def _hwpx_clone_or_create_header_ref(
    archive: _EditableHwpxArchive,
    *,
    container_name: str,
    child_name: str,
    source_id: str | None,
    edit: StyleEdit,
    default_factory: Callable[[str], ET.Element],
) -> tuple[ET.Element, str]:
    header = _hwpx_header_or_raise(archive, edit)
    container = _ensure_hwpx_header_style_container(header, container_name)
    new_id = _hwpx_next_header_ref_id(container, child_name)
    source = _hwpx_header_child_by_id(container, child_name, source_id)
    cloned = deepcopy(source) if source is not None else default_factory(new_id)
    cloned.set("id", new_id)
    container.append(cloned)
    container.set("itemCnt", str(len(container.findall(f"{_HH}{child_name}"))))
    return cloned, new_id


def _hwpx_clone_char_pr_for_run(
    archive: _EditableHwpxArchive,
    run_el: ET.Element,
    edit: StyleEdit,
) -> ET.Element:
    char_pr, new_id = _hwpx_clone_or_create_header_ref(
        archive,
        container_name="charProperties",
        child_name="charPr",
        source_id=run_el.get("charPrIDRef"),
        edit=edit,
        default_factory=_hwpx_default_char_pr,
    )
    run_el.set("charPrIDRef", new_id)
    return char_pr


def _hwpx_clone_para_pr_for_paragraph(
    archive: _EditableHwpxArchive,
    paragraph_el: ET.Element,
    edit: StyleEdit,
) -> ET.Element:
    para_pr, new_id = _hwpx_clone_or_create_header_ref(
        archive,
        container_name="paraProperties",
        child_name="paraPr",
        source_id=paragraph_el.get("paraPrIDRef"),
        edit=edit,
        default_factory=_hwpx_default_para_pr,
    )
    paragraph_el.set("paraPrIDRef", new_id)
    return para_pr


def _hwpx_style_unsupported_fields(edit: StyleEdit) -> set[str]:
    supplied = _style_edit_supplied_fields(edit)
    if edit.target_kind == "run":
        supported = {
            "bold",
            "italic",
            "underline",
            "strikethrough",
            "color",
            "font_size_pt",
        }
        return supplied - supported
    if edit.target_kind == "paragraph":
        supported = {
            "paragraph_align",
            "left_indent_pt",
            "right_indent_pt",
            "first_line_indent_pt",
            "hanging_indent_pt",
        }
        return supplied - supported
    if edit.target_kind == "cell":
        supported = {
            "width_pt",
            "height_pt",
            "background",
            "vertical_align",
            "horizontal_align",
            "padding_top_pt",
            "padding_right_pt",
            "padding_bottom_pt",
            "padding_left_pt",
            "border_top",
            "border_right",
            "border_bottom",
            "border_left",
        }
        return supplied - supported
    if edit.target_kind == "table":
        return supplied - set(_PLACEMENT_FIELD_MAP)
    if edit.target_kind == "image":
        return supplied - ({"width_pt", "height_pt"} | set(_PLACEMENT_FIELD_MAP))
    return supplied


def _hwpx_set_size_el(object_el: ET.Element, edit: StyleEdit) -> None:
    if not any(field in _style_edit_supplied_fields(edit) for field in ("width_pt", "height_pt")):
        return
    size_el = _hwpx_find_or_create_child(object_el, "sz")
    if "width_pt" in edit.clear_fields:
        size_el.attrib.pop("width", None)
    elif edit.width_pt is not None:
        size_el.set("width", str(_pt_to_hwpunit(edit.width_pt) or 0))
        size_el.set("widthRelTo", "ABSOLUTE")
    if "height_pt" in edit.clear_fields:
        size_el.attrib.pop("height", None)
    elif edit.height_pt is not None:
        size_el.set("height", str(_pt_to_hwpunit(edit.height_pt) or 0))
        size_el.set("heightRelTo", "ABSOLUTE")
    if "protect" not in size_el.attrib:
        size_el.set("protect", "0")


def _hwpx_cell_colspan(cell_el: ET.Element) -> int:
    cell_span = cell_el.find(f"{_HP}cellSpan")
    return max(_safe_int(cell_span.get("colSpan")) if cell_span is not None else 1, 1)


def _hwpx_cell_margin_units(cell_el: ET.Element, table_el: ET.Element) -> dict[str, int]:
    margin = cell_el.find(f"{_HP}cellMargin")
    if margin is None:
        margin = table_el.find(f"{_HP}inMargin")
    return {
        side: _safe_int(margin.get(side)) if margin is not None else 0
        for side in ("left", "right", "top", "bottom")
    }


def _hwpx_update_cell_text_area(cell_el: ET.Element, table_el: ET.Element) -> None:
    sub_list = cell_el.find(f"{_HP}subList")
    size_el = cell_el.find(f"{_HP}cellSz")
    if sub_list is None or size_el is None:
        return
    margins = _hwpx_cell_margin_units(cell_el, table_el)
    width = _safe_int(size_el.get("width"))
    height = _safe_int(size_el.get("height"))
    if width is not None:
        sub_list.set("textWidth", str(max(width - (margins["left"] or 0) - (margins["right"] or 0), 0)))
    if height is not None:
        sub_list.set("textHeight", str(max(height - (margins["top"] or 0) - (margins["bottom"] or 0), 0)))


def _hwpx_set_cell_size_attrs(
    cell_el: ET.Element,
    table_el: ET.Element,
    *,
    width: int | None = None,
    height: int | None = None,
    clear_width: bool = False,
    clear_height: bool = False,
) -> None:
    size_el = _hwpx_find_or_create_child(cell_el, "cellSz")
    if clear_width:
        size_el.attrib.pop("width", None)
    elif width is not None:
        size_el.set("width", str(width))
    if clear_height:
        size_el.attrib.pop("height", None)
    elif height is not None:
        size_el.set("height", str(height))
    _hwpx_update_cell_text_area(cell_el, table_el)


def _hwpx_apply_cell_width(location: _HwpxCellLocation, edit: StyleEdit) -> None:
    if "width_pt" not in _style_edit_supplied_fields(edit):
        return

    table_el = location.table
    target_colspan = _hwpx_cell_colspan(location.element)
    width = None if "width_pt" in edit.clear_fields else _pt_to_hwpunit(edit.width_pt)
    for row_el in _hwpx_table_rows(table_el):
        for logical_col, cell_el in _logical_table_cells(row_el):
            if logical_col != location.col_index:
                continue
            if cell_el is not location.element and (target_colspan > 1 or _hwpx_cell_colspan(cell_el) > 1):
                continue
            _hwpx_set_cell_size_attrs(
                cell_el,
                table_el,
                width=width,
                clear_width="width_pt" in edit.clear_fields,
            )


def _hwpx_apply_cell_height(location: _HwpxCellLocation, edit: StyleEdit) -> None:
    if "height_pt" not in _style_edit_supplied_fields(edit):
        return

    rows = _hwpx_table_rows(location.table)
    if location.row_index < 1 or location.row_index > len(rows):
        return
    height = None if "height_pt" in edit.clear_fields else _pt_to_hwpunit(edit.height_pt)
    for cell_el in _hwpx_row_cells(rows[location.row_index - 1]):
        _hwpx_set_cell_size_attrs(
            cell_el,
            location.table,
            height=height,
            clear_height="height_pt" in edit.clear_fields,
        )


def _hwpx_set_cell_size_el(location: _HwpxCellLocation, edit: StyleEdit) -> None:
    if not any(field in _style_edit_supplied_fields(edit) for field in ("width_pt", "height_pt")):
        return
    _hwpx_apply_cell_width(location, edit)
    _hwpx_apply_cell_height(location, edit)


def _hwpx_set_cell_margin_el(cell_el: ET.Element, edit: StyleEdit) -> None:
    margin_fields = {
        "padding_top_pt": "top",
        "padding_right_pt": "right",
        "padding_bottom_pt": "bottom",
        "padding_left_pt": "left",
    }
    if not any(field in _style_edit_supplied_fields(edit) for field in margin_fields):
        return
    margin_el = _hwpx_find_or_create_child(cell_el, "cellMargin")
    for field_name, attr in margin_fields.items():
        if field_name in edit.clear_fields:
            margin_el.attrib.pop(attr, None)
        elif (value := getattr(edit, field_name)) is not None:
            margin_el.set(attr, str(_pt_to_hwpunit(value) or 0))


def _hwpx_set_presence_style_child(parent: ET.Element, name: str, value: bool | None, *, clear: bool) -> None:
    child = parent.find(f"{_HH}{name}")
    if clear or value is False:
        if child is not None:
            parent.remove(child)
        return
    if value is True and child is None:
        parent.append(_hh_el(name))


def _hwpx_apply_run_style(archive: _EditableHwpxArchive, location: _HwpxRunLocation, edit: StyleEdit) -> None:
    char_pr = _hwpx_clone_char_pr_for_run(archive, location.element, edit)
    _hwpx_set_presence_style_child(char_pr, "bold", edit.bold, clear="bold" in edit.clear_fields)
    _hwpx_set_presence_style_child(char_pr, "italic", edit.italic, clear="italic" in edit.clear_fields)

    if "underline" in edit.clear_fields or edit.underline is False:
        underline = char_pr.find(f"{_HH}underline")
        if underline is not None:
            underline.set("type", "NONE")
    elif edit.underline is True:
        underline = _hwpx_find_or_create_hh_child(char_pr, "underline")
        underline.set("type", "BOTTOM")
        underline.set("shape", "SOLID")
        underline.set("color", "#000000")

    if "strikethrough" in edit.clear_fields or edit.strikethrough is False:
        strikeout = char_pr.find(f"{_HH}strikeout")
        if strikeout is not None:
            strikeout.set("type", "NONE")
            strikeout.set("shape", "NONE")
    elif edit.strikethrough is True:
        strikeout = _hwpx_find_or_create_hh_child(char_pr, "strikeout")
        strikeout.set("type", "CONTINUOUS")
        strikeout.set("shape", "SOLID")
        strikeout.set("color", "#000000")

    if "color" in edit.clear_fields:
        char_pr.set("textColor", "#000000")
    elif edit.color is not None:
        char_pr.set("textColor", f"#{_docx_hex_color(edit.color)}")

    if "font_size_pt" in edit.clear_fields:
        char_pr.attrib.pop("height", None)
    elif edit.font_size_pt is not None:
        char_pr.set("height", str(int(round(edit.font_size_pt * 100))))


def _hwpx_align_value(value: str | None) -> str | None:
    if value is None:
        return None
    return {
        "left": "LEFT",
        "center": "CENTER",
        "right": "RIGHT",
        "justify": "JUSTIFY",
    }[value]


def _hwpx_set_para_align(para_pr: ET.Element, align: str | None, *, clear: bool = False) -> None:
    align_el = _hwpx_find_or_create_hh_child(para_pr, "align")
    if clear:
        align_el.set("horizontal", "LEFT")
    elif align is not None:
        align_el.set("horizontal", _hwpx_align_value(align) or "LEFT")
    if "vertical" not in align_el.attrib:
        align_el.set("vertical", "BASELINE")


def _hwpx_set_margin_value(margin_el: ET.Element, name: str, value_pt: float | None, *, clear: bool = False) -> None:
    child = _hwpx_find_or_create_hc_child(margin_el, name)
    child.set("unit", "HWPUNIT")
    child.set("value", "0" if clear or value_pt is None else str(_pt_to_hwpunit(value_pt) or 0))


def _hwpx_apply_para_margin_edits(para_pr: ET.Element, edit: StyleEdit) -> None:
    margin_fields = {
        "left_indent_pt",
        "right_indent_pt",
        "first_line_indent_pt",
        "hanging_indent_pt",
    }
    if not any(field in _style_edit_supplied_fields(edit) for field in margin_fields):
        return

    margin_elements = para_pr.findall(f".//{_HH}margin")
    if not margin_elements:
        margin_elements = [_hwpx_find_or_create_hh_child(para_pr, "margin")]

    for margin_el in margin_elements:
        if "left_indent_pt" in edit.clear_fields:
            _hwpx_set_margin_value(margin_el, "left", None, clear=True)
        elif edit.left_indent_pt is not None:
            _hwpx_set_margin_value(margin_el, "left", edit.left_indent_pt)

        if "right_indent_pt" in edit.clear_fields:
            _hwpx_set_margin_value(margin_el, "right", None, clear=True)
        elif edit.right_indent_pt is not None:
            _hwpx_set_margin_value(margin_el, "right", edit.right_indent_pt)

        if "first_line_indent_pt" in edit.clear_fields or "hanging_indent_pt" in edit.clear_fields:
            _hwpx_set_margin_value(margin_el, "intent", None, clear=True)
        if edit.first_line_indent_pt is not None:
            _hwpx_set_margin_value(margin_el, "intent", edit.first_line_indent_pt)
        if edit.hanging_indent_pt is not None:
            _hwpx_set_margin_value(margin_el, "intent", -edit.hanging_indent_pt)


def _hwpx_apply_paragraph_style(archive: _EditableHwpxArchive, location: _HwpxParagraphLocation, edit: StyleEdit) -> None:
    para_pr = _hwpx_clone_para_pr_for_paragraph(archive, location.element, edit)
    if "paragraph_align" in edit.clear_fields:
        _hwpx_set_para_align(para_pr, None, clear=True)
    elif edit.paragraph_align is not None:
        _hwpx_set_para_align(para_pr, edit.paragraph_align)
    _hwpx_apply_para_margin_edits(para_pr, edit)


def _hwpx_clone_border_fill_for_cell(
    archive: _EditableHwpxArchive,
    cell_el: ET.Element,
    edit: StyleEdit,
) -> ET.Element:
    border_fill, new_id = _hwpx_clone_or_create_header_ref(
        archive,
        container_name="borderFills",
        child_name="borderFill",
        source_id=cell_el.get("borderFillIDRef"),
        edit=edit,
        default_factory=_hwpx_create_default_border_fill,
    )
    cell_el.set("borderFillIDRef", new_id)
    return border_fill


def _hwpx_border_attrs(border_value: str) -> dict[str, str]:
    attrs = _docx_border_attrs(border_value)
    val = {
        "nil": "NONE",
        "dashed": "DASH",
        "dotted": "DOT",
        "single": "SOLID",
    }.get(attrs["val"], "SOLID")
    pt_match = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*pt", border_value, flags=re.IGNORECASE)
    px_match = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*px", border_value, flags=re.IGNORECASE)
    mm_match = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*mm", border_value, flags=re.IGNORECASE)
    if pt_match:
        width_mm = float(pt_match.group(1)) * 0.352777778
    elif px_match:
        width_mm = float(px_match.group(1)) / _CSS_PX_PER_MM
    elif mm_match:
        width_mm = float(mm_match.group(1))
    else:
        width_mm = 0.12
    width = _hwpx_line_width_value(width_mm)
    return {"type": val, "width": width, "color": f"#{attrs['color']}"}


def _hwpx_remove_fill_brushes(border_fill: ET.Element) -> None:
    for child in list(border_fill):
        if child.tag.rsplit("}", 1)[-1] == "fillBrush":
            border_fill.remove(child)


def _hwpx_set_border_fill_background(border_fill: ET.Element, color: str) -> None:
    _hwpx_remove_fill_brushes(border_fill)
    fill_brush = _hc_el("fillBrush")
    fill_brush.append(
        _hc_el(
            "winBrush",
            {
                "faceColor": f"#{_docx_hex_color(color)}",
                "hatchColor": "#999999",
                "alpha": "0",
            },
        )
    )
    border_fill.append(fill_brush)


def _hwpx_apply_cell_border_fill_style(
    archive: _EditableHwpxArchive,
    cell_el: ET.Element,
    edit: StyleEdit,
) -> None:
    border_fields = {
        "border_top": "topBorder",
        "border_right": "rightBorder",
        "border_bottom": "bottomBorder",
        "border_left": "leftBorder",
    }
    if not any(
        field in _style_edit_supplied_fields(edit)
        for field in {"background", *border_fields}
    ):
        return

    border_fill = _hwpx_clone_border_fill_for_cell(archive, cell_el, edit)
    if "background" in edit.clear_fields:
        _hwpx_remove_fill_brushes(border_fill)
    elif edit.background is not None:
        _hwpx_set_border_fill_background(border_fill, edit.background)

    for field_name, child_name in border_fields.items():
        border_el = _hwpx_find_or_create_hh_child(border_fill, child_name)
        if field_name in edit.clear_fields:
            border_el.set("type", "NONE")
            border_el.set("width", "0.12 mm")
            border_el.set("color", "#000000")
            continue
        value = getattr(edit, field_name)
        if value is None:
            continue
        for attr, attr_value in _hwpx_border_attrs(value).items():
            border_el.set(attr, attr_value)


def _hwpx_apply_cell_horizontal_align(archive: _EditableHwpxArchive, cell_el: ET.Element, edit: StyleEdit) -> None:
    if "horizontal_align" not in _style_edit_supplied_fields(edit):
        return
    sub_list = cell_el.find(f"{_HP}subList")
    if sub_list is None:
        return
    for paragraph_el in sub_list.findall(f"{_HP}p"):
        para_pr = _hwpx_clone_para_pr_for_paragraph(archive, paragraph_el, edit)
        if "horizontal_align" in edit.clear_fields:
            _hwpx_set_para_align(para_pr, None, clear=True)
        elif edit.horizontal_align is not None:
            _hwpx_set_para_align(para_pr, edit.horizontal_align)


def _hwpx_apply_cell_style(archive: _EditableHwpxArchive, location: _HwpxCellLocation, edit: StyleEdit) -> None:
    cell_el = location.element
    _hwpx_set_cell_size_el(location, edit)
    _hwpx_set_cell_margin_el(cell_el, edit)
    if _style_edit_supplied_fields(edit).intersection({"width_pt", "height_pt", "padding_top_pt", "padding_right_pt", "padding_bottom_pt", "padding_left_pt"}):
        for row_el in _hwpx_table_rows(location.table):
            for existing_cell in _hwpx_row_cells(row_el):
                _hwpx_update_cell_text_area(existing_cell, location.table)
    _hwpx_apply_cell_border_fill_style(archive, cell_el, edit)
    _hwpx_apply_cell_horizontal_align(archive, cell_el, edit)
    if "vertical_align" in edit.clear_fields:
        sub_list = cell_el.find(f"{_HP}subList")
        if sub_list is not None:
            sub_list.attrib.pop("vertAlign", None)
    elif edit.vertical_align is not None:
        sub_list = _hwpx_find_or_create_child(cell_el, "subList")
        sub_list.set("vertAlign", {"top": "TOP", "middle": "CENTER", "bottom": "BOTTOM"}[edit.vertical_align])
    _update_hwpx_table_size_from_cells(location.table)


def _hwpx_wrap_value(value: str | None) -> str | None:
    if value is None:
        return None
    return {
        "none": "TOP_AND_BOTTOM",
        "square": "SQUARE",
        "tight": "TIGHT",
        "through": "THROUGH",
        "top_bottom": "TOP_AND_BOTTOM",
        "behind_text": "BEHIND_TEXT",
        "in_front_of_text": "IN_FRONT_OF_TEXT",
    }[value]


def _hwpx_text_flow_value(value: str | None) -> str | None:
    if value is None:
        return None
    return {
        "both_sides": "BOTH_SIDES",
        "left": "LEFT_ONLY",
        "right": "RIGHT_ONLY",
        "largest": "LARGEST_ONLY",
    }[value]


def _hwpx_rel_to_x(value: str | None) -> str | None:
    if value is None:
        return None
    return {
        "page": "PAPER",
        "margin": "PAGE",
        "column": "COLUMN",
        "paragraph": "PARA",
        "character": "PARA",
    }[value]


def _hwpx_rel_to_y(value: str | None) -> str | None:
    if value is None:
        return None
    return {
        "page": "PAPER",
        "margin": "PAGE",
        "paragraph": "PARA",
        "line": "PARA",
    }[value]


def _hwpx_apply_object_placement(object_el: ET.Element, edit: StyleEdit) -> None:
    supplied = _style_edit_supplied_fields(edit)
    if not supplied.intersection(_PLACEMENT_FIELD_MAP):
        return

    if "wrap" in edit.clear_fields:
        object_el.attrib.pop("textWrap", None)
    elif edit.wrap is not None:
        object_el.set("textWrap", _hwpx_wrap_value(edit.wrap) or "")
    if "text_flow" in edit.clear_fields:
        object_el.attrib.pop("textFlow", None)
    elif edit.text_flow is not None:
        object_el.set("textFlow", _hwpx_text_flow_value(edit.text_flow) or "")
    if "z_order" in edit.clear_fields:
        object_el.attrib.pop("zOrder", None)
    elif edit.z_order is not None:
        object_el.set("zOrder", str(edit.z_order))

    pos_el = _hwpx_find_or_create_child(object_el, "pos")
    if "placement_mode" in edit.clear_fields:
        pos_el.attrib.pop("treatAsChar", None)
    elif edit.placement_mode is not None:
        pos_el.set("treatAsChar", "1" if edit.placement_mode == "inline" else "0")
    if "flow_with_text" in edit.clear_fields:
        pos_el.attrib.pop("flowWithText", None)
    elif edit.flow_with_text is not None:
        pos_el.set("flowWithText", "1" if edit.flow_with_text else "0")
    if "allow_overlap" in edit.clear_fields:
        pos_el.attrib.pop("allowOverlap", None)
    elif edit.allow_overlap is not None:
        pos_el.set("allowOverlap", "1" if edit.allow_overlap else "0")
    if edit.x_relative_to is not None:
        pos_el.set("horzRelTo", _hwpx_rel_to_x(edit.x_relative_to) or "PARA")
    if edit.y_relative_to is not None:
        pos_el.set("vertRelTo", _hwpx_rel_to_y(edit.y_relative_to) or "PARA")
    if edit.x_align is not None:
        pos_el.set("horzAlign", edit.x_align.upper())
    if edit.y_align is not None:
        pos_el.set("vertAlign", edit.y_align.upper())
    if "x_offset_pt" in edit.clear_fields:
        pos_el.attrib.pop("horzOffset", None)
    elif edit.x_offset_pt is not None:
        pos_el.set("horzOffset", str(_pt_to_hwpunit(edit.x_offset_pt) or 0))
    if "y_offset_pt" in edit.clear_fields:
        pos_el.attrib.pop("vertOffset", None)
    elif edit.y_offset_pt is not None:
        pos_el.set("vertOffset", str(_pt_to_hwpunit(edit.y_offset_pt) or 0))
    for attr, default in (
        ("affectLSpacing", "0"),
        ("holdAnchorAndSO", "0"),
        ("flowWithText", "1"),
        ("allowOverlap", "0"),
        ("vertRelTo", "PARA"),
        ("horzRelTo", "PARA"),
        ("vertAlign", "TOP"),
        ("horzAlign", "LEFT"),
        ("vertOffset", "0"),
        ("horzOffset", "0"),
    ):
        if attr not in pos_el.attrib:
            pos_el.set(attr, default)

    out_margin = _hwpx_find_or_create_child(object_el, "outMargin")
    for field_name, attr in (
        ("margin_top_pt", "top"),
        ("margin_right_pt", "right"),
        ("margin_bottom_pt", "bottom"),
        ("margin_left_pt", "left"),
    ):
        if field_name in edit.clear_fields:
            out_margin.attrib.pop(attr, None)
        elif (value := getattr(edit, field_name)) is not None:
            out_margin.set(attr, str(_pt_to_hwpunit(value) or 0))


def _apply_hwpx_style_edit(archive: _EditableHwpxArchive, edit: StyleEdit, result: _EditEngineResult) -> None:
    unsupported = _hwpx_style_unsupported_fields(edit)
    if unsupported:
        raise EditValidationError(
            f"HWPX native style write-back does not support {edit.target_kind} field(s): {sorted(unsupported)}.",
            code="invalid_style",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )

    index = _build_hwpx_structural_index(archive)
    if edit.target_kind == "run" and (location := index.runs.get(edit.target_id)):
        _hwpx_apply_run_style(archive, location, edit)
    elif edit.target_kind == "paragraph" and (location := index.paragraphs.get(edit.target_id)):
        _hwpx_apply_paragraph_style(archive, location, edit)
    elif edit.target_kind == "cell" and (location := index.cells.get(edit.target_id)):
        _hwpx_apply_cell_style(archive, location, edit)
    elif edit.target_kind == "table" and (location := index.tables.get(edit.target_id)):
        _hwpx_apply_object_placement(location.element, edit)
    elif edit.target_kind == "image" and (location := index.images.get(edit.target_id)):
        _hwpx_set_size_el(location.element, edit)
        _hwpx_apply_object_placement(location.element, edit)
    else:
        raise EditValidationError(
            f"{edit.target_kind.capitalize()} style target does not exist: {edit.target_id}.",
            code="target_not_found",
            target_kind=edit.target_kind,
            target_id=edit.target_id,
        )

    _append_unique(result.modified_target_ids, edit.target_id)
    result.styles_applied += 1


def _hwpx_el(name: str, attrs: dict[str, str] | None = None) -> ET.Element:
    return ET.Element(f"{_HP}{name}", attrs or {})


def _hh_el(name: str, attrs: dict[str, str] | None = None) -> ET.Element:
    return ET.Element(f"{_HH}{name}", attrs or {})


def _hc_el(name: str, attrs: dict[str, str] | None = None) -> ET.Element:
    return ET.Element(f"{_HC}{name}", attrs or {})


def _hwpx_default_column_widths(col_count: int) -> list[int]:
    if col_count <= 0:
        return []
    width = max(_HWPX_MIN_CELL_WIDTH, _HWPX_DEFAULT_TABLE_WIDTH // col_count)
    return [width for _ in range(col_count)]


def _hwpx_sub_list_el() -> ET.Element:
    return _hwpx_el(
        "subList",
        {
            "id": "",
            "textDirection": "HORIZONTAL",
            "lineWrap": "BREAK",
            "vertAlign": "TOP",
            "linkListIDRef": "0",
            "linkListNextIDRef": "0",
            "textWidth": "0",
            "textHeight": "0",
            "hasTextRef": "0",
            "hasNumRef": "0",
        },
    )


def _hwpx_next_shape_id(archive: _EditableHwpxArchive) -> str:
    max_id = 0
    for section in archive.section_entries:
        for element in section.root.iter():
            value = _safe_int(element.get("id"))
            if value is not None:
                max_id = max(max_id, value)
    return str(max(max_id + 1, 1))


def _hwpx_border_fill_is_default(border_fill: ET.Element) -> bool:
    for side in ("leftBorder", "rightBorder", "topBorder", "bottomBorder"):
        border = border_fill.find(f"{_HH}{side}")
        if border is None:
            return False
        if border.get("type") != "SOLID":
            return False
        if (border.get("color") or "#000000").lower() != "#000000":
            return False
    return True


def _hwpx_create_default_border_fill(border_fill_id: str) -> ET.Element:
    border_fill = _hh_el(
        "borderFill",
        {
            "id": border_fill_id,
            "threeD": "0",
            "shadow": "0",
            "centerLine": "NONE",
            "breakCellSeparateLine": "0",
        },
    )
    border_fill.append(_hh_el("slash", {"type": "NONE", "Crooked": "0", "isCounter": "0"}))
    border_fill.append(_hh_el("backSlash", {"type": "NONE", "Crooked": "0", "isCounter": "0"}))
    for side in ("leftBorder", "rightBorder", "topBorder", "bottomBorder"):
        border_fill.append(_hh_el(side, {"type": "SOLID", "width": "0.12 mm", "color": "#000000"}))
    border_fill.append(_hh_el("diagonal", {"type": "SOLID", "width": "0.1 mm", "color": "#000000"}))
    return border_fill


def _ensure_hwpx_ref_list(header: _EditableHwpxHeader) -> ET.Element:
    ref_list = header.root.find(f"{_HH}refList")
    if ref_list is not None:
        return ref_list
    ref_list = _hh_el("refList")
    children = list(header.root)
    begin_num = header.root.find(f"{_HH}beginNum")
    insert_at = children.index(begin_num) + 1 if begin_num in children else 0
    header.root.insert(insert_at, ref_list)
    return ref_list


def _ensure_hwpx_border_fills(header: _EditableHwpxHeader) -> ET.Element:
    ref_list = _ensure_hwpx_ref_list(header)
    border_fills = ref_list.find(f"{_HH}borderFills")
    if border_fills is not None:
        return border_fills
    border_fills = _hh_el("borderFills", {"itemCnt": "0"})
    children = list(ref_list)
    insert_at = 0
    for index, child in enumerate(children):
        if child.tag == f"{_HH}fontfaces":
            insert_at = index + 1
            break
    ref_list.insert(insert_at, border_fills)
    return border_fills


def _ensure_hwpx_default_border_fill_id(archive: _EditableHwpxArchive) -> str:
    if archive.header_entry is None:
        return "0"
    border_fills = _ensure_hwpx_border_fills(archive.header_entry)
    existing_ids: list[int] = []
    for border_fill in border_fills.findall(f"{_HH}borderFill"):
        border_fill_id = _safe_int(border_fill.get("id"))
        if border_fill_id is not None:
            existing_ids.append(border_fill_id)
        if border_fill_id is not None and _hwpx_border_fill_is_default(border_fill):
            return str(border_fill_id)

    new_id = str((max(existing_ids) + 1) if existing_ids else 1)
    border_fills.append(_hwpx_create_default_border_fill(new_id))
    border_fills.set("itemCnt", str(len(border_fills.findall(f"{_HH}borderFill"))))
    return new_id


def _hwpx_paragraph_el(text: str) -> ET.Element:
    paragraph = _hwpx_el("p")
    run = _hwpx_el("run")
    text_el = _hwpx_el("t")
    text_el.text = text
    run.append(text_el)
    paragraph.append(run)
    return paragraph


def _hwpx_cell_el(
    text: str,
    *,
    row_index: int,
    col_index: int,
    border_fill_id: str,
    width: int,
    height: int = _HWPX_DEFAULT_CELL_HEIGHT,
) -> ET.Element:
    cell = _hwpx_el(
        "tc",
        {
            "name": "",
            "header": "0",
            "hasMargin": "0",
            "protect": "0",
            "editable": "0",
            "dirty": "0",
            "borderFillIDRef": border_fill_id,
        },
    )
    sub_list = _hwpx_sub_list_el()
    sub_list.append(_hwpx_paragraph_el(text))
    cell.append(sub_list)
    cell.append(_hwpx_el("cellAddr", {"colAddr": str(col_index - 1), "rowAddr": str(row_index - 1)}))
    cell.append(_hwpx_el("cellSpan", {"colSpan": "1", "rowSpan": "1"}))
    cell.append(_hwpx_el("cellSz", {"width": str(width), "height": str(height)}))
    cell.append(
        _hwpx_el(
            "cellMargin",
            {
                "left": str(_HWPX_DEFAULT_CELL_MARGIN_X),
                "right": str(_HWPX_DEFAULT_CELL_MARGIN_X),
                "top": str(_HWPX_DEFAULT_CELL_MARGIN_Y),
                "bottom": str(_HWPX_DEFAULT_CELL_MARGIN_Y),
            },
        )
    )
    return cell


def _hwpx_clone_cell_for_text(cell_el: ET.Element, text: str) -> ET.Element:
    cloned = deepcopy(cell_el)
    _hwpx_set_cell_text(cloned, text)
    return cloned


def _hwpx_row_el(
    values: list[str],
    *,
    row_index: int,
    border_fill_id: str,
    col_widths: list[int] | None = None,
    template_row: ET.Element | None = None,
) -> ET.Element:
    row = _hwpx_el("tr")
    template_cells = _hwpx_row_cells(template_row) if template_row is not None else []
    col_widths = col_widths or _hwpx_default_column_widths(len(values))
    for col_index, value in enumerate(values, start=1):
        if col_index <= len(template_cells):
            row.append(_hwpx_clone_cell_for_text(template_cells[col_index - 1], value))
        else:
            row.append(
                _hwpx_cell_el(
                    value,
                    row_index=row_index,
                    col_index=col_index,
                    border_fill_id=border_fill_id,
                    width=col_widths[col_index - 1] if col_index - 1 < len(col_widths) else _HWPX_MIN_CELL_WIDTH,
                )
            )
    return row


def _hwpx_table_el(rows: list[list[str]], *, border_fill_id: str, object_id: str) -> ET.Element:
    col_widths = _hwpx_default_column_widths(len(rows[0]))
    table_width = sum(col_widths)
    table_height = len(rows) * _HWPX_DEFAULT_CELL_HEIGHT
    table = _hwpx_el(
        "tbl",
        {
            "id": object_id,
            "zOrder": "0",
            "numberingType": "TABLE",
            "textWrap": "TOP_AND_BOTTOM",
            "textFlow": "BOTH_SIDES",
            "lock": "0",
            "dropcapstyle": "None",
            "pageBreak": "CELL",
            "repeatHeader": "0",
            "rowCnt": str(len(rows)),
            "colCnt": str(len(rows[0])),
            "cellSpacing": "0",
            "borderFillIDRef": border_fill_id,
            "noAdjust": "0",
        },
    )
    table.append(_hwpx_el("sz", {"width": str(table_width), "widthRelTo": "ABSOLUTE", "height": str(table_height), "heightRelTo": "ABSOLUTE", "protect": "0"}))
    table.append(
        _hwpx_el(
            "pos",
            {
                "treatAsChar": "1",
                "affectLSpacing": "0",
                "flowWithText": "1",
                "allowOverlap": "0",
                "holdAnchorAndSO": "0",
                "vertRelTo": "PARA",
                "horzRelTo": "PARA",
                "vertAlign": "TOP",
                "horzAlign": "LEFT",
                "vertOffset": "0",
                "horzOffset": "0",
            },
        )
    )
    table.append(_hwpx_el("outMargin", {"left": "0", "right": "0", "top": "0", "bottom": "0"}))
    table.append(
        _hwpx_el(
            "inMargin",
            {
                "left": str(_HWPX_DEFAULT_CELL_MARGIN_X),
                "right": str(_HWPX_DEFAULT_CELL_MARGIN_X),
                "top": str(_HWPX_DEFAULT_CELL_MARGIN_Y),
                "bottom": str(_HWPX_DEFAULT_CELL_MARGIN_Y),
            },
        )
    )
    for row_index, row in enumerate(rows, start=1):
        table.append(_hwpx_row_el(row, row_index=row_index, border_fill_id=border_fill_id, col_widths=col_widths))
    return table


def _hwpx_table_rows(table_el: ET.Element) -> list[ET.Element]:
    return table_el.findall(f"{_HP}tr")


def _hwpx_row_cells(row_el: ET.Element) -> list[ET.Element]:
    return row_el.findall(f"{_HP}tc")


def _renumber_hwpx_table(table_el: ET.Element) -> None:
    rows = _hwpx_table_rows(table_el)
    max_cols = 0
    for row_index, row in enumerate(rows, start=1):
        cells = _hwpx_row_cells(row)
        max_cols = max(max_cols, len(cells))
        for col_index, cell in enumerate(cells, start=1):
            cell_addr = cell.find(f"{_HP}cellAddr")
            if cell_addr is None:
                cell_addr = _hwpx_el("cellAddr")
                cell.append(cell_addr)
            cell_addr.set("rowAddr", str(row_index - 1))
            cell_addr.set("colAddr", str(col_index - 1))
            cell_span = cell.find(f"{_HP}cellSpan")
            if cell_span is None:
                cell_span = _hwpx_el("cellSpan")
                cell.append(cell_span)
            cell_span.set("rowSpan", cell_span.get("rowSpan") or "1")
            cell_span.set("colSpan", cell_span.get("colSpan") or "1")
    table_el.set("rowCnt", str(len(rows)))
    table_el.set("colCnt", str(max_cols))
    _update_hwpx_table_size_from_cells(table_el)


def _update_hwpx_table_size_from_cells(table_el: ET.Element) -> None:
    max_row_width = 0
    total_height = 0
    for row in _hwpx_table_rows(table_el):
        row_width = 0
        row_height = 0
        for cell in _hwpx_row_cells(row):
            cell_size = cell.find(f"{_HP}cellSz")
            if cell_size is None:
                continue
            row_width += _safe_int(cell_size.get("width")) or 0
            row_height = max(row_height, _safe_int(cell_size.get("height")) or 0)
        max_row_width = max(max_row_width, row_width)
        total_height += row_height

    if max_row_width <= 0 and total_height <= 0:
        return
    size_el = _hwpx_find_or_create_child(table_el, "sz")
    if max_row_width > 0:
        size_el.set("width", str(max_row_width))
        size_el.set("widthRelTo", "ABSOLUTE")
    if total_height > 0:
        size_el.set("height", str(total_height))
        size_el.set("heightRelTo", "ABSOLUTE")
    if "protect" not in size_el.attrib:
        size_el.set("protect", "0")


def _hwpx_set_cell_text(cell_el: ET.Element, text: str) -> None:
    sub_list = cell_el.find(f"{_HP}subList")
    if sub_list is None:
        sub_list = _hwpx_sub_list_el()
        cell_el.insert(0, sub_list)
    for child in list(sub_list):
        if child.tag == f"{_HP}p":
            sub_list.remove(child)
    for line in (text.split("\n") if text != "" else [""]):
        sub_list.append(_hwpx_paragraph_el(line))


def _ensure_hwpx_parent_has_paragraph(parent_el: ET.Element) -> None:
    if parent_el.findall(f"{_HP}p"):
        return
    parent_el.append(_hwpx_paragraph_el(""))


def _resolve_hwpx_table_axis(
    index: _HwpxStructuralIndex,
    operation: StructuralEdit,
    *,
    axis: str,
) -> tuple[ET.Element, int]:
    if cell_location := index.cells.get(operation.target_id):
        return cell_location.table, cell_location.row_index if axis == "row" else cell_location.col_index
    if table_location := index.tables.get(operation.target_id):
        axis_index = operation.row_index if axis == "row" else operation.column_index
        if axis_index is None:
            raise EditValidationError(
                f"{operation.operation} with a table target requires {axis}_index.",
                code="index_out_of_bounds",
                target_kind="table",
                target_id=operation.target_id,
                operation=operation.operation,
            )
        return table_location.element, axis_index
    raise EditValidationError(
        f"{operation.operation} target must be a table or cell.",
        code="target_not_found",
        target_id=operation.target_id,
        operation=operation.operation,
    )


def _insert_child_relative(parent: ET.Element, target: ET.Element, new_child: ET.Element, *, position: str) -> None:
    children = list(parent)
    target_index = children.index(target)
    parent.insert(target_index if position in {"before", "start"} else target_index + 1, new_child)


def _apply_hwpx_structural_operation(archive: _EditableHwpxArchive, operation: StructuralEdit, result: _EditEngineResult) -> None:
    index = _build_hwpx_structural_index(archive)

    if operation.operation == "insert_paragraph":
        paragraph_el = _hwpx_paragraph_el(operation.text or "")
        if paragraph_location := index.paragraphs.get(operation.target_id):
            if operation.position not in {"before", "after"}:
                raise EditValidationError("insert_paragraph with a paragraph target requires before/after.", code="invalid_position")
            _insert_child_relative(paragraph_location.parent, paragraph_location.element, paragraph_el, position=operation.position)
        elif cell_location := index.cells.get(operation.target_id):
            sub_list = cell_location.element.find(f"{_HP}subList")
            if sub_list is None:
                sub_list = _hwpx_el("subList")
                cell_location.element.insert(0, sub_list)
            if operation.position in {"start", "before"}:
                sub_list.insert(0, paragraph_el)
            else:
                sub_list.append(paragraph_el)
        else:
            raise EditValidationError("insert_paragraph target must be a paragraph or cell.", code="target_not_found", target_id=operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_paragraph":
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is None:
            raise EditValidationError("Paragraph does not exist.", code="target_not_found", target_kind="paragraph", target_id=operation.target_id)
        current_text = _hwpx_paragraph_visible_text(paragraph_location.element)
        current_text_hash = _text_hash(current_text)
        if operation.expected_text_hash is not None and current_text_hash != operation.expected_text_hash:
            raise EditValidationError(
                "Paragraph text hash mismatch.",
                code="text_hash_mismatch",
                target_kind="paragraph",
                target_id=operation.target_id,
                operation=operation.operation,
                current_text=current_text,
                expected_text_hash=operation.expected_text_hash,
                current_text_hash=current_text_hash,
            )
        paragraph_location.parent.remove(paragraph_location.element)
        _ensure_hwpx_parent_has_paragraph(paragraph_location.parent)
        result.operations_applied += 1
        return

    if operation.operation == "insert_run":
        run_el = _hwpx_el("run")
        text_el = _hwpx_el("t")
        text_el.text = operation.text or ""
        run_el.append(text_el)
        if run_location := index.runs.get(operation.target_id):
            if operation.position not in {"before", "after"}:
                raise EditValidationError("insert_run with a run target requires before/after.", code="invalid_position")
            _insert_child_relative(run_location.parent, run_location.element, run_el, position=operation.position)
        elif paragraph_location := index.paragraphs.get(operation.target_id):
            if operation.position in {"start", "before"}:
                paragraph_location.element.insert(0, run_el)
            else:
                paragraph_location.element.append(run_el)
        else:
            raise EditValidationError("insert_run target must be a run or paragraph.", code="target_not_found", target_id=operation.target_id)
        result.operations_applied += 1
        return

    if operation.operation == "remove_run":
        run_location = index.runs.get(operation.target_id)
        if run_location is None:
            raise EditValidationError("Run does not exist.", code="target_not_found", target_kind="run", target_id=operation.target_id)
        current_text = _run_text(run_location.element)
        current_text_hash = _text_hash(current_text)
        if operation.expected_text_hash is not None and current_text_hash != operation.expected_text_hash:
            raise EditValidationError(
                "Run text hash mismatch.",
                code="text_hash_mismatch",
                target_kind="run",
                target_id=operation.target_id,
                operation=operation.operation,
                current_text=current_text,
                expected_text_hash=operation.expected_text_hash,
                current_text_hash=current_text_hash,
            )
        run_location.parent.remove(run_location.element)
        if not run_location.parent.findall(f"{_HP}run"):
            run_location.parent.append(_hwpx_el("run"))
        result.operations_applied += 1
        return

    if operation.operation == "insert_table":
        paragraph_location = index.paragraphs.get(operation.target_id)
        if paragraph_location is None:
            raise EditValidationError("insert_table target must be a paragraph.", code="target_not_found", target_kind="paragraph", target_id=operation.target_id)
        table_paragraph = _hwpx_el("p")
        run = _hwpx_el("run")
        run.append(
            _hwpx_table_el(
                _normalize_table_rows(operation.rows),
                border_fill_id=_ensure_hwpx_default_border_fill_id(archive),
                object_id=_hwpx_next_shape_id(archive),
            )
        )
        table_paragraph.append(run)
        if operation.position == "before":
            _insert_child_relative(paragraph_location.parent, paragraph_location.element, table_paragraph, position="before")
        elif operation.position == "after":
            _insert_child_relative(paragraph_location.parent, paragraph_location.element, table_paragraph, position="after")
        else:
            raise EditValidationError("insert_table requires before/after.", code="invalid_position")
        result.operations_applied += 1
        return

    if operation.operation == "remove_table":
        table_location = index.tables.get(operation.target_id)
        if table_location is None:
            raise EditValidationError("Table does not exist.", code="target_not_found", target_kind="table", target_id=operation.target_id)
        table_location.parent.remove(table_location.element)
        result.operations_applied += 1
        return

    if operation.operation == "set_cell_text":
        cell_location = index.cells.get(operation.target_id)
        if cell_location is None:
            raise EditValidationError("Cell does not exist.", code="target_not_found", target_kind="cell", target_id=operation.target_id)
        current_text = _hwpx_cell_visible_text(cell_location.element)
        current_text_hash = _text_hash(current_text)
        if operation.expected_text_hash is not None and current_text_hash != operation.expected_text_hash:
            raise EditValidationError(
                "Cell text hash mismatch.",
                code="text_hash_mismatch",
                target_kind="cell",
                target_id=operation.target_id,
                operation=operation.operation,
                current_text=current_text,
                expected_text_hash=operation.expected_text_hash,
                current_text_hash=current_text_hash,
            )
        _hwpx_set_cell_text(cell_location.element, operation.text or "")
        result.operations_applied += 1
        return

    if operation.operation in {"insert_table_row", "remove_table_row"}:
        table_el, row_index = _resolve_hwpx_table_axis(index, operation, axis="row")
        rows = _hwpx_table_rows(table_el)
        if row_index < 1 or row_index > len(rows):
            raise EditValidationError("Table row index is out of bounds.", code="index_out_of_bounds")
        if operation.operation == "remove_table_row":
            if len(rows) <= 1:
                raise EditValidationError("Cannot remove the only table row.", code="invalid_table_shape")
            table_el.remove(rows[row_index - 1])
        else:
            col_count = max((len(_hwpx_row_cells(row)) for row in rows), default=0)
            values = operation.values or ["" for _ in range(col_count)]
            if len(values) != col_count:
                raise EditValidationError("Inserted row values must match table column count.", code="invalid_table_shape")
            template_row = rows[row_index - 1]
            col_widths: list[int] = []
            for cell in _hwpx_row_cells(template_row):
                cell_size = cell.find(f"{_HP}cellSz")
                col_widths.append((_safe_int(cell_size.get("width")) if cell_size is not None else None) or _HWPX_MIN_CELL_WIDTH)
            if len(col_widths) < col_count:
                col_widths.extend(_hwpx_default_column_widths(col_count)[len(col_widths) :])
            new_row = _hwpx_row_el(
                values,
                row_index=row_index,
                border_fill_id=_ensure_hwpx_default_border_fill_id(archive),
                col_widths=col_widths,
                template_row=template_row,
            )
            _insert_child_relative(table_el, rows[row_index - 1], new_row, position=operation.position)
        _renumber_hwpx_table(table_el)
        result.operations_applied += 1
        return

    if operation.operation in {"insert_table_column", "remove_table_column"}:
        table_el, column_index = _resolve_hwpx_table_axis(index, operation, axis="column")
        rows = _hwpx_table_rows(table_el)
        col_count = max((len(_hwpx_row_cells(row)) for row in rows), default=0)
        if column_index < 1 or column_index > col_count:
            raise EditValidationError("Table column index is out of bounds.", code="index_out_of_bounds")
        if operation.operation == "remove_table_column":
            if col_count <= 1:
                raise EditValidationError("Cannot remove the only table column.", code="invalid_table_shape")
            for row in rows:
                cells = _hwpx_row_cells(row)
                if len(cells) >= column_index:
                    row.remove(cells[column_index - 1])
        else:
            values = operation.values or ["" for _ in rows]
            if len(values) != len(rows):
                raise EditValidationError("Inserted column values must match table row count.", code="invalid_table_shape")
            for row_index, row in enumerate(rows, start=1):
                cells = _hwpx_row_cells(row)
                if len(cells) >= column_index:
                    new_cell = _hwpx_clone_cell_for_text(cells[column_index - 1], values[row_index - 1])
                else:
                    new_cell = _hwpx_cell_el(
                        values[row_index - 1],
                        row_index=row_index,
                        col_index=column_index,
                        border_fill_id=_ensure_hwpx_default_border_fill_id(archive),
                        width=_HWPX_MIN_CELL_WIDTH,
                    )
                insert_index = column_index if operation.position in {"after", "end"} else column_index - 1
                row.insert(insert_index, new_cell)
        _renumber_hwpx_table(table_el)
        result.operations_applied += 1
        return

    raise EditValidationError(f"Unsupported structural operation: {operation.operation!r}.", code="invalid_operation")


def _apply_document_edits_to_file(
    source_path: str | Path,
    operations: list[StructuralEdit],
    *,
    output_path: str | Path | None = None,
) -> _EditEngineResult:
    source = Path(source_path)
    doc = DocIR.from_file(source)
    result = _EditEngineResult(source_doc_type=doc.source_doc_type)
    target_suffix = ".hwpx" if doc.source_doc_type == "hwp" else None
    target_path = (
        Path(output_path)
        if output_path is not None
        else _default_output_path(source, output_suffix=target_suffix)
    )
    target_path = _normalize_output_path_for_source_doc_type(
        target_path,
        source_doc_type=doc.source_doc_type,
        result=result,
    )

    if _same_path(source, target_path):
        raise EditValidationError(
            f"Refusing to overwrite source file {source}; choose a different output path.",
            code="output_path_conflicts_with_source",
        )

    if doc.source_doc_type == "docx":
        from docx import Document as load_docx

        native_doc = load_docx(str(source))
        for operation in operations:
            _apply_docx_structural_operation(native_doc, operation, result)
        native_doc.save(str(target_path))
    elif doc.source_doc_type == "hwpx":
        archive = _EditableHwpxArchive.open(source)
        for operation in operations:
            _apply_hwpx_structural_operation(archive, operation, result)
        archive.write_to(target_path)
    elif doc.source_doc_type == "hwp":
        archive = _EditableHwpxArchive.from_bytes(
            convert_hwp_to_hwpx_bytes(source),
            source_path=source.with_suffix(".hwpx"),
        )
        for operation in operations:
            _apply_hwpx_structural_operation(archive, operation, result)
        archive.write_to(target_path)
    else:
        raise EditValidationError(
            f"Native write-back is currently supported only for docx/hwp/hwpx, got {doc.source_doc_type!r}.",
            code="unsupported_source_doc_type",
        )

    result.output_path = str(target_path)
    result.output_filename = target_path.name
    return result


def _apply_document_edits_to_bytes(
    source_bytes: bytes,
    operations: list[StructuralEdit],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_filename: str | None = None,
) -> _EditEngineResult:
    resolved_doc_type = _resolve_bytes_doc_type(
        source_bytes,
        doc_type=doc_type,
        source_name=source_name,
    )
    with TemporarySourcePath(source_bytes, suffix=_source_suffix_for_doc_type(resolved_doc_type)) as source_path:
        default_filename = _default_output_filename(
            source_name=source_name,
            source_doc_type=resolved_doc_type,
        )
        chosen_filename = output_filename or default_filename
        with tempfile.TemporaryDirectory() as tmp_dir:
            target_path = Path(tmp_dir) / chosen_filename
            result = _apply_document_edits_to_file(source_path, operations, output_path=target_path)
            output_path = Path(result.output_path) if result.output_path is not None else target_path
            result.output_bytes = output_path.read_bytes()
            result.output_filename = output_path.name
            result.output_path = None
            return result


def _apply_document_edits_to_source(
    source: DocIR | str | Path | bytes | BinaryIO,
    operations: list[StructuralEdit],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_path: str | Path | None = None,
    output_filename: str | None = None,
) -> _EditEngineResult:
    if isinstance(source, DocIR):
        return _apply_structural_edits_to_doc_ir(source, operations)

    if output_path is not None and output_filename is not None:
        raise ValueError("Specify either output_path or output_filename, not both.")

    if isinstance(source, (str, Path)):
        resolved_output_path = output_path
        if resolved_output_path is None and output_filename is not None:
            resolved_output_path = Path(source).with_name(output_filename)
        result = _apply_document_edits_to_file(source, operations, output_path=resolved_output_path)
        if result.output_path is not None:
            result.output_filename = Path(result.output_path).name
        return result

    source_bytes = coerce_source_to_supported_value(source, doc_type=infer_doc_type(source, doc_type))
    if not isinstance(source_bytes, bytes):
        raise TypeError("Expected bytes-like source after coercion.")
    return _apply_document_edits_to_bytes(
        source_bytes,
        operations,
        doc_type=doc_type,
        source_name=source_name,
        output_filename=output_filename,
    )


def _apply_style_edits_to_file(
    source_path: str | Path,
    edits: list[StyleEdit],
    *,
    output_path: str | Path | None = None,
) -> _EditEngineResult:
    source = Path(source_path)
    doc = DocIR.from_file(source)
    result = _EditEngineResult(source_doc_type=doc.source_doc_type)
    target_suffix = ".hwpx" if doc.source_doc_type == "hwp" else None
    target_path = (
        Path(output_path)
        if output_path is not None
        else _default_output_path(source, output_suffix=target_suffix)
    )
    target_path = _normalize_output_path_for_source_doc_type(
        target_path,
        source_doc_type=doc.source_doc_type,
        result=result,
    )

    if _same_path(source, target_path):
        raise EditValidationError(
            f"Refusing to overwrite source file {source}; choose a different output path.",
            code="output_path_conflicts_with_source",
        )

    if doc.source_doc_type == "docx":
        from docx import Document as load_docx

        native_doc = load_docx(str(source))
        for edit in edits:
            _apply_docx_style_edit(native_doc, edit, result)
        native_doc.save(str(target_path))
    elif doc.source_doc_type == "hwpx":
        archive = _EditableHwpxArchive.open(source)
        for edit in edits:
            _apply_hwpx_style_edit(archive, edit, result)
        archive.write_to(target_path)
    elif doc.source_doc_type == "hwp":
        archive = _EditableHwpxArchive.from_bytes(
            convert_hwp_to_hwpx_bytes(source),
            source_path=source.with_suffix(".hwpx"),
        )
        for edit in edits:
            _apply_hwpx_style_edit(archive, edit, result)
        archive.write_to(target_path)
    else:
        raise EditValidationError(
            f"Native style write-back is currently supported only for docx/hwp/hwpx, got {doc.source_doc_type!r}.",
            code="unsupported_source_doc_type",
        )

    result.output_path = str(target_path)
    result.output_filename = target_path.name
    return result


def _apply_style_edits_to_bytes(
    source_bytes: bytes,
    edits: list[StyleEdit],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_filename: str | None = None,
) -> _EditEngineResult:
    resolved_doc_type = _resolve_bytes_doc_type(
        source_bytes,
        doc_type=doc_type,
        source_name=source_name,
    )
    with TemporarySourcePath(source_bytes, suffix=_source_suffix_for_doc_type(resolved_doc_type)) as source_path:
        default_filename = _default_output_filename(
            source_name=source_name,
            source_doc_type=resolved_doc_type,
        )
        chosen_filename = output_filename or default_filename
        with tempfile.TemporaryDirectory() as tmp_dir:
            target_path = Path(tmp_dir) / chosen_filename
            result = _apply_style_edits_to_file(source_path, edits, output_path=target_path)
            output_path = Path(result.output_path) if result.output_path is not None else target_path
            result.output_bytes = output_path.read_bytes()
            result.output_filename = output_path.name
            result.output_path = None
            return result


def _apply_style_edits_to_source(
    source: DocIR | str | Path | bytes | BinaryIO,
    edits: list[StyleEdit],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_path: str | Path | None = None,
    output_filename: str | None = None,
) -> _EditEngineResult:
    if isinstance(source, DocIR):
        return _apply_style_edits_to_doc_ir(source, edits)

    if output_path is not None and output_filename is not None:
        raise ValueError("Specify either output_path or output_filename, not both.")

    if isinstance(source, (str, Path)):
        resolved_output_path = output_path
        if resolved_output_path is None and output_filename is not None:
            resolved_output_path = Path(source).with_name(output_filename)
        result = _apply_style_edits_to_file(source, edits, output_path=resolved_output_path)
        if result.output_path is not None:
            result.output_filename = Path(result.output_path).name
        return result

    source_bytes = coerce_source_to_supported_value(source, doc_type=infer_doc_type(source, doc_type))
    if not isinstance(source_bytes, bytes):
        raise TypeError("Expected bytes-like source after coercion.")
    return _apply_style_edits_to_bytes(
        source_bytes,
        edits,
        doc_type=doc_type,
        source_name=source_name,
        output_filename=output_filename,
    )


def _apply_text_edits_to_bytes(
    source_bytes: bytes,
    edits: list[TextEdit],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_filename: str | None = None,
) -> _EditEngineResult:
    resolved_doc_type = _resolve_bytes_doc_type(
        source_bytes,
        doc_type=doc_type,
        source_name=source_name,
    )

    with TemporarySourcePath(source_bytes, suffix=_source_suffix_for_doc_type(resolved_doc_type)) as source_path:
        default_filename = _default_output_filename(
            source_name=source_name,
            source_doc_type=resolved_doc_type,
        )
        chosen_filename = output_filename or default_filename
        with tempfile.TemporaryDirectory() as tmp_dir:
            target_path = Path(tmp_dir) / chosen_filename
            result = _apply_text_edits_to_file(source_path, edits, output_path=target_path)
            output_path = Path(result.output_path) if result.output_path is not None else target_path
            result.output_bytes = output_path.read_bytes()
            result.output_filename = output_path.name
            result.output_path = None
            return result


def _apply_text_edits_to_source(
    source: DocIR | str | Path | bytes | BinaryIO,
    edits: list[TextEdit],
    *,
    doc_type: SourceDocType = "auto",
    source_name: str | None = None,
    output_path: str | Path | None = None,
    output_filename: str | None = None,
) -> _EditEngineResult:
    if isinstance(source, DocIR):
        return _apply_text_edits_to_doc_ir(source, edits)

    if output_path is not None and output_filename is not None:
        raise ValueError("Specify either output_path or output_filename, not both.")

    if isinstance(source, (str, Path)):
        resolved_output_path = output_path
        if resolved_output_path is None and output_filename is not None:
            resolved_output_path = Path(source).with_name(output_filename)
        result = _apply_text_edits_to_file(source, edits, output_path=resolved_output_path)
        if result.output_path is not None:
            result.output_filename = Path(result.output_path).name
        return result

    source_bytes = coerce_source_to_supported_value(source, doc_type=infer_doc_type(source, doc_type))
    if not isinstance(source_bytes, bytes):
        raise TypeError("Expected bytes-like source after coercion.")
    return _apply_text_edits_to_bytes(
        source_bytes,
        edits,
        doc_type=doc_type,
        source_name=source_name,
        output_filename=output_filename,
    )


def _apply_text_edits_to_file(
    source_path: str | Path,
    edits: list[TextEdit],
    *,
    output_path: str | Path | None = None,
) -> _EditEngineResult:
    source = Path(source_path)
    doc = DocIR.from_file(source)
    result = _EditEngineResult(source_doc_type=doc.source_doc_type)
    target_suffix = ".hwpx" if doc.source_doc_type == "hwp" else None
    target_path = (
        Path(output_path)
        if output_path is not None
        else _default_output_path(source, output_suffix=target_suffix)
    )
    target_path = _normalize_output_path_for_source_doc_type(
        target_path,
        source_doc_type=doc.source_doc_type,
        result=result,
    )

    if _same_path(source, target_path):
        raise EditValidationError(
            f"Refusing to overwrite source file {source}; choose a different output path."
        )

    if doc.source_doc_type == "docx":
        from docx import Document as load_docx

        native_doc = load_docx(str(source))
        index = _build_docx_index(native_doc)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        native_doc.save(str(target_path))
    elif doc.source_doc_type == "hwpx":
        archive = _EditableHwpxArchive.open(source)
        index = _build_hwpx_index(archive)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        archive.write_to(target_path)
    elif doc.source_doc_type == "hwp":
        archive = _EditableHwpxArchive.from_bytes(
            convert_hwp_to_hwpx_bytes(source),
            source_path=source.with_suffix(".hwpx"),
        )
        index = _build_hwpx_index(archive)
        for edit in edits:
            _apply_single_edit(index, edit, result)
        archive.write_to(target_path)
    else:
        raise EditValidationError(
            f"Native write-back is currently supported only for docx/hwp/hwpx, got {doc.source_doc_type!r}."
        )

    result.output_path = str(target_path)
    result.output_filename = target_path.name
    return result


__all__: list[str] = []
