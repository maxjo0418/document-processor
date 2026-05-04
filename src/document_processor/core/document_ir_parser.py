"""Direct DOCX/HWPX/PDF file parsing into structural document IR."""

from __future__ import annotations

from collections import OrderedDict
import hashlib
import mimetypes
import struct
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable, Literal
from xml.etree import ElementTree as ET
import zipfile

from ..io_utils import coerce_source_to_supported_value
from ..models import (
    DocIR,
    ImageAsset,
    ImageIR,
    PageInfo,
    ParagraphIR,
    RunIR,
    TableCellIR,
    TableIR,
    _anchored_node_id,
    _make_native_anchor,
)
from ..style_types import ColumnLayoutInfo, ParaStyleInfo
from .docx_structured_exporter import _iter_blocks, _iter_blocks_from_element, _load_docx_source
from .hwp_converter import convert_hwp_to_hwpx_bytes
from .hwpx_structured_exporter import _HP, _logical_table_cells, _paragraph_text, _run_text, _safe_int, _section_roots_from_bytes

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from ..hwpx import HwpxDocument


DocType = Literal["hwp", "hwpx", "docx", "pdf"]

_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_DOCX_EMBED_ATTR = f"{{{_REL_NS}}}embed"
_HC_IMG_TAG = "img"
_EMU_PER_PT = 12700.0
_HWPUNIT_PER_PT = 100.0
_TWIPS_PER_PT = 20.0


def _node_kwargs(
    kind,
    structural_path: str,
    *,
    source_doc_type: str,
    parent_debug_path: str | None = None,
    part_name: str | None = None,
    text: str | None = None,
) -> dict[str, object]:
    return {
        "node_id": _anchored_node_id(kind, structural_path),
        "native_anchor": _make_native_anchor(
            kind,
            structural_path,
            source_doc_type=source_doc_type,
            parent_debug_path=parent_debug_path,
            part_name=part_name,
            text=text,
        ),
    }


def _emu_to_pt(value: str | int | None) -> float | None:
    if value is None:
        return None
    try:
        return int(value) / _EMU_PER_PT
    except (TypeError, ValueError):
        return None


def _hwpunit_to_pt(value: str | int | None) -> float | None:
    if value is None:
        return None
    try:
        return int(value) / _HWPUNIT_PER_PT
    except (TypeError, ValueError):
        return None


def _twips_to_pt(value: str | int | None) -> float | None:
    if value is None:
        return None
    try:
        return int(value) / _TWIPS_PER_PT
    except (TypeError, ValueError):
        return None


def _page_layout(
    *,
    width_pt: float | None,
    height_pt: float | None,
    margin_left_pt: float | None = None,
    margin_right_pt: float | None = None,
    margin_top_pt: float | None = None,
    margin_bottom_pt: float | None = None,
) -> dict[str, float | None]:
    return {
        "width_pt": width_pt,
        "height_pt": height_pt,
        "margin_left_pt": margin_left_pt,
        "margin_right_pt": margin_right_pt,
        "margin_top_pt": margin_top_pt,
        "margin_bottom_pt": margin_bottom_pt,
    }


def _copy_column_style(column_style: ColumnLayoutInfo | None) -> ColumnLayoutInfo | None:
    return column_style.model_copy(deep=True) if column_style is not None else None


def _has_meaningful_column_layout(column_style: ColumnLayoutInfo | None) -> bool:
    if column_style is None:
        return False
    return (column_style.count or 1) > 1 or bool(column_style.widths_pt) or bool(column_style.gaps_pt)


def _para_style_with_columns(column_style: ColumnLayoutInfo | None) -> ParaStyleInfo | None:
    copied_column_style = _copy_column_style(column_style)
    if not _has_meaningful_column_layout(copied_column_style):
        return None
    return ParaStyleInfo(column_layout=copied_column_style)


def _ensure_page_info(
    pages: "OrderedDict[int, PageInfo]",
    *,
    page_number: int,
    layout: dict[str, Any],
) -> None:
    if page_number not in pages:
        page_layout = {
            key: layout.get(key)
            for key in (
                "width_pt",
                "height_pt",
                "margin_left_pt",
                "margin_right_pt",
                "margin_top_pt",
                "margin_bottom_pt",
            )
        }
        pages[page_number] = PageInfo(page_number=page_number, **page_layout)


def _assign_page_number_to_paragraph(paragraph: ParagraphIR, page_number: int | None) -> None:
    paragraph.page_number = page_number
    for node in paragraph.content:
        if isinstance(node, TableIR):
            for cell in node.cells:
                for cell_paragraph in cell.paragraphs:
                    _assign_page_number_to_paragraph(cell_paragraph, page_number)


def _image_dimensions_from_bytes(data: bytes) -> tuple[int | None, int | None]:
    if data.startswith(b"\x89PNG\r\n\x1a\n") and len(data) >= 24:
        return struct.unpack(">II", data[16:24])

    if data.startswith((b"GIF87a", b"GIF89a")) and len(data) >= 10:
        return struct.unpack("<HH", data[6:10])

    if len(data) >= 4 and data[:2] == b"\xff\xd8":
        offset = 2
        while offset + 9 < len(data):
            if data[offset] != 0xFF:
                offset += 1
                continue
            marker = data[offset + 1]
            offset += 2
            if marker in (0xD8, 0xD9):
                continue
            if offset + 2 > len(data):
                break
            segment_length = int.from_bytes(data[offset:offset + 2], "big")
            if segment_length < 2 or offset + segment_length > len(data):
                break
            if marker in {
                0xC0, 0xC1, 0xC2, 0xC3,
                0xC5, 0xC6, 0xC7,
                0xC9, 0xCA, 0xCB,
                0xCD, 0xCE, 0xCF,
            } and offset + 7 < len(data):
                height = int.from_bytes(data[offset + 3:offset + 5], "big")
                width = int.from_bytes(data[offset + 5:offset + 7], "big")
                return width, height
            offset += segment_length

    return None, None


def _mime_type_for_filename(filename: str | None) -> str:
    if not filename:
        return "application/octet-stream"
    mime_type, _ = mimetypes.guess_type(filename)
    return mime_type or "application/octet-stream"


def _register_image_asset(
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
    *,
    data: bytes,
    mime_type: str,
    filename: str | None,
) -> str:
    digest = hashlib.sha1(data).hexdigest()
    cache_key = (digest, mime_type)
    existing = asset_lookup.get(cache_key)
    if existing is not None:
        return existing

    image_id = f"img{len(assets) + 1}"
    intrinsic_width_px, intrinsic_height_px = _image_dimensions_from_bytes(data)
    assets[image_id] = ImageAsset.from_bytes(
        data=data,
        mime_type=mime_type,
        filename=filename,
        intrinsic_width_px=intrinsic_width_px,
        intrinsic_height_px=intrinsic_height_px,
    )
    asset_lookup[cache_key] = image_id
    return image_id


def _resolve_doc_metadata(
    *,
    source_path: str | Path | None,
    source_doc_type: str,
    metadata: dict[str, Any] | None,
    doc_id: str | None,
) -> tuple[str | None, str | None, dict[str, Any]]:
    resolved_source_path = str(source_path) if source_path is not None else None
    resolved_doc_id = doc_id
    if resolved_doc_id is None and source_path is not None:
        resolved_doc_id = Path(source_path).stem
    return resolved_doc_id, resolved_source_path, metadata or {}


def _docx_section_layouts(doc) -> list[dict[str, Any]]:
    from docx.oxml.ns import qn

    def _docx_column_style(sect_pr) -> ColumnLayoutInfo | None:
        cols = sect_pr.find(qn("w:cols")) if sect_pr is not None else None
        if cols is None:
            return None

        col_els = cols.findall(qn("w:col"))
        explicit_count = _safe_int(cols.get(qn("w:num")))
        count = explicit_count or len(col_els) or 1
        gap_pt = _twips_to_pt(cols.get(qn("w:space")))
        widths_pt = [
            width
            for col_el in col_els
            if (width := _twips_to_pt(col_el.get(qn("w:w")))) is not None
        ]
        gaps_pt = [
            gap
            for col_el in col_els
            if (gap := _twips_to_pt(col_el.get(qn("w:space")))) is not None
        ]
        equal_width_value = cols.get(qn("w:equalWidth"))
        equal_width = None if equal_width_value is None else equal_width_value not in {"0", "false", "False"}

        return ColumnLayoutInfo(
            count=max(count, 1),
            gap_pt=gap_pt,
            widths_pt=widths_pt,
            gaps_pt=gaps_pt,
            equal_width=equal_width,
        )

    return [
        {
            **_page_layout(
                width_pt=_emu_to_pt(int(section.page_width)) if section.page_width is not None else None,
                height_pt=_emu_to_pt(int(section.page_height)) if section.page_height is not None else None,
                margin_left_pt=_emu_to_pt(int(section.left_margin)) if section.left_margin is not None else None,
                margin_right_pt=_emu_to_pt(int(section.right_margin)) if section.right_margin is not None else None,
                margin_top_pt=_emu_to_pt(int(section.top_margin)) if section.top_margin is not None else None,
                margin_bottom_pt=_emu_to_pt(int(section.bottom_margin)) if section.bottom_margin is not None else None,
            ),
            "column_style": _docx_column_style(section._sectPr),
        }
        for section in doc.sections
    ] or [{**_page_layout(width_pt=None, height_pt=None), "column_style": None}]


def _docx_paragraph_has_page_break_before(paragraph) -> bool:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    page_break_before = p_pr.find(qn("w:pageBreakBefore"))
    if page_break_before is None:
        return False
    value = page_break_before.get(qn("w:val"))
    return value not in {"0", "false", "False"}


def _docx_paragraph_page_break_count(paragraph) -> int:
    count = 0
    for element in paragraph._p.iter():
        tag = getattr(element, "tag", None)
        if not isinstance(tag, str):
            continue
        local_name = tag.rsplit("}", 1)[-1]
        if local_name == "br" and element.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
            count += 1
    return count


def _docx_paragraph_section_break_type(paragraph) -> str | None:
    from docx.oxml.ns import qn

    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    sect_pr = p_pr.find(qn("w:sectPr"))
    if sect_pr is None:
        return None
    type_el = sect_pr.find(qn("w:type"))
    return type_el.get(qn("w:val")) if type_el is not None else "nextPage"


def _hwpx_section_page_layout(section_root: ET.Element) -> dict[str, float | None]:
    sec_pr = section_root.find(f".//{_HP}secPr")
    if sec_pr is None:
        return _page_layout(width_pt=None, height_pt=None)

    page_pr = sec_pr.find(f"{_HP}pagePr")
    if page_pr is None:
        return _page_layout(width_pt=None, height_pt=None)

    margin = page_pr.find(f"{_HP}margin")
    return _page_layout(
        width_pt=_hwpunit_to_pt(page_pr.get("width")),
        height_pt=_hwpunit_to_pt(page_pr.get("height")),
        margin_left_pt=_hwpunit_to_pt(margin.get("left")) if margin is not None else None,
        margin_right_pt=_hwpunit_to_pt(margin.get("right")) if margin is not None else None,
        margin_top_pt=_hwpunit_to_pt(margin.get("top")) if margin is not None else None,
        margin_bottom_pt=_hwpunit_to_pt(margin.get("bottom")) if margin is not None else None,
    )


def _hwpx_column_style_from_col_pr(col_pr: ET.Element) -> ColumnLayoutInfo:
    count = _safe_int(col_pr.get("colCount")) or 1
    same_gap_pt = _hwpunit_to_pt(col_pr.get("sameGap"))
    same_sz = col_pr.get("sameSz")
    equal_width = None if same_sz is None else same_sz not in {"0", "false", "False"}

    widths_pt: list[float] = []
    gaps_pt: list[float] = []
    for child in list(col_pr):
        local_name = child.tag.rsplit("}", 1)[-1]
        if local_name not in {"colSz", "col"}:
            continue
        width = _hwpunit_to_pt(child.get("width") or child.get("w"))
        gap = _hwpunit_to_pt(child.get("gap") or child.get("space"))
        if width is not None:
            widths_pt.append(width)
        if gap is not None:
            gaps_pt.append(gap)

    return ColumnLayoutInfo(
        count=max(count, 1),
        gap_pt=same_gap_pt,
        widths_pt=widths_pt,
        gaps_pt=gaps_pt,
        equal_width=equal_width,
    )


def _hwpx_paragraph_column_style(paragraph_el: ET.Element) -> ColumnLayoutInfo | None:
    col_prs = paragraph_el.findall(f"{_HP}run/{_HP}secPr/{_HP}colPr")
    col_prs.extend(paragraph_el.findall(f"{_HP}run/{_HP}ctrl/{_HP}colPr"))
    if not col_prs:
        return None
    return _hwpx_column_style_from_col_pr(col_prs[-1])


def _hwpx_section_column_style(section_root: ET.Element) -> ColumnLayoutInfo | None:
    sec_pr = section_root.find(f".//{_HP}secPr")
    col_pr = sec_pr.find(f"{_HP}colPr") if sec_pr is not None else None
    return _hwpx_column_style_from_col_pr(col_pr) if col_pr is not None else None


def _hwpx_paragraph_has_page_break_before(paragraph_el: ET.Element) -> bool:
    value = paragraph_el.get("pageBreak")
    return value not in {None, "", "0", "false", "False"}


def _hwpx_paragraph_vertpos(paragraph_el: ET.Element) -> int | None:
    line_seg = paragraph_el.find(f"{_HP}linesegarray/{_HP}lineseg")
    return _safe_int(line_seg.get("vertpos")) if line_seg is not None else None


def _parse_docx_run_images(
    run,
    *,
    paragraph_id: str,
    image_counter: int,
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[ImageIR], int]:
    images: list[ImageIR] = []

    for element in run._r.iter():
        tag = getattr(element, "tag", "")
        if not isinstance(tag, str) or not tag.endswith("}blip"):
            continue

        rel_id = element.get(_DOCX_EMBED_ATTR)
        if not rel_id:
            continue

        image_part = run.part.related_parts.get(rel_id)
        if image_part is None or not hasattr(image_part, "blob"):
            continue

        part_name = getattr(image_part, "partname", None)
        filename = Path(str(part_name)).name if part_name is not None else None
        mime_type = getattr(image_part, "content_type", None) or _mime_type_for_filename(filename)
        image_id = _register_image_asset(
            assets,
            asset_lookup,
            data=image_part.blob,
            mime_type=mime_type,
            filename=filename,
        )

        image_counter += 1
        drawing_parent = element
        while drawing_parent is not None and not (
            isinstance(getattr(drawing_parent, "tag", None), str)
            and drawing_parent.tag.rsplit("}", 1)[-1] in {"inline", "anchor"}
        ):
            drawing_parent = drawing_parent.getparent()

        extent_el = None if drawing_parent is None else next(
            (
                child
                for child in drawing_parent.iter()
                if isinstance(getattr(child, "tag", None), str)
                and child.tag.rsplit("}", 1)[-1] == "extent"
                and child.get("cx") is not None
            ),
            None,
        )
        doc_pr_el = None if drawing_parent is None else next(
            (
                child
                for child in drawing_parent.iter()
                if isinstance(getattr(child, "tag", None), str)
                and child.tag.rsplit("}", 1)[-1] == "docPr"
            ),
            None,
        )
        images.append(
            ImageIR(
                **_node_kwargs(
                    "image",
                    f"{paragraph_id}.img{image_counter}",
                    source_doc_type="docx",
                    parent_debug_path=paragraph_id,
                    part_name=str(part_name) if part_name is not None else None,
                ),
                image_id=image_id,
                alt_text=doc_pr_el.get("descr") if doc_pr_el is not None else None,
                title=doc_pr_el.get("name") if doc_pr_el is not None else None,
                display_width_pt=_emu_to_pt(extent_el.get("cx")) if extent_el is not None else None,
                display_height_pt=_emu_to_pt(extent_el.get("cy")) if extent_el is not None else None,
            )
        )

    return images, image_counter


def _parse_docx_paragraph_content(
    paragraph,
    paragraph_id: str,
    *,
    skip_empty: bool,
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[RunIR], list[ImageIR], list[object]]:
    runs: list[RunIR] = []
    images: list[ImageIR] = []
    content: list[object] = []
    image_counter = 0

    if not paragraph.runs:
        text = paragraph.text or ""
        if text or not skip_empty:
            run = RunIR(
                **_node_kwargs(
                    "run",
                    f"{paragraph_id}.r1",
                    source_doc_type="docx",
                    parent_debug_path=paragraph_id,
                    part_name="word/document.xml",
                    text=text,
                ),
                text=text,
            )
            runs.append(run)
            content.append(run)
        return runs, images, content

    for run_index, run in enumerate(paragraph.runs, start=1):
        text = run.text or ""
        run_images, image_counter = _parse_docx_run_images(
            run,
            paragraph_id=paragraph_id,
            image_counter=image_counter,
            assets=assets,
            asset_lookup=asset_lookup,
        )
        if text or (not skip_empty and not run_images):
            run_ir = RunIR(
                **_node_kwargs(
                    "run",
                    f"{paragraph_id}.r{run_index}",
                    source_doc_type="docx",
                    parent_debug_path=paragraph_id,
                    part_name="word/document.xml",
                    text=text,
                ),
                text=text,
            )
            runs.append(run_ir)
            content.append(run_ir)
        images.extend(run_images)
        content.extend(run_images)

    return runs, images, content


def _parse_docx_table(
    table,
    table_id: str,
    *,
    include_tables: bool,
    skip_empty: bool,
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
    CT_P,
    CT_Tbl,
    Paragraph,
    Table,
) -> TableIR:
    table_ir = TableIR(
        **_node_kwargs(
            "table",
            table_id,
            source_doc_type="docx",
            parent_debug_path=table_id.rsplit(".", 1)[0] if "." in table_id else None,
            part_name="word/document.xml",
        )
    )

    for tr_idx, row in enumerate(table.rows, start=1):
        for tc_idx, cell in enumerate(row.cells, start=1):
            cell_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}"
            cell_ir = TableCellIR(
                **_node_kwargs(
                    "cell",
                    cell_id,
                    source_doc_type="docx",
                    parent_debug_path=table_id,
                    part_name="word/document.xml",
                ),
                row_index=tr_idx,
                col_index=tc_idx,
            )
            table_ir.cells.append(cell_ir)

            cp_idx = 0
            current_paragraph: ParagraphIR | None = None
            nested_table_counter_by_paragraph: dict[str, int] = {}

            for block in _iter_blocks_from_element(
                cell,
                cell._tc,
                CT_P=CT_P,
                CT_Tbl=CT_Tbl,
                Paragraph=Paragraph,
                Table=Table,
            ):
                if isinstance(block, Paragraph):
                    cp_idx += 1
                    paragraph_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                    runs, images, content = _parse_docx_paragraph_content(
                        block,
                        paragraph_id,
                        skip_empty=skip_empty,
                        assets=assets,
                        asset_lookup=asset_lookup,
                    )
                    current_paragraph = ParagraphIR(
                        **_node_kwargs(
                            "paragraph",
                            paragraph_id,
                            source_doc_type="docx",
                            parent_debug_path=cell_id,
                            part_name="word/document.xml",
                        ),
                        content=content,
                    )
                    current_paragraph.recompute_text()
                    if current_paragraph.content or current_paragraph.text or not skip_empty:
                        cell_ir.paragraphs.append(current_paragraph)
                    continue

                if not include_tables or not isinstance(block, Table):
                    continue

                if current_paragraph is None:
                    cp_idx += 1
                    paragraph_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                    current_paragraph = ParagraphIR(
                        **_node_kwargs(
                            "paragraph",
                            paragraph_id,
                            source_doc_type="docx",
                            parent_debug_path=cell_id,
                            part_name="word/document.xml",
                        ),
                    )
                    cell_ir.paragraphs.append(current_paragraph)

                current_paragraph_id = current_paragraph.native_anchor.debug_path
                table_counter = nested_table_counter_by_paragraph.get(current_paragraph_id, 0) + 1
                nested_table_counter_by_paragraph[current_paragraph_id] = table_counter
                nested_table = _parse_docx_table(
                    block,
                    f"{current_paragraph_id}.tbl{table_counter}",
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    assets=assets,
                    asset_lookup=asset_lookup,
                    CT_P=CT_P,
                    CT_Tbl=CT_Tbl,
                    Paragraph=Paragraph,
                    Table=Table,
                )
                current_paragraph.append_content(nested_table)
                current_paragraph.recompute_text()

            cell_ir.recompute_text()

    return table_ir


def _build_docx_doc_ir(
    source: "DocxDocument | str | Path | bytes",
    *,
    include_tables: bool,
    skip_empty: bool,
    source_path: str | Path | None,
    metadata: dict[str, Any] | None,
    doc_id: str | None,
    doc_cls: type[DocIR] | None,
    **doc_kwargs: Any,
) -> DocIR:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = _load_docx_source(source)
    paragraphs: list[ParagraphIR] = []
    assets: dict[str, ImageAsset] = {}
    asset_lookup: dict[tuple[str, str], str] = {}
    pages: "OrderedDict[int, PageInfo]" = OrderedDict()
    section_layouts = _docx_section_layouts(doc)
    current_section_index = 0
    current_page_number = 1
    has_seen_content = False

    p_idx = 0
    table_counter = 0

    for block in _iter_blocks(
        doc,
        CT_P=CT_P,
        CT_Tbl=CT_Tbl,
        Paragraph=Paragraph,
        Table=Table,
    ):
        current_layout = section_layouts[min(current_section_index, len(section_layouts) - 1)]
        if isinstance(block, Paragraph):
            if has_seen_content and _docx_paragraph_has_page_break_before(block):
                current_page_number += 1

            p_idx += 1
            paragraph_id = f"s1.p{p_idx}"
            _ensure_page_info(
                pages,
                page_number=current_page_number,
                layout=current_layout,
            )
            runs, images, content = _parse_docx_paragraph_content(
                block,
                paragraph_id,
                skip_empty=skip_empty,
                assets=assets,
                asset_lookup=asset_lookup,
            )
            paragraph_ir = ParagraphIR(
                **_node_kwargs(
                    "paragraph",
                    paragraph_id,
                    source_doc_type="docx",
                    part_name="word/document.xml",
                ),
                page_number=current_page_number,
                para_style=_para_style_with_columns(current_layout.get("column_style")),
                content=content,
            )
            _assign_page_number_to_paragraph(paragraph_ir, current_page_number)
            paragraph_ir.recompute_text()
            if paragraph_ir.content or paragraph_ir.text or not skip_empty:
                paragraphs.append(paragraph_ir)

            has_seen_content = True
            current_page_number += _docx_paragraph_page_break_count(block)
            section_break_type = _docx_paragraph_section_break_type(block)
            if section_break_type is not None:
                if current_section_index < len(section_layouts) - 1:
                    current_section_index += 1
                if section_break_type != "continuous":
                    current_page_number += 1
            continue

        if not include_tables or not isinstance(block, Table):
            continue

        table_counter += 1
        p_idx += 1
        paragraph_id = f"s1.p{p_idx}"
        _ensure_page_info(
            pages,
            page_number=current_page_number,
            layout=current_layout,
        )
        table_ir = _parse_docx_table(
            block,
            f"{paragraph_id}.r1.tbl{table_counter}",
            include_tables=include_tables,
            skip_empty=skip_empty,
            assets=assets,
            asset_lookup=asset_lookup,
            CT_P=CT_P,
            CT_Tbl=CT_Tbl,
            Paragraph=Paragraph,
            Table=Table,
        )
        paragraph_ir = ParagraphIR(
            **_node_kwargs(
                "paragraph",
                paragraph_id,
                source_doc_type="docx",
                part_name="word/document.xml",
            ),
            page_number=current_page_number,
            para_style=_para_style_with_columns(current_layout.get("column_style")),
            content=[table_ir],
        )
        _assign_page_number_to_paragraph(paragraph_ir, current_page_number)
        paragraph_ir.recompute_text()
        paragraphs.append(paragraph_ir)
        has_seen_content = True

    resolved_doc_id, resolved_source_path, resolved_metadata = _resolve_doc_metadata(
        source_path=source_path,
        source_doc_type="docx",
        metadata=metadata,
        doc_id=doc_id,
    )
    resolved_doc_cls = doc_cls or DocIR
    return resolved_doc_cls(
        doc_id=resolved_doc_id,
        source_path=resolved_source_path,
        source_doc_type="docx",
        metadata=resolved_metadata,
        assets=assets,
        pages=list(pages.values()),
        paragraphs=paragraphs,
        **doc_kwargs,
    )


def _hwpx_binary_name_map(zf: zipfile.ZipFile) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for name in zf.namelist():
        if not name.startswith("BinData/"):
            continue
        mapping[Path(name).stem.lower()] = name
    return mapping


def _find_hwpx_binary_path(binary_name_map: dict[str, str], binary_item_id: str | None) -> str | None:
    if not binary_item_id:
        return None
    return binary_name_map.get(binary_item_id.lower())


def _hwpx_picture_display_size(pic_el: ET.Element) -> tuple[float | None, float | None]:
    size_el = pic_el.find(f"{_HP}sz")
    if size_el is not None:
        width_pt = _hwpunit_to_pt(size_el.get("width"))
        height_pt = _hwpunit_to_pt(size_el.get("height"))
        if width_pt is not None or height_pt is not None:
            return width_pt, height_pt

    dim_el = pic_el.find(f"{_HP}imgDim")
    if dim_el is None:
        return None, None
    return _hwpunit_to_pt(dim_el.get("dimwidth")), _hwpunit_to_pt(dim_el.get("dimheight"))


def _parse_hwpx_picture_image(
    pic_el: ET.Element,
    *,
    paragraph_id: str,
    image_counter: int,
    archive: zipfile.ZipFile,
    binary_name_map: dict[str, str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[ImageIR | None, int]:
    img_el = next(
        (
            element
            for element in pic_el.iter()
            if isinstance(element.tag, str)
            and element.tag.rsplit("}", 1)[-1] == _HC_IMG_TAG
            and element.get("binaryItemIDRef")
        ),
        None,
    )
    if img_el is None:
        return None, image_counter

    binary_path = _find_hwpx_binary_path(binary_name_map, img_el.get("binaryItemIDRef"))
    if binary_path is None:
        return None, image_counter

    data = archive.read(binary_path)
    filename = Path(binary_path).name
    image_id = _register_image_asset(
        assets,
        asset_lookup,
        data=data,
        mime_type=_mime_type_for_filename(filename),
        filename=filename,
    )
    display_width_pt, display_height_pt = _hwpx_picture_display_size(pic_el)

    image_counter += 1
    return ImageIR(
        **_node_kwargs(
            "image",
            f"{paragraph_id}.img{image_counter}",
            source_doc_type="hwpx",
            parent_debug_path=paragraph_id,
            part_name=binary_path,
        ),
        image_id=image_id,
        display_width_pt=display_width_pt,
        display_height_pt=display_height_pt,
    ), image_counter


def _hwpx_text_element_text(text_el: ET.Element) -> str:
    return "".join(text_el.itertext())


def _parse_hwpx_paragraph_content(
    paragraph_el: ET.Element,
    paragraph_id: str,
    *,
    table_id_builder: Callable[[int], str],
    include_tables: bool,
    skip_empty: bool,
    archive: zipfile.ZipFile,
    binary_name_map: dict[str, str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> tuple[list[RunIR], list[ImageIR], list[TableIR], list[object]]:
    runs: list[RunIR] = []
    images: list[ImageIR] = []
    tables: list[TableIR] = []
    content: list[object] = []
    image_counter = 0
    table_counter = 0
    content_position = 0

    run_els = paragraph_el.findall(f"{_HP}run")
    if not run_els:
        text = _paragraph_text(paragraph_el)
        if text or not skip_empty:
            run = RunIR(
                **_node_kwargs(
                    "run",
                    f"{paragraph_id}.r1",
                    source_doc_type="hwpx",
                    parent_debug_path=paragraph_id,
                    text=text,
                ),
                text=text,
            )
            runs.append(run)
            content.append(run)
        return runs, images, tables, content

    for run_el in run_els:
        pending_text_parts: list[str] = []
        emitted_run_content = False
        saw_run_level_content = False

        def flush_pending_text() -> None:
            nonlocal content_position, emitted_run_content
            if not pending_text_parts:
                return

            text = "".join(pending_text_parts)
            pending_text_parts.clear()
            if not text and (skip_empty or saw_run_level_content):
                return

            content_position += 1
            run = RunIR(
                **_node_kwargs(
                    "run",
                    f"{paragraph_id}.r{content_position}",
                    source_doc_type="hwpx",
                    parent_debug_path=paragraph_id,
                    text=text,
                ),
                text=text,
            )
            runs.append(run)
            content.append(run)
            emitted_run_content = True

        for child in list(run_el):
            tag = child.tag
            if tag == f"{_HP}t":
                pending_text_parts.append(_hwpx_text_element_text(child))
                continue

            if tag not in {f"{_HP}pic", f"{_HP}tbl"}:
                continue

            flush_pending_text()
            saw_run_level_content = True

            if tag == f"{_HP}pic":
                image_ir, image_counter = _parse_hwpx_picture_image(
                    child,
                    paragraph_id=paragraph_id,
                    image_counter=image_counter,
                    archive=archive,
                    binary_name_map=binary_name_map,
                    assets=assets,
                    asset_lookup=asset_lookup,
                )
                if image_ir is not None:
                    content_position += 1
                    images.append(image_ir)
                    content.append(image_ir)
                    emitted_run_content = True
                continue

            if not include_tables:
                continue

            table_counter += 1
            table_ir = _parse_hwpx_table(
                child,
                table_id_builder(table_counter),
                include_tables=include_tables,
                skip_empty=skip_empty,
                archive=archive,
                binary_name_map=binary_name_map,
                assets=assets,
                asset_lookup=asset_lookup,
            )
            content_position += 1
            tables.append(table_ir)
            content.append(table_ir)
            emitted_run_content = True

        flush_pending_text()

        if not emitted_run_content and not skip_empty:
            content_position += 1
            run = RunIR(
                **_node_kwargs(
                    "run",
                    f"{paragraph_id}.r{content_position}",
                    source_doc_type="hwpx",
                    parent_debug_path=paragraph_id,
                    text="",
                ),
                text="",
            )
            runs.append(run)
            content.append(run)

    return runs, images, tables, content


def _parse_hwpx_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    sub_list = cell_el.find(f"{_HP}subList")
    if sub_list is None:
        return []
    return [child for child in list(sub_list) if child.tag == f"{_HP}p"]


def _parse_hwpx_table(
    table_el: ET.Element,
    table_id: str,
    *,
    include_tables: bool,
    skip_empty: bool,
    archive: zipfile.ZipFile,
    binary_name_map: dict[str, str],
    assets: dict[str, ImageAsset],
    asset_lookup: dict[tuple[str, str], str],
) -> TableIR:
    table_ir = TableIR(
        **_node_kwargs(
            "table",
            table_id,
            source_doc_type="hwpx",
            parent_debug_path=table_id.rsplit(".", 1)[0] if "." in table_id else None,
        )
    )

    for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
        for tc_idx, cell_el in _logical_table_cells(row_el):
            cell_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}"
            cell_ir = TableCellIR(
                **_node_kwargs(
                    "cell",
                    cell_id,
                    source_doc_type="hwpx",
                    parent_debug_path=table_id,
                ),
                row_index=tr_idx,
                col_index=tc_idx,
            )
            table_ir.cells.append(cell_ir)

            cell_paragraphs = _parse_hwpx_cell_paragraphs(cell_el)
            if not cell_paragraphs:
                if not skip_empty:
                    paragraph_ir = ParagraphIR(
                        **_node_kwargs(
                            "paragraph",
                            f"{cell_id}.p1",
                            source_doc_type="hwpx",
                            parent_debug_path=cell_id,
                        ),
                        content=[
                            RunIR(
                                **_node_kwargs(
                                    "run",
                                    f"{cell_id}.p1.r1",
                                    source_doc_type="hwpx",
                                    parent_debug_path=f"{cell_id}.p1",
                                    text="",
                                ),
                                text="",
                            )
                        ],
                    )
                    paragraph_ir.recompute_text()
                    cell_ir.paragraphs.append(paragraph_ir)
                cell_ir.recompute_text()
                continue

            for cp_idx, paragraph_el in enumerate(cell_paragraphs, start=1):
                paragraph_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}.p{cp_idx}"
                runs, images, tables, content = _parse_hwpx_paragraph_content(
                    paragraph_el,
                    paragraph_id,
                    table_id_builder=lambda counter, base=paragraph_id: f"{base}.tbl{counter}",
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    archive=archive,
                    binary_name_map=binary_name_map,
                    assets=assets,
                    asset_lookup=asset_lookup,
                )
                paragraph_ir = ParagraphIR(
                    **_node_kwargs(
                        "paragraph",
                        paragraph_id,
                        source_doc_type="hwpx",
                        parent_debug_path=cell_id,
                    ),
                    content=content,
                )
                paragraph_ir.recompute_text()
                if paragraph_ir.content or paragraph_ir.text or not skip_empty:
                    cell_ir.paragraphs.append(paragraph_ir)

            cell_ir.recompute_text()

    return table_ir


def _build_hwpx_doc_ir(
    source: "HwpxDocument | str | Path | bytes",
    *,
    include_tables: bool,
    skip_empty: bool,
    source_path: str | Path | None,
    metadata: dict[str, Any] | None,
    doc_id: str | None,
    doc_cls: type[DocIR] | None,
    **doc_kwargs: Any,
) -> DocIR:
    if isinstance(source, bytes):
        hwpx_bytes = source
    elif isinstance(source, (str, Path)):
        hwpx_bytes = Path(source).read_bytes()
    else:
        hwpx_bytes = coerce_source_to_supported_value(source, doc_type="hwpx")  # type: ignore[arg-type]
        if not isinstance(hwpx_bytes, bytes):
            raise TypeError(f"Unsupported HWPX source type: {type(source)!r}")

    section_roots = _section_roots_from_bytes(hwpx_bytes)
    assets: dict[str, ImageAsset] = {}
    asset_lookup: dict[tuple[str, str], str] = {}
    paragraphs: list[ParagraphIR] = []
    pages: "OrderedDict[int, PageInfo]" = OrderedDict()
    page_offset = 0

    with zipfile.ZipFile(BytesIO(hwpx_bytes)) as archive:
        binary_name_map = _hwpx_binary_name_map(archive)

        for s_idx, section_root in enumerate(section_roots, start=1):
            section_layout = _hwpx_section_page_layout(section_root)
            current_column_style = _hwpx_section_column_style(section_root)
            section_page_number = 1
            last_vertpos: int | None = None
            saw_paragraph = False

            for p_idx, paragraph_el in enumerate(section_root.findall(f"{_HP}p"), start=1):
                paragraph_column_style = _hwpx_paragraph_column_style(paragraph_el)
                if paragraph_column_style is not None:
                    current_column_style = paragraph_column_style

                if saw_paragraph and _hwpx_paragraph_has_page_break_before(paragraph_el):
                    section_page_number += 1
                    last_vertpos = None

                vertpos = _hwpx_paragraph_vertpos(paragraph_el)
                if saw_paragraph and last_vertpos is not None and vertpos is not None and vertpos < last_vertpos:
                    section_page_number += 1

                absolute_page_number = page_offset + section_page_number
                _ensure_page_info(
                    pages,
                    page_number=absolute_page_number,
                    layout=section_layout,
                )
                paragraph_id = f"s{s_idx}.p{p_idx}"
                runs, images, tables, content = _parse_hwpx_paragraph_content(
                    paragraph_el,
                    paragraph_id,
                    table_id_builder=lambda counter, base=paragraph_id: f"{base}.r1.tbl{counter}",
                    include_tables=include_tables,
                    skip_empty=skip_empty,
                    archive=archive,
                    binary_name_map=binary_name_map,
                    assets=assets,
                    asset_lookup=asset_lookup,
                )

                paragraph_ir = ParagraphIR(
                    **_node_kwargs(
                        "paragraph",
                        paragraph_id,
                        source_doc_type="hwpx",
                        part_name=f"Contents/section{s_idx - 1}.xml",
                    ),
                    page_number=absolute_page_number,
                    para_style=_para_style_with_columns(current_column_style),
                    content=content,
                )
                _assign_page_number_to_paragraph(paragraph_ir, absolute_page_number)
                paragraph_ir.recompute_text()
                if paragraph_ir.content or paragraph_ir.text or not skip_empty:
                    paragraphs.append(paragraph_ir)

                if vertpos is not None:
                    last_vertpos = vertpos
                saw_paragraph = True

            page_offset += section_page_number

    resolved_doc_id, resolved_source_path, resolved_metadata = _resolve_doc_metadata(
        source_path=source_path,
        source_doc_type="hwpx",
        metadata=metadata,
        doc_id=doc_id,
    )
    resolved_doc_cls = doc_cls or DocIR
    return resolved_doc_cls(
        doc_id=resolved_doc_id,
        source_path=resolved_source_path,
        source_doc_type="hwpx",
        metadata=resolved_metadata,
        assets=assets,
        pages=list(pages.values()),
        paragraphs=paragraphs,
        **doc_kwargs,
    )


def build_doc_ir_from_file(
    source: "HwpxDocument | DocxDocument | str | Path | bytes",
    *,
    doc_type: DocType,
    include_tables: bool = True,
    skip_empty: bool = False,
    source_path: str | Path | None = None,
    metadata: dict[str, Any] | None = None,
    doc_id: str | None = None,
    doc_cls: type[DocIR] | None = None,
    **doc_kwargs: Any,
) -> DocIR:
    """Build document IR directly from a document source."""
    if doc_type == "pdf":
        from ..pdf.pipeline import parse_pdf_to_doc_ir

        if not isinstance(source, (str, Path)):
            raise TypeError("PDF parsing currently requires a filesystem path.")
        return parse_pdf_to_doc_ir(
            source,
            doc_id=doc_id,
            doc_cls=doc_cls,
            **doc_kwargs,
        )

    if doc_type == "docx":
        return _build_docx_doc_ir(
            source,
            include_tables=include_tables,
            skip_empty=skip_empty,
            source_path=source_path,
            metadata=metadata,
            doc_id=doc_id,
            doc_cls=doc_cls,
            **doc_kwargs,
        )

    if doc_type == "hwp":
        if not isinstance(source, (str, Path)):
            raise TypeError("HWP conversion currently requires a filesystem path.")
        return _build_hwpx_doc_ir(
            convert_hwp_to_hwpx_bytes(source),
            include_tables=include_tables,
            skip_empty=skip_empty,
            source_path=source_path,
            metadata=metadata,
            doc_id=doc_id,
            doc_cls=doc_cls,
            **doc_kwargs,
        )

    return _build_hwpx_doc_ir(
        source,
        include_tables=include_tables,
        skip_empty=skip_empty,
        source_path=source_path,
        metadata=metadata,
        doc_id=doc_id,
        doc_cls=doc_cls,
        **doc_kwargs,
    )


__all__ = ["build_doc_ir_from_file"]
