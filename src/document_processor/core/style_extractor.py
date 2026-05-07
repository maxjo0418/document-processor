"""Unified style extraction for HWP/HWPX/DOCX documents."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from typing import TYPE_CHECKING, Literal
from xml.etree import ElementTree as ET
import zipfile

from ..io_utils import TemporarySourcePath, infer_doc_type
from ..style_types import (
    CellStyleInfo,
    ColumnLayoutInfo,
    ListItemInfo,
    ParaStyleInfo,
    RunStyleInfo,
    StyleMap,
    TableStyleInfo,
)
from .hwp_converter import convert_hwp_to_hwpx_bytes

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from ..hwpx import HwpxDocument


DocType = Literal["auto", "hwp", "hwpx", "docx", "pdf"]

_NS_HH = "http://www.hancom.co.kr/hwpml/2011/head"
_NS_HC = "http://www.hancom.co.kr/hwpml/2011/core"
_NS_HP = "http://www.hancom.co.kr/hwpml/2011/paragraph"
_HP = f"{{{_NS_HP}}}"

_HWPUNIT_PER_PT = 100.0
_HWPX_INHERIT_UINT = "4294967295"

_HWPX_BORDER_STYLE = {
    "SOLID": "solid",
    "DASH": "dashed",
    "DOT": "dotted",
    "DASH_DOT": "dashed",
    "DASH_DOT_DOT": "dotted",
    "DOUBLE": "double",
    "NONE": "none",
}

_HWPX_HALIGN = {
    "LEFT": "left",
    "CENTER": "center",
    "RIGHT": "right",
    "JUSTIFY": "justify",
    "DISTRIBUTE": "justify",
}

_HWPX_VALIGN = {
    "TOP": "top",
    "CENTER": "center",
    "BOTTOM": "bottom",
    "BASELINE": "top",
}

_DOCX_ALIGN = {0: "left", 1: "center", 2: "right", 3: "justify"}

_HWPX_VISIBLE_LINE_SHAPES = {
    "SOLID",
    "DASH",
    "DOT",
    "DASH_DOT",
    "DASH_DOT_DOT",
    "LONG_DASH",
    "CIRCLE",
    "DOUBLE",
    "SLIM_THICK",
    "THICK_SLIM",
    "SLIM_THICK_SLIM",
}

_XML_1_0_INVALID_CHAR_RE = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F\uD800-\uDFFF\uFFFE\uFFFF]"
)
_BARE_XML_AMPERSAND_RE = re.compile(
    r"&(?!#\d+;|#x[0-9A-Fa-f]+;|[A-Za-z][A-Za-z0-9]*;)"
)


@dataclass(frozen=True)
class _ListLevelDefinition:
    list_id: str
    level: int
    marker_type: str | None = None
    marker_text: str | None = None
    start: int = 1
    bullet_char: str | None = None
    left_indent_pt: float | None = None
    first_line_indent_pt: float | None = None
    hanging_indent_pt: float | None = None


@dataclass
class _ListCounterState:
    counters: dict[str, dict[int, int]]

    @classmethod
    def create(cls) -> "_ListCounterState":
        return cls(counters={})


def _alpha_counter(value: int, *, uppercase: bool) -> str:
    if value <= 0:
        return str(value)
    chars: list[str] = []
    current = value
    while current:
        current -= 1
        chars.append(chr(ord("A") + current % 26))
        current //= 26
    text = "".join(reversed(chars))
    return text if uppercase else text.lower()


def _roman_counter(value: int, *, uppercase: bool) -> str:
    if value <= 0 or value >= 4000:
        return str(value)
    numerals = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    parts: list[str] = []
    remaining = value
    for number, marker in numerals:
        while remaining >= number:
            parts.append(marker)
            remaining -= number
    text = "".join(parts)
    return text if uppercase else text.lower()


def _hangul_syllable_counter(value: int) -> str:
    sequence = ("가", "나", "다", "라", "마", "바", "사", "아", "자", "차", "카", "타", "파", "하")
    if 1 <= value <= len(sequence):
        return sequence[value - 1]
    return str(value)


def _circled_digit_counter(value: int) -> str:
    if 1 <= value <= 20:
        return chr(0x2460 + value - 1)
    return str(value)


def _normalize_bullet_marker(value: str | None) -> str:
    if not value:
        return "\u2022"
    return {
        "\uf0b7": "\u2022",
        "\uf0a7": "\u25aa",
        "\uf0d8": "\u25e6",
    }.get(value, value)


def _format_counter(value: int, marker_type: str | None) -> str:
    normalized = (marker_type or "decimal").lower()
    if normalized in {"lowerletter", "lower_letter", "lower-alpha", "loweralpha"}:
        return _alpha_counter(value, uppercase=False)
    if normalized in {"upperletter", "upper_letter", "upper-alpha", "upperalpha"}:
        return _alpha_counter(value, uppercase=True)
    if normalized in {"lowerroman", "lower_roman"}:
        return _roman_counter(value, uppercase=False)
    if normalized in {"upperroman", "upper_roman"}:
        return _roman_counter(value, uppercase=True)
    if normalized in {"decimalzero", "decimal_zero"}:
        return f"{value:02d}"
    if normalized in {"hangul_syllable", "hangulsyllable"}:
        return _hangul_syllable_counter(value)
    if normalized in {"circled_digit", "decimalenclosedcircle"}:
        return _circled_digit_counter(value)
    return str(value)


def _list_marker_from_pattern(
    marker_text: str | None,
    *,
    current_level: int,
    level_definitions: dict[int, _ListLevelDefinition],
    level_counters: dict[int, int],
    fallback: str,
    source: Literal["docx", "hwpx"],
) -> str:
    if not marker_text:
        return fallback

    def marker_type_for_level(level: int) -> str | None:
        definition = level_definitions.get(level) or level_definitions.get(current_level)
        return definition.marker_type if definition is not None else None

    if source == "docx":
        def replace_docx(match: re.Match[str]) -> str:
            level = int(match.group(1)) - 1
            value = level_counters.get(level)
            if value is None:
                value = level_definitions.get(level, _ListLevelDefinition("", level)).start
            return _format_counter(value, marker_type_for_level(level))

        return re.sub(r"%([1-9])", replace_docx, marker_text)

    def replace_hwpx(match: re.Match[str]) -> str:
        token = match.group(1)
        if token in {"n", "N"}:
            numbers = [
                _format_counter(
                    level_counters.get(level, definition.start),
                    definition.marker_type,
                )
                for level, definition in sorted(level_definitions.items())
                if level <= current_level
            ]
            text = ".".join(numbers)
            return f"{text}." if token == "N" and text else text

        level = int(token) - 1
        value = level_counters.get(level)
        if value is None:
            value = level_definitions.get(level, _ListLevelDefinition("", level)).start
        return _format_counter(value, marker_type_for_level(level))

    return re.sub(r"\^([1-7nN])", replace_hwpx, marker_text)


def _advance_list_counter(
    state: _ListCounterState,
    definition: _ListLevelDefinition,
    *,
    level_definitions: dict[int, _ListLevelDefinition],
    source: Literal["docx", "hwpx"],
) -> ListItemInfo:
    list_counters = state.counters.setdefault(definition.list_id, {})
    if definition.marker_type and definition.marker_type.lower() == "bullet":
        marker = _normalize_bullet_marker(definition.bullet_char or definition.marker_text)
        return ListItemInfo(
            list_id=definition.list_id,
            level=max(definition.level, 0),
            marker=marker,
            marker_type="bullet",
            marker_text=definition.marker_text,
        )

    previous = list_counters.get(definition.level)
    list_counters[definition.level] = definition.start if previous is None else previous + 1
    for level in list(list_counters):
        if level > definition.level:
            del list_counters[level]

    current_value = list_counters[definition.level]
    fallback = _format_counter(current_value, definition.marker_type)
    marker = _list_marker_from_pattern(
        definition.marker_text,
        current_level=definition.level,
        level_definitions=level_definitions,
        level_counters=list_counters,
        fallback=fallback,
        source=source,
    )
    return ListItemInfo(
        list_id=definition.list_id,
        level=max(definition.level, 0),
        marker=marker,
        marker_type=definition.marker_type,
        marker_text=definition.marker_text,
    )


def _has_para_style(info: ParaStyleInfo) -> bool:
    return any(
        value is not None
        for value in (
            info.align,
            info.left_indent_pt,
            info.right_indent_pt,
            info.first_line_indent_pt,
            info.hanging_indent_pt,
            info.column_layout,
            info.list_info,
        )
    )


def _safe_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _xml_attr(el: ET.Element | None, *names: str) -> str | None:
    if el is None:
        return None
    for name in names:
        value = el.get(name)
        if value is not None:
            return value
    lower_names = {name.lower() for name in names}
    for key, value in el.attrib.items():
        if key.lower() in lower_names:
            return value
    return None


def _length_to_pt(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value) / 12700.0
    except (TypeError, ValueError):
        return None


def _docx_measure_to_pt(raw_value: str | None, measurement_type: str | None) -> float | None:
    if raw_value is None:
        return None
    try:
        number = float(raw_value)
    except (TypeError, ValueError):
        return None

    normalized_type = (measurement_type or "dxa").lower()
    if normalized_type == "dxa":
        return number / 20.0
    if normalized_type == "nil":
        return 0.0
    return None


def _docx_cell_margin_to_padding(margin_el) -> dict[str, float]:
    from docx.oxml.ns import qn

    padding: dict[str, float] = {}
    if margin_el is None:
        return padding

    for side in ("top", "right", "bottom", "left"):
        side_el = margin_el.find(qn(f"w:{side}"))
        if side_el is None:
            continue
        value_pt = _docx_measure_to_pt(
            side_el.get(qn("w:w")),
            side_el.get(qn("w:type")),
        )
        if value_pt is not None:
            padding[side] = value_pt

    return padding


def _apply_cell_padding(info: CellStyleInfo, padding: dict[str, float]) -> None:
    if "top" in padding:
        info.padding_top_pt = padding["top"]
    if "right" in padding:
        info.padding_right_pt = padding["right"]
    if "bottom" in padding:
        info.padding_bottom_pt = padding["bottom"]
    if "left" in padding:
        info.padding_left_pt = padding["left"]


def _hwp_numeric_to_pt(raw_value: str | None) -> float | None:
    if raw_value is None:
        return None
    raw_text = str(raw_value).strip()
    if raw_text in {_HWPX_INHERIT_UINT, "-1"}:
        return None
    try:
        return float(raw_text) / _HWPUNIT_PER_PT
    except (TypeError, ValueError):
        return None


def _hwp_margin_value_to_pt(el: ET.Element | None) -> float | None:
    if el is None:
        return None

    raw = el.get("value")
    if raw is None:
        return None
    raw_text = str(raw).strip()
    if raw_text == _HWPX_INHERIT_UINT:
        return None

    try:
        number = float(raw_text)
    except (TypeError, ValueError):
        return None

    unit = (el.get("unit") or "HWPUNIT").upper()
    if unit == "HWPUNIT":
        return number / _HWPUNIT_PER_PT
    if unit == "PT":
        return number
    return number


def _hwpx_table_cell_padding_defaults(table_el: ET.Element) -> dict[str, float]:
    margin_el = table_el.find(f"{_HP}inMargin")
    if margin_el is None:
        return {}

    padding: dict[str, float] = {}
    for side in ("top", "right", "bottom", "left"):
        value_pt = _hwp_numeric_to_pt(margin_el.get(side))
        if value_pt is not None:
            padding[side] = value_pt
    return padding


def _apply_hwpx_cell_margin(
    info: CellStyleInfo,
    margin_el: ET.Element | None,
    *,
    defaults: dict[str, float] | None = None,
) -> None:
    padding = dict(defaults or {})
    if margin_el is None and not padding:
        return

    if margin_el is not None:
        for side in ("top", "right", "bottom", "left"):
            value_pt = _hwp_numeric_to_pt(margin_el.get(side))
            if value_pt is not None:
                padding[side] = value_pt

    _apply_cell_padding(info, padding)


def _hwpx_border_css(border_el: ET.Element | None) -> str | None:
    if border_el is None:
        return None
    btype = border_el.get("type", "NONE")
    if btype == "NONE":
        return None
    width = border_el.get("width", "0.12 mm")
    try:
        mm_val = float(width.replace("mm", "").strip())
        px = max(1, round(mm_val * 3.78))
    except (ValueError, AttributeError):
        px = 1
    color = border_el.get("color", "#000000")
    style = _HWPX_BORDER_STYLE.get(btype, "solid")
    return f"{px}px {style} {color}"


def _hwpx_diagonal_border_css(
    border_fill_el: ET.Element,
    *,
    direction: Literal["slash", "backslash"],
) -> str | None:
    if direction == "slash":
        direction_el = border_fill_el.find(f"{{{_NS_HH}}}slash")
    else:
        direction_el = border_fill_el.find(f"{{{_NS_HH}}}backSlash")

    if direction_el is None or direction_el.get("type", "NONE") == "NONE":
        return None

    diagonal_el = border_fill_el.find(f"{{{_NS_HH}}}diagonal")
    return _hwpx_border_css(diagonal_el)


def _map_by_id(root: ET.Element | None, tag: str) -> dict[str, ET.Element]:
    if root is None:
        return {}
    out: dict[str, ET.Element] = {}
    for el in root.findall(f".//{{{_NS_HH}}}{tag}"):
        el_id = el.get("id")
        if el_id:
            out[el_id] = el
    return out


def _normalize_hwpx_marker_type(value: str | None) -> str | None:
    normalized = (value or "").strip().upper()
    return {
        "DIGIT": "decimal",
        "DECIMAL": "decimal",
        "HANGUL_SYLLABLE": "hangul_syllable",
        "CIRCLED_DIGIT": "circled_digit",
        "ROMAN_CAPITAL": "upperRoman",
        "ROMAN_SMALL": "lowerRoman",
        "LATIN_CAPITAL": "upperLetter",
        "LATIN_SMALL": "lowerLetter",
    }.get(normalized, normalized.lower() or None)


def _hwpx_para_head_text(para_head_el: ET.Element | None) -> str | None:
    if para_head_el is None:
        return None
    text = "".join(para_head_el.itertext())
    return text or None


def _hwpx_list_level_from_para_head(para_head_el: ET.Element | None) -> int:
    raw_level = _safe_int(_xml_attr(para_head_el, "level")) or 0
    if raw_level <= 0:
        return 0
    return raw_level - 1


def _hwpx_numbering_definitions(header_root: ET.Element | None) -> dict[str, dict[int, _ListLevelDefinition]]:
    if header_root is None:
        return {}

    definitions: dict[str, dict[int, _ListLevelDefinition]] = {}
    for numbering_el in header_root.findall(f".//{{{_NS_HH}}}numbering"):
        native_id = _xml_attr(numbering_el, "id")
        if not native_id:
            continue
        list_id = f"hwpx_number_{native_id}"
        start = _safe_int(_xml_attr(numbering_el, "start")) or 1
        levels: dict[int, _ListLevelDefinition] = {}
        for para_head_el in numbering_el.findall(f"{{{_NS_HH}}}paraHead"):
            level = _hwpx_list_level_from_para_head(para_head_el)
            levels[level] = _ListLevelDefinition(
                list_id=list_id,
                level=level,
                marker_type=_normalize_hwpx_marker_type(_xml_attr(para_head_el, "numFormat", "numformat")),
                marker_text=_hwpx_para_head_text(para_head_el),
                start=_safe_int(_xml_attr(para_head_el, "start")) or start,
            )
        if levels:
            definitions[native_id] = levels
    return definitions


def _hwpx_bullet_definitions(header_root: ET.Element | None) -> dict[str, dict[int, _ListLevelDefinition]]:
    if header_root is None:
        return {}

    definitions: dict[str, dict[int, _ListLevelDefinition]] = {}
    for bullet_el in header_root.findall(f".//{{{_NS_HH}}}bullet"):
        native_id = _xml_attr(bullet_el, "id")
        if not native_id:
            continue
        list_id = f"hwpx_bullet_{native_id}"
        bullet_char = _xml_attr(bullet_el, "char", "checkedChar") or "\u2022"
        levels: dict[int, _ListLevelDefinition] = {}
        para_head_els = bullet_el.findall(f"{{{_NS_HH}}}paraHead")
        if not para_head_els:
            levels[0] = _ListLevelDefinition(
                list_id=list_id,
                level=0,
                marker_type="bullet",
                marker_text=bullet_char,
                bullet_char=bullet_char,
            )
        for para_head_el in para_head_els:
            level = _hwpx_list_level_from_para_head(para_head_el)
            levels[level] = _ListLevelDefinition(
                list_id=list_id,
                level=level,
                marker_type="bullet",
                marker_text=_hwpx_para_head_text(para_head_el),
                bullet_char=bullet_char,
            )
        definitions[native_id] = levels
    return definitions


def _hwpx_resolve_list_info(
    para_pr_el: ET.Element,
    *,
    numbering_definitions: dict[str, dict[int, _ListLevelDefinition]],
    bullet_definitions: dict[str, dict[int, _ListLevelDefinition]],
    list_counter_state: _ListCounterState,
) -> ListItemInfo | None:
    heading_el = para_pr_el.find(f"{{{_NS_HH}}}heading")
    if heading_el is None:
        return None

    heading_type = (_xml_attr(heading_el, "type") or "NONE").upper()
    if heading_type not in {"NUMBER", "BULLET"}:
        return None

    native_id = _xml_attr(heading_el, "idRef", "idref")
    if not native_id:
        return None

    level = _safe_int(_xml_attr(heading_el, "level")) or 0
    level = max(level, 0)
    definitions = numbering_definitions.get(native_id) if heading_type == "NUMBER" else bullet_definitions.get(native_id)
    if not definitions:
        return None

    definition = definitions.get(level)
    if definition is None and level > 0:
        definition = definitions.get(level - 1)
    if definition is None:
        definition = definitions.get(0) or next(iter(definitions.values()))

    if definition.level != level:
        definition = _ListLevelDefinition(
            list_id=definition.list_id,
            level=level,
            marker_type=definition.marker_type,
            marker_text=definition.marker_text,
            start=definition.start,
            bullet_char=definition.bullet_char,
            left_indent_pt=definition.left_indent_pt,
            first_line_indent_pt=definition.first_line_indent_pt,
            hanging_indent_pt=definition.hanging_indent_pt,
        )

    return _advance_list_counter(
        list_counter_state,
        definition,
        level_definitions=definitions,
        source="hwpx",
    )


def _hwpx_para_style_from_pr(
    para_pr_el: ET.Element | None,
    *,
    numbering_definitions: dict[str, dict[int, _ListLevelDefinition]] | None = None,
    bullet_definitions: dict[str, dict[int, _ListLevelDefinition]] | None = None,
    list_counter_state: _ListCounterState | None = None,
) -> ParaStyleInfo | None:
    if para_pr_el is None:
        return None

    info = ParaStyleInfo()

    align_el = para_pr_el.find(f"{{{_NS_HH}}}align")
    if align_el is not None:
        info.align = _HWPX_HALIGN.get(align_el.get("horizontal", ""))

    margin_el = para_pr_el.find(f"{{{_NS_HH}}}margin")
    if margin_el is None:
        margin_el = para_pr_el.find(f".//{{{_NS_HH}}}margin")
    if margin_el is not None:
        first_line = _hwp_margin_value_to_pt(margin_el.find(f"{{{_NS_HC}}}intent"))
        info.first_line_indent_pt = first_line
        info.left_indent_pt = _hwp_margin_value_to_pt(margin_el.find(f"{{{_NS_HC}}}left"))
        info.right_indent_pt = _hwp_margin_value_to_pt(margin_el.find(f"{{{_NS_HC}}}right"))
        if first_line is not None and first_line < 0:
            info.hanging_indent_pt = abs(first_line)

    if list_counter_state is not None:
        info.list_info = _hwpx_resolve_list_info(
            para_pr_el,
            numbering_definitions=numbering_definitions or {},
            bullet_definitions=bullet_definitions or {},
            list_counter_state=list_counter_state,
        )

    return info if _has_para_style(info) else None


def _hwpx_run_style_from_char_pr(char_pr_el: ET.Element | None) -> RunStyleInfo:
    info = RunStyleInfo()
    if char_pr_el is None:
        return info

    info.bold = char_pr_el.find(f"{{{_NS_HH}}}bold") is not None
    info.italic = char_pr_el.find(f"{{{_NS_HH}}}italic") is not None

    underline_el = char_pr_el.find(f"{{{_NS_HH}}}underline")
    if underline_el is not None and underline_el.get("type", "NONE") != "NONE":
        info.underline = True

    strike_el = char_pr_el.find(f"{{{_NS_HH}}}strikeout")
    if strike_el is not None:
        strike_type = (strike_el.get("type") or "").upper()
        strike_shape = (strike_el.get("shape") or "").upper()
        if strike_type and strike_type != "NONE":
            info.strikethrough = True
        elif strike_shape in _HWPX_VISIBLE_LINE_SHAPES:
            info.strikethrough = True

    color = char_pr_el.get("textColor")
    if color and color != "#000000":
        info.color = color

    height = char_pr_el.get("height")
    if height is not None:
        try:
            info.size_pt = int(height) / 100.0
        except (TypeError, ValueError):
            pass

    return info


def _iter_section_paragraphs(section_root: ET.Element) -> list[ET.Element]:
    return section_root.findall(f"{_HP}p")


def _iter_paragraph_tables(paragraph_el: ET.Element) -> list[ET.Element]:
    return paragraph_el.findall(f"{_HP}run/{_HP}tbl")


def _iter_cell_paragraphs(cell_el: ET.Element) -> list[ET.Element]:
    direct = cell_el.findall(f"{_HP}subList/{_HP}p")
    if direct:
        return direct
    return cell_el.findall(f".//{_HP}p")


def _logical_table_cells(row_el: ET.Element) -> list[tuple[int, ET.Element]]:
    """Return logical 1-based column indices for row cells."""
    logical_cells: list[tuple[int, ET.Element]] = []
    fallback_col = 1

    for cell_el in row_el.findall(f"{_HP}tc"):
        cell_addr = cell_el.find(f"{_HP}cellAddr")
        col_addr = _safe_int(cell_addr.get("colAddr")) if cell_addr is not None else None
        logical_col = (col_addr + 1) if col_addr is not None else fallback_col
        logical_cells.append((logical_col, cell_el))

        cell_span = cell_el.find(f"{_HP}cellSpan")
        colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
        fallback_col = max(fallback_col, logical_col + max(colspan or 1, 1))

    return logical_cells


def _hwpx_table_dimensions(table_el: ET.Element) -> tuple[int, int]:
    row_els = table_el.findall(f"{_HP}tr")
    table_row_count = 0
    table_col_count = 0
    for row_el in row_els:
        for logical_col, cell_el in _logical_table_cells(row_el):
            cell_addr = cell_el.find(f"{_HP}cellAddr")
            row_addr = _safe_int(cell_addr.get("rowAddr")) if cell_addr is not None else None
            logical_row = (row_addr + 1) if row_addr is not None else 1
            cell_span = cell_el.find(f"{_HP}cellSpan")
            rowspan = _safe_int(cell_span.get("rowSpan")) if cell_span is not None else None
            colspan = _safe_int(cell_span.get("colSpan")) if cell_span is not None else None
            table_row_count = max(table_row_count, logical_row + max(rowspan or 1, 1) - 1)
            table_col_count = max(table_col_count, logical_col + max(colspan or 1, 1) - 1)
    return table_row_count or len(row_els), table_col_count


def _hwpx_table_size(table_el: ET.Element) -> tuple[float | None, float | None]:
    size_el = table_el.find(f"{_HP}sz")
    if size_el is None:
        return None, None
    return (
        _hwp_numeric_to_pt(size_el.get("width")),
        _hwp_numeric_to_pt(size_el.get("height")),
    )


def _extract_hwpx_table_styles(
    style_map: StyleMap,
    table_el: ET.Element,
    table_id: str,
    *,
    para_pr_map: dict[str, ET.Element],
    char_pr_map: dict[str, ET.Element],
    border_fill_map: dict[str, ET.Element],
    numbering_definitions: dict[str, dict[int, _ListLevelDefinition]],
    bullet_definitions: dict[str, dict[int, _ListLevelDefinition]],
    list_counter_state: _ListCounterState,
) -> None:
    row_count, col_count = _hwpx_table_dimensions(table_el)
    width_pt, height_pt = _hwpx_table_size(table_el)
    table_cell_padding_defaults = _hwpx_table_cell_padding_defaults(table_el)
    style_map.tables[table_id] = TableStyleInfo(
        row_count=row_count,
        col_count=col_count,
        width_pt=width_pt,
        height_pt=height_pt,
    )

    for tr_idx, row_el in enumerate(table_el.findall(f"{_HP}tr"), start=1):
        for tc_idx, cell_el in _logical_table_cells(row_el):
            cell_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}"
            style_map.cells[cell_id] = _hwpx_cell_style(
                cell_el,
                para_pr_map=para_pr_map,
                border_fill_map=border_fill_map,
                table_cell_padding_defaults=table_cell_padding_defaults,
            )

            cell_paragraphs = _iter_cell_paragraphs(cell_el)
            if not cell_paragraphs:
                style_map.runs[f"{cell_id}.p1.r1"] = RunStyleInfo()
                continue

            for cp_idx, cell_para_el in enumerate(cell_paragraphs, start=1):
                cell_paragraph_id = f"{cell_id}.p{cp_idx}"
                cell_para_pr_ref = cell_para_el.get("paraPrIDRef")
                if cell_para_pr_ref and cell_para_pr_ref in para_pr_map:
                    cp_style = _hwpx_para_style_from_pr(
                        para_pr_map[cell_para_pr_ref],
                        numbering_definitions=numbering_definitions,
                        bullet_definitions=bullet_definitions,
                        list_counter_state=list_counter_state,
                    )
                    if cp_style is not None:
                        style_map.paragraphs[cell_paragraph_id] = cp_style

                cell_run_els = cell_para_el.findall(f"{_HP}run")
                if not cell_run_els:
                    style_map.runs[f"{cell_paragraph_id}.r1"] = RunStyleInfo()
                else:
                    for cr_idx, cell_run_el in enumerate(cell_run_els, start=1):
                        char_pr_ref = cell_run_el.get("charPrIDRef")
                        char_pr_el = char_pr_map.get(char_pr_ref) if char_pr_ref else None
                        style_map.runs[f"{cell_paragraph_id}.r{cr_idx}"] = (
                            _hwpx_run_style_from_char_pr(char_pr_el)
                        )

                for nested_t_idx, nested_table_el in enumerate(_iter_paragraph_tables(cell_para_el), start=1):
                    nested_table_id = f"{cell_paragraph_id}.tbl{nested_t_idx}"
                    _extract_hwpx_table_styles(
                        style_map,
                        nested_table_el,
                        nested_table_id,
                        para_pr_map=para_pr_map,
                        char_pr_map=char_pr_map,
                        border_fill_map=border_fill_map,
                        numbering_definitions=numbering_definitions,
                        bullet_definitions=bullet_definitions,
                        list_counter_state=list_counter_state,
                    )


def _section_roots_from_bytes(source: bytes) -> list[ET.Element]:
    section_name_pattern = re.compile(r"^Contents/section\d+\.xml$")

    with zipfile.ZipFile(BytesIO(source)) as zf:
        def _section_order(name: str) -> int:
            match = re.search(r"section(\d+)\.xml$", name)
            return int(match.group(1)) if match else -1

        names = sorted(
            (name for name in zf.namelist() if section_name_pattern.match(name)),
            key=_section_order,
        )
        return [_parse_hwpx_xml_part(zf.read(name)) for name in names]


def _parse_hwpx_xml_part(data: bytes) -> ET.Element:
    try:
        return ET.fromstring(data)
    except ET.ParseError as parse_error:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            raise parse_error from None
        cleaned = _repair_hwpx_xml_text(text)
        if cleaned == text:
            raise
        return ET.fromstring(cleaned.encode("utf-8"))


def _repair_hwpx_xml_text(text: str) -> str:
    cleaned = _XML_1_0_INVALID_CHAR_RE.sub("", text)
    cleaned = _BARE_XML_AMPERSAND_RE.sub("&amp;", cleaned)
    return _escape_attribute_angle_brackets(cleaned)


def _escape_attribute_angle_brackets(text: str) -> str:
    out: list[str] = []
    in_tag = False
    quote: str | None = None

    for ch in text:
        if not in_tag:
            if ch == "<":
                in_tag = True
            out.append(ch)
            continue

        if quote is not None:
            if ch == quote:
                quote = None
                out.append(ch)
            elif ch == "<":
                out.append("&lt;")
            elif ch == ">":
                out.append("&gt;")
            else:
                out.append(ch)
            continue

        if ch in ("'", '"'):
            quote = ch
        elif ch == ">":
            in_tag = False
        out.append(ch)

    return "".join(out)


def _header_root_from_bytes(source: bytes) -> ET.Element | None:
    with zipfile.ZipFile(BytesIO(source)) as zf:
        try:
            return _parse_hwpx_xml_part(zf.read("Contents/header.xml"))
        except KeyError:
            return None


def _hwpx_cell_style(
    cell_el: ET.Element,
    *,
    para_pr_map: dict[str, ET.Element],
    border_fill_map: dict[str, ET.Element],
    table_cell_padding_defaults: dict[str, float] | None = None,
) -> CellStyleInfo:
    info = CellStyleInfo(vertical_align="center")

    span_el = cell_el.find(f"{_HP}cellSpan")
    if span_el is not None:
        info.rowspan = _safe_int(span_el.get("rowSpan")) or 1
        info.colspan = _safe_int(span_el.get("colSpan")) or 1

    sub_list = cell_el.find(f"{_HP}subList")
    if sub_list is not None:
        valign = sub_list.get("vertAlign", "")
        info.vertical_align = _HWPX_VALIGN.get(valign) or info.vertical_align

    bf_ref = cell_el.get("borderFillIDRef")
    if bf_ref and bf_ref in border_fill_map:
        border_fill = border_fill_map[bf_ref]
        info.border_top = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}topBorder"))
        info.border_bottom = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}bottomBorder"))
        info.border_left = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}leftBorder"))
        info.border_right = _hwpx_border_css(border_fill.find(f"{{{_NS_HH}}}rightBorder"))
        info.diagonal_tr_bl = _hwpx_diagonal_border_css(border_fill, direction="slash")
        info.diagonal_tl_br = _hwpx_diagonal_border_css(border_fill, direction="backslash")

        fill_brush = border_fill.find(f"{{{_NS_HH}}}fillBrush")
        if fill_brush is None:
            fill_brush = border_fill.find(f"{{{_NS_HC}}}fillBrush")
        if fill_brush is not None:
            face_color = fill_brush.get("faceColor")
            if not face_color:
                win_brush = fill_brush.find(f"{{{_NS_HH}}}winBrush")
                if win_brush is None:
                    win_brush = fill_brush.find(f"{{{_NS_HC}}}winBrush")
                if win_brush is not None:
                    face_color = win_brush.get("faceColor")
            if face_color and face_color.lower() not in ("none", "#ffffff", "transparent"):
                info.background = face_color

    cell_size = cell_el.find(f"{_HP}cellSz")
    if cell_size is not None:
        info.width_pt = _hwp_numeric_to_pt(cell_size.get("width"))
        info.height_pt = _hwp_numeric_to_pt(cell_size.get("height"))

    _apply_hwpx_cell_margin(
        info,
        cell_el.find(f"{_HP}cellMargin"),
        defaults=table_cell_padding_defaults,
    )

    return info


def _extract_styles_hwpx_from_roots(
    section_roots: list[ET.Element],
    *,
    header_root: ET.Element | None,
) -> StyleMap:
    style_map = StyleMap()

    para_pr_map = _map_by_id(header_root, "paraPr")
    char_pr_map = _map_by_id(header_root, "charPr")
    border_fill_map = _map_by_id(header_root, "borderFill")
    numbering_definitions = _hwpx_numbering_definitions(header_root)
    bullet_definitions = _hwpx_bullet_definitions(header_root)
    list_counter_state = _ListCounterState.create()

    for s_idx, section_root in enumerate(section_roots, start=1):
        for p_idx, para_el in enumerate(_iter_section_paragraphs(section_root), start=1):
            paragraph_id = f"s{s_idx}.p{p_idx}"

            para_pr_ref = para_el.get("paraPrIDRef")
            if para_pr_ref and para_pr_ref in para_pr_map:
                para_style = _hwpx_para_style_from_pr(
                    para_pr_map[para_pr_ref],
                    numbering_definitions=numbering_definitions,
                    bullet_definitions=bullet_definitions,
                    list_counter_state=list_counter_state,
                )
                if para_style is not None:
                    style_map.paragraphs[paragraph_id] = para_style

            run_els = para_el.findall(f"{_HP}run")
            if not run_els:
                style_map.runs[f"{paragraph_id}.r1"] = RunStyleInfo()
            else:
                for r_idx, run_el in enumerate(run_els, start=1):
                    char_pr_ref = run_el.get("charPrIDRef")
                    char_pr_el = char_pr_map.get(char_pr_ref) if char_pr_ref else None
                    style_map.runs[f"{paragraph_id}.r{r_idx}"] = _hwpx_run_style_from_char_pr(
                        char_pr_el
                    )

            for t_idx, table_el in enumerate(_iter_paragraph_tables(para_el), start=1):
                table_id = f"{paragraph_id}.r1.tbl{t_idx}"
                _extract_hwpx_table_styles(
                    style_map,
                    table_el,
                    table_id,
                    para_pr_map=para_pr_map,
                    char_pr_map=char_pr_map,
                    border_fill_map=border_fill_map,
                    numbering_definitions=numbering_definitions,
                    bullet_definitions=bullet_definitions,
                    list_counter_state=list_counter_state,
                )

    return style_map


def extract_styles_hwpx(source: "HwpxDocument | str | Path | bytes") -> StyleMap:
    """Extract style map from HWPX source."""
    from ..hwpx import HwpxDocument

    if isinstance(source, bytes):
        return _extract_styles_hwpx_from_roots(
            _section_roots_from_bytes(source),
            header_root=_header_root_from_bytes(source),
        )

    if isinstance(source, (str, Path)):
        return extract_styles_hwpx(Path(source).read_bytes())

    if isinstance(source, HwpxDocument):
        section_roots = [section.element for section in source.sections]
        header_root = source.headers[0].element if source.headers else None
        return _extract_styles_hwpx_from_roots(section_roots, header_root=header_root)

    raise TypeError(
        "source must be HwpxDocument, bytes, or a .hwpx path, "
        f"got {type(source)!r}"
    )


def _docx_run_style(run) -> RunStyleInfo:
    info = RunStyleInfo(
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
    )

    font = run.font

    if font.color and font.color.rgb:
        rgb = str(font.color.rgb)
        if rgb != "000000":
            info.color = f"#{rgb}"

    if font.size is not None:
        info.size_pt = _length_to_pt(font.size)

    if font.strike:
        info.strikethrough = True
    if font.superscript:
        info.superscript = True
    if font.subscript:
        info.subscript = True
    if font.highlight_color is not None:
        info.highlight = str(font.highlight_color)

    return info


def _docx_numbering_root(doc):
    try:
        numbering_part = getattr(doc.part, "numbering_part", None)
    except (AttributeError, KeyError, NotImplementedError, ValueError):
        return None
    return getattr(numbering_part, "element", None)


def _docx_val(el, name: str = "w:val") -> str | None:
    if el is None:
        return None
    from docx.oxml.ns import qn

    return el.get(qn(name))


def _docx_num_pr_values(num_pr) -> tuple[str | None, int | None]:
    if num_pr is None:
        return None, None
    from docx.oxml.ns import qn

    ilvl_el = num_pr.find(qn("w:ilvl"))
    num_id_el = num_pr.find(qn("w:numId"))
    return _docx_val(num_id_el), _safe_int(_docx_val(ilvl_el))


def _merge_docx_numbering_values(
    primary: tuple[str | None, int | None],
    fallback: tuple[str | None, int | None],
) -> tuple[str | None, int | None]:
    return (
        primary[0] if primary[0] is not None else fallback[0],
        primary[1] if primary[1] is not None else fallback[1],
    )


def _docx_style_numbering_values(
    style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, tuple[str | None, int | None]],
) -> tuple[str | None, int | None]:
    from docx.oxml.ns import qn

    if not style_id:
        return None, None
    cached = cache.get(style_id)
    if cached is not None:
        return cached

    style_el = style_elements.get(style_id)
    if style_el is None:
        cache[style_id] = (None, None)
        return None, None

    based_on_el = style_el.find(qn("w:basedOn"))
    base_style_id = _docx_val(based_on_el)
    base_values = _docx_style_numbering_values(
        base_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    p_pr = style_el.find(qn("w:pPr"))
    num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
    values = _merge_docx_numbering_values(_docx_num_pr_values(num_pr), base_values)
    cache[style_id] = values
    return values


def _docx_paragraph_numbering_values(
    paragraph,
    *,
    style_elements: dict[str, object],
    style_numbering_cache: dict[str, tuple[str | None, int | None]],
) -> tuple[str | None, int | None]:
    from docx.oxml.ns import qn

    style_id = paragraph.style.style_id if paragraph.style is not None else None
    style_values = _docx_style_numbering_values(
        style_id,
        style_elements=style_elements,
        cache=style_numbering_cache,
    )
    p_pr = paragraph._p.find(qn("w:pPr"))
    num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
    return _merge_docx_numbering_values(_docx_num_pr_values(num_pr), style_values)


def _docx_indentation_from_p_pr(p_pr) -> tuple[float | None, float | None, float | None]:
    from docx.oxml.ns import qn

    if p_pr is None:
        return None, None, None
    ind_el = p_pr.find(qn("w:ind"))
    if ind_el is None:
        return None, None, None

    left_indent = _docx_measure_to_pt(
        ind_el.get(qn("w:left")) or ind_el.get(qn("w:start")),
        "dxa",
    )
    first_line_indent = _docx_measure_to_pt(ind_el.get(qn("w:firstLine")), "dxa")
    hanging_indent = _docx_measure_to_pt(ind_el.get(qn("w:hanging")), "dxa")
    if hanging_indent is not None:
        first_line_indent = -hanging_indent
    return left_indent, first_line_indent, hanging_indent


def _docx_level_definition(
    lvl_el,
    *,
    list_id: str,
    level: int,
    default_start: int = 1,
) -> _ListLevelDefinition:
    from docx.oxml.ns import qn

    start_el = lvl_el.find(qn("w:start")) if lvl_el is not None else None
    num_fmt_el = lvl_el.find(qn("w:numFmt")) if lvl_el is not None else None
    lvl_text_el = lvl_el.find(qn("w:lvlText")) if lvl_el is not None else None
    p_pr = lvl_el.find(qn("w:pPr")) if lvl_el is not None else None
    left_indent, first_line_indent, hanging_indent = _docx_indentation_from_p_pr(p_pr)
    marker_type = _docx_val(num_fmt_el) or "decimal"
    marker_text = _docx_val(lvl_text_el)
    return _ListLevelDefinition(
        list_id=list_id,
        level=level,
        marker_type=marker_type,
        marker_text=marker_text,
        start=_safe_int(_docx_val(start_el)) or default_start,
        bullet_char=marker_text if marker_type == "bullet" else None,
        left_indent_pt=left_indent,
        first_line_indent_pt=first_line_indent,
        hanging_indent_pt=hanging_indent,
    )


def _docx_numbering_definitions(doc) -> dict[str, dict[int, _ListLevelDefinition]]:
    from docx.oxml.ns import qn

    numbering_root = _docx_numbering_root(doc)
    if numbering_root is None:
        return {}

    abstract_levels: dict[str, dict[int, _ListLevelDefinition]] = {}
    for abstract_el in numbering_root.findall(qn("w:abstractNum")):
        abstract_id = abstract_el.get(qn("w:abstractNumId"))
        if abstract_id is None:
            continue
        levels: dict[int, _ListLevelDefinition] = {}
        for lvl_el in abstract_el.findall(qn("w:lvl")):
            level = _safe_int(lvl_el.get(qn("w:ilvl"))) or 0
            levels[level] = _docx_level_definition(
                lvl_el,
                list_id=f"docx_abstract_{abstract_id}",
                level=level,
            )
        abstract_levels[abstract_id] = levels

    definitions: dict[str, dict[int, _ListLevelDefinition]] = {}
    for num_el in numbering_root.findall(qn("w:num")):
        num_id = num_el.get(qn("w:numId"))
        if num_id is None:
            continue
        abstract_id_el = num_el.find(qn("w:abstractNumId"))
        abstract_id = _docx_val(abstract_id_el)
        inherited = abstract_levels.get(abstract_id or "", {})
        list_id = f"docx_num_{num_id}"
        levels = {
            level: _ListLevelDefinition(
                list_id=list_id,
                level=definition.level,
                marker_type=definition.marker_type,
                marker_text=definition.marker_text,
                start=definition.start,
                bullet_char=definition.bullet_char,
                left_indent_pt=definition.left_indent_pt,
                first_line_indent_pt=definition.first_line_indent_pt,
                hanging_indent_pt=definition.hanging_indent_pt,
            )
            for level, definition in inherited.items()
        }

        for override_el in num_el.findall(qn("w:lvlOverride")):
            level = _safe_int(override_el.get(qn("w:ilvl"))) or 0
            start_override_el = override_el.find(qn("w:startOverride"))
            default_start = _safe_int(_docx_val(start_override_el)) or levels.get(
                level,
                _ListLevelDefinition(list_id, level),
            ).start
            lvl_el = override_el.find(qn("w:lvl"))
            if lvl_el is not None:
                levels[level] = _docx_level_definition(
                    lvl_el,
                    list_id=list_id,
                    level=level,
                    default_start=default_start,
                )
            elif level in levels:
                previous = levels[level]
                levels[level] = _ListLevelDefinition(
                    list_id=list_id,
                    level=previous.level,
                    marker_type=previous.marker_type,
                    marker_text=previous.marker_text,
                    start=default_start,
                    bullet_char=previous.bullet_char,
                    left_indent_pt=previous.left_indent_pt,
                    first_line_indent_pt=previous.first_line_indent_pt,
                    hanging_indent_pt=previous.hanging_indent_pt,
                )
        if levels:
            definitions[num_id] = levels
    return definitions


def _docx_resolve_list_info(
    paragraph,
    *,
    numbering_definitions: dict[str, dict[int, _ListLevelDefinition]],
    style_elements: dict[str, object],
    style_numbering_cache: dict[str, tuple[str | None, int | None]],
    list_counter_state: _ListCounterState,
) -> tuple[ListItemInfo | None, _ListLevelDefinition | None]:
    num_id, level = _docx_paragraph_numbering_values(
        paragraph,
        style_elements=style_elements,
        style_numbering_cache=style_numbering_cache,
    )
    if not num_id or num_id == "0":
        return None, None

    level = max(level or 0, 0)
    definitions = numbering_definitions.get(num_id)
    if not definitions:
        return None, None
    definition = definitions.get(level) or definitions.get(0) or next(iter(definitions.values()))
    if definition.level != level:
        definition = _ListLevelDefinition(
            list_id=definition.list_id,
            level=level,
            marker_type=definition.marker_type,
            marker_text=definition.marker_text,
            start=definition.start,
            bullet_char=definition.bullet_char,
            left_indent_pt=definition.left_indent_pt,
            first_line_indent_pt=definition.first_line_indent_pt,
            hanging_indent_pt=definition.hanging_indent_pt,
        )
    list_info = _advance_list_counter(
        list_counter_state,
        definition,
        level_definitions=definitions,
        source="docx",
    )
    return list_info, definition


def _docx_para_style(
    paragraph,
    *,
    numbering_definitions: dict[str, dict[int, _ListLevelDefinition]] | None = None,
    style_elements: dict[str, object] | None = None,
    style_numbering_cache: dict[str, tuple[str | None, int | None]] | None = None,
    list_counter_state: _ListCounterState | None = None,
) -> ParaStyleInfo | None:
    info = ParaStyleInfo()

    if paragraph.alignment is not None:
        info.align = _DOCX_ALIGN.get(paragraph.alignment)

    pf = paragraph.paragraph_format
    if pf is not None:
        info.left_indent_pt = _length_to_pt(pf.left_indent)
        info.right_indent_pt = _length_to_pt(pf.right_indent)

        first_line = _length_to_pt(pf.first_line_indent)
        info.first_line_indent_pt = first_line
        if first_line is not None and first_line < 0:
            info.hanging_indent_pt = abs(first_line)

    list_definition = None
    if list_counter_state is not None:
        info.list_info, list_definition = _docx_resolve_list_info(
            paragraph,
            numbering_definitions=numbering_definitions or {},
            style_elements={} if style_elements is None else style_elements,
            style_numbering_cache={} if style_numbering_cache is None else style_numbering_cache,
            list_counter_state=list_counter_state,
        )
    if list_definition is not None:
        if info.left_indent_pt is None:
            info.left_indent_pt = list_definition.left_indent_pt
        if info.first_line_indent_pt is None:
            info.first_line_indent_pt = list_definition.first_line_indent_pt
        if info.hanging_indent_pt is None:
            info.hanging_indent_pt = list_definition.hanging_indent_pt

    return info if _has_para_style(info) else None


def _docx_border_css(tc_borders, side: str) -> str | None:
    if tc_borders is None:
        return None

    from docx.oxml.ns import qn

    el = tc_borders.find(qn(f"w:{side}"))
    if el is None:
        return None
    val = el.get(qn("w:val"), "none")
    if val in ("none", "nil"):
        return None
    sz = el.get(qn("w:sz"), "4")
    try:
        px = max(1, round(int(sz) / 8 * 1.333))
    except (TypeError, ValueError):
        px = 1
    color = el.get(qn("w:color"), "000000")
    if color.lower() == "auto":
        color = "000000"
    style_map = {"single": "solid", "double": "double", "dashed": "dashed", "dotted": "dotted"}
    css_style = style_map.get(val, "solid")
    return f"{px}px {css_style} #{color}"


def _docx_table_style_border_defaults(
    style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, dict[str, str | None]],
) -> dict[str, str | None]:
    from docx.oxml.ns import qn

    if not style_id:
        return {}
    cached = cache.get(style_id)
    if cached is not None:
        return dict(cached)

    style_el = style_elements.get(style_id)
    if style_el is None:
        cache[style_id] = {}
        return {}

    based_on_el = style_el.find(qn("w:basedOn"))
    base_style_id = based_on_el.get(qn("w:val")) if based_on_el is not None else None
    merged = _docx_table_style_border_defaults(
        base_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    tbl_pr = style_el.find(qn("w:tblPr"))
    tbl_borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    if tbl_borders is not None:
        for border_name, border_key in (
            ("top", "top"),
            ("bottom", "bottom"),
            ("left", "left"),
            ("right", "right"),
            ("insideH", "inside_h"),
            ("insideV", "inside_v"),
        ):
            border_css = _docx_border_css(tbl_borders, border_name)
            if border_css is not None:
                merged[border_key] = border_css

    cache[style_id] = dict(merged)
    return merged


def _docx_table_style_cell_padding_defaults(
    style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, dict[str, float]],
) -> dict[str, float]:
    from docx.oxml.ns import qn

    if not style_id:
        return {}
    cached = cache.get(style_id)
    if cached is not None:
        return dict(cached)

    style_el = style_elements.get(style_id)
    if style_el is None:
        cache[style_id] = {}
        return {}

    based_on_el = style_el.find(qn("w:basedOn"))
    base_style_id = based_on_el.get(qn("w:val")) if based_on_el is not None else None
    merged = _docx_table_style_cell_padding_defaults(
        base_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    tbl_pr = style_el.find(qn("w:tblPr"))
    cell_margin = tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None
    merged.update(_docx_cell_margin_to_padding(cell_margin))

    cache[style_id] = dict(merged)
    return merged


def _docx_table_cell_padding_defaults(
    table,
    table_style_id: str | None,
    *,
    style_elements: dict[str, object],
    cache: dict[str, dict[str, float]],
) -> dict[str, float]:
    from docx.oxml.ns import qn

    defaults = _docx_table_style_cell_padding_defaults(
        table_style_id,
        style_elements=style_elements,
        cache=cache,
    )

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    cell_margin = tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None
    defaults.update(_docx_cell_margin_to_padding(cell_margin))
    return defaults


def _docx_default_cell_border(
    side: str,
    *,
    row_index: int,
    col_index: int,
    row_count: int,
    col_count: int,
    table_border_defaults: dict[str, str | None],
) -> str | None:
    if side == "top":
        return table_border_defaults.get("top") if row_index == 1 else table_border_defaults.get("inside_h")
    if side == "bottom":
        return table_border_defaults.get("bottom") if row_index == row_count else table_border_defaults.get("inside_h")
    if side == "left":
        return table_border_defaults.get("left") if col_index == 1 else table_border_defaults.get("inside_v")
    if side == "right":
        return table_border_defaults.get("right") if col_index == col_count else table_border_defaults.get("inside_v")
    return None


def _docx_table_direct_border_defaults(table) -> dict[str, str | None]:
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    tbl_borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    if tbl_borders is None:
        return {}

    defaults: dict[str, str | None] = {}
    for border_name, border_key in (
        ("top", "top"),
        ("bottom", "bottom"),
        ("left", "left"),
        ("right", "right"),
        ("insideH", "inside_h"),
        ("insideV", "inside_v"),
    ):
        border_css = _docx_border_css(tbl_borders, border_name)
        if border_css is not None:
            defaults[border_key] = border_css
    return defaults


def _docx_table_size(table) -> tuple[float | None, float | None]:
    from docx.oxml.ns import qn

    table_width_pt: float | None = None
    table_height_pt: float | None = None

    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is not None:
            table_width_pt = _docx_measure_to_pt(
                tbl_w.get(qn("w:w")),
                tbl_w.get(qn("w:type")),
            )

    if table_width_pt is None:
        tbl_grid = table._tbl.find(qn("w:tblGrid"))
        if tbl_grid is not None:
            grid_width_pt = 0.0
            has_grid = False
            for grid_col in tbl_grid.findall(qn("w:gridCol")):
                width_pt = _docx_measure_to_pt(grid_col.get(qn("w:w")), "dxa")
                if width_pt is not None:
                    has_grid = True
                    grid_width_pt += width_pt
            if has_grid:
                table_width_pt = grid_width_pt

    row_height_total = 0.0
    has_row_height = False
    for row in table.rows:
        tr_pr = row._tr.find(qn("w:trPr"))
        tr_height = tr_pr.find(qn("w:trHeight")) if tr_pr is not None else None
        if tr_height is None:
            continue
        height_pt = _docx_measure_to_pt(tr_height.get(qn("w:val")), "dxa")
        if height_pt is None:
            continue
        has_row_height = True
        row_height_total += height_pt
    if has_row_height:
        table_height_pt = row_height_total

    return table_width_pt, table_height_pt


def _docx_cell_style(
    cell,
    *,
    row_index: int,
    col_index: int,
    row_count: int,
    col_count: int,
    row_height_pt: float | None = None,
    table_border_defaults: dict[str, str | None] | None = None,
    table_cell_padding_defaults: dict[str, float] | None = None,
) -> CellStyleInfo:
    from docx.oxml.ns import qn

    info = CellStyleInfo(vertical_align="center")
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return info

    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is not None:
        try:
            info.colspan = int(grid_span.get(qn("w:val"), "1"))
        except (TypeError, ValueError):
            pass

    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is not None:
        val = v_align.get(qn("w:val"), "")
        info.vertical_align = {"top": "top", "center": "center", "bottom": "bottom"}.get(val) or info.vertical_align

    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is not None:
        info.width_pt = _docx_measure_to_pt(
            tc_width.get(qn("w:w")),
            tc_width.get(qn("w:type")),
        )
    if row_height_pt is not None:
        info.height_pt = row_height_pt

    table_cell_padding_defaults = table_cell_padding_defaults or {}
    if table_cell_padding_defaults:
        _apply_cell_padding(info, table_cell_padding_defaults)
    cell_margin = tc_pr.find(qn("w:tcMar"))
    if cell_margin is not None:
        _apply_cell_padding(info, _docx_cell_margin_to_padding(cell_margin))

    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill.lower() not in ("auto", "ffffff", "none"):
            info.background = f"#{fill}"

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    table_border_defaults = table_border_defaults or {}
    info.border_top = _docx_border_css(tc_borders, "top") or _docx_default_cell_border(
        "top",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.border_bottom = _docx_border_css(tc_borders, "bottom") or _docx_default_cell_border(
        "bottom",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.border_left = _docx_border_css(tc_borders, "left") or _docx_default_cell_border(
        "left",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.border_right = _docx_border_css(tc_borders, "right") or _docx_default_cell_border(
        "right",
        row_index=row_index,
        col_index=col_index,
        row_count=row_count,
        col_count=col_count,
        table_border_defaults=table_border_defaults,
    )
    info.diagonal_tl_br = _docx_border_css(tc_borders, "tl2br")
    info.diagonal_tr_bl = _docx_border_css(tc_borders, "tr2bl")
    return info


def extract_styles_docx(
    source: "DocxDocument | str | Path | bytes",
    *,
    include_tables: bool = True,
) -> StyleMap:
    """Extract style map from DOCX source."""
    from docx import Document as load_docx
    from docx.document import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    from .docx_structured_exporter import _iter_blocks, _iter_blocks_from_element

    if isinstance(source, DocxDocument):
        doc = source
    elif isinstance(source, bytes):
        doc = load_docx(BytesIO(source))
    else:
        doc = load_docx(str(source))

    style_map = StyleMap()
    style_elements: dict[str, object] = {}
    table_style_elements: dict[str, object] = {}
    border_defaults_cache: dict[str, dict[str, str | None]] = {}
    cell_padding_defaults_cache: dict[str, dict[str, float]] = {}
    style_numbering_cache: dict[str, tuple[str | None, int | None]] = {}
    numbering_definitions = _docx_numbering_definitions(doc)
    list_counter_state = _ListCounterState.create()

    for style_el in doc.styles.element.findall(qn("w:style")):
        style_id = style_el.get(qn("w:styleId"))
        if style_id:
            style_elements[style_id] = style_el
            if style_el.get(qn("w:type")) == "table":
                table_style_elements[style_id] = style_el

    p_idx = 0
    tbl_counter = 0

    def _extract_docx_table_styles(table, table_id: str) -> None:
        vmerge_starts: dict[tuple[str, int], str] = {}
        table_style_id = table.style.style_id if table.style is not None else None
        table_border_defaults = _docx_table_style_border_defaults(
            table_style_id,
            style_elements=table_style_elements,
            cache=border_defaults_cache,
        )
        table_border_defaults.update(_docx_table_direct_border_defaults(table))
        table_cell_padding_defaults = _docx_table_cell_padding_defaults(
            table,
            table_style_id,
            style_elements=table_style_elements,
            cache=cell_padding_defaults_cache,
        )
        table_width_pt, table_height_pt = _docx_table_size(table)

        style_map.tables[table_id] = TableStyleInfo(
            row_count=len(table.rows),
            col_count=len(table.columns),
            width_pt=table_width_pt,
            height_pt=table_height_pt,
        )

        for tr_idx, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.find(qn("w:trPr"))
            tr_height = tr_pr.find(qn("w:trHeight")) if tr_pr is not None else None
            row_height_pt = None
            if tr_height is not None:
                row_height_pt = _docx_measure_to_pt(
                    tr_height.get(qn("w:val")),
                    "dxa",
                )
            for tc_idx, cell in enumerate(row.cells, start=1):
                cell_id = f"{table_id}.tr{tr_idx}.tc{tc_idx}"

                tc_pr = cell._tc.find(qn("w:tcPr"))
                if tc_pr is not None:
                    v_merge = tc_pr.find(qn("w:vMerge"))
                    if v_merge is not None:
                        val = v_merge.get(qn("w:val"), "")
                        col_key = (table_id, tc_idx)
                        if val == "restart":
                            vmerge_starts[col_key] = cell_id
                        elif col_key in vmerge_starts:
                            start_cell_id = vmerge_starts[col_key]
                            start_style = style_map.cells.get(start_cell_id)
                            if start_style is not None:
                                start_style.rowspan += 1
                            continue
                    else:
                        vmerge_starts.pop((table_id, tc_idx), None)

                style_map.cells[cell_id] = _docx_cell_style(
                    cell,
                    row_index=tr_idx,
                    col_index=tc_idx,
                    row_count=len(table.rows),
                    col_count=len(table.columns),
                    row_height_pt=row_height_pt,
                    table_border_defaults=table_border_defaults,
                    table_cell_padding_defaults=table_cell_padding_defaults,
                )

                cp_idx = 0
                current_paragraph_id: str | None = None
                nested_table_counter_by_paragraph: dict[str, int] = {}

                for block in _iter_blocks_from_element(
                    cell,
                    cell._tc,
                    CT_P=CT_P,
                    CT_Tbl=CT_Tbl,
                    Paragraph=Paragraph,
                    Table=Table,
                ):
                    if isinstance(block, Paragraph):
                        cp_idx += 1
                        current_paragraph_id = f"{cell_id}.p{cp_idx}"
                        cp_style = _docx_para_style(
                            block,
                            numbering_definitions=numbering_definitions,
                            style_elements=style_elements,
                            style_numbering_cache=style_numbering_cache,
                            list_counter_state=list_counter_state,
                        )
                        if cp_style is not None:
                            style_map.paragraphs[current_paragraph_id] = cp_style

                        if not block.runs:
                            style_map.runs[f"{current_paragraph_id}.r1"] = RunStyleInfo()
                            continue

                        for cr_idx, cell_run in enumerate(block.runs, start=1):
                            style_map.runs[f"{current_paragraph_id}.r{cr_idx}"] = _docx_run_style(
                                cell_run
                            )
                        continue

                    if current_paragraph_id is None:
                        cp_idx += 1
                        current_paragraph_id = f"{cell_id}.p{cp_idx}"

                    nested_tbl_counter = nested_table_counter_by_paragraph.get(current_paragraph_id, 0) + 1
                    nested_table_counter_by_paragraph[current_paragraph_id] = nested_tbl_counter
                    nested_table_id = f"{current_paragraph_id}.tbl{nested_tbl_counter}"
                    _extract_docx_table_styles(block, nested_table_id)

    for block in _iter_blocks(
        doc,
        CT_P=CT_P,
        CT_Tbl=CT_Tbl,
        Paragraph=Paragraph,
        Table=Table,
    ):
        if isinstance(block, Paragraph):
            p_idx += 1
            paragraph_id = f"s1.p{p_idx}"

            pstyle = _docx_para_style(
                block,
                numbering_definitions=numbering_definitions,
                style_elements=style_elements,
                style_numbering_cache=style_numbering_cache,
                list_counter_state=list_counter_state,
            )
            if pstyle is not None:
                style_map.paragraphs[paragraph_id] = pstyle

            if not block.runs:
                style_map.runs[f"{paragraph_id}.r1"] = RunStyleInfo()
            else:
                for r_idx, run in enumerate(block.runs, start=1):
                    style_map.runs[f"{paragraph_id}.r{r_idx}"] = _docx_run_style(run)
            continue

        if not include_tables or not isinstance(block, Table):
            continue

        tbl_counter += 1
        p_idx += 1
        table_id = f"s1.p{p_idx}.r1.tbl{tbl_counter}"
        _extract_docx_table_styles(block, table_id)

    return style_map


def _collect_style_map_from_doc_ir(doc_ir) -> StyleMap:
    from ..models import _node_debug_path

    style_map = StyleMap()

    def collect_paragraph(paragraph) -> None:
        if paragraph.para_style is not None:
            style_map.paragraphs[_node_debug_path(paragraph)] = paragraph.para_style.model_copy(deep=True)
        for run in paragraph.runs:
            if run.run_style is not None:
                style_map.runs[_node_debug_path(run)] = run.run_style.model_copy(deep=True)
        for table in paragraph.tables:
            collect_table(table)

    def collect_table(table) -> None:
        if table.table_style is not None:
            style_map.tables[_node_debug_path(table)] = table.table_style.model_copy(deep=True)
        for cell in table.iter_cells():
            if cell.cell_style is not None:
                style_map.cells[_node_debug_path(cell)] = cell.cell_style.model_copy(deep=True)
            for paragraph in cell.paragraphs:
                collect_paragraph(paragraph)

    for paragraph in doc_ir.paragraphs:
        collect_paragraph(paragraph)

    return style_map


def extract_styles_pdf(source: str | Path | bytes) -> StyleMap:
    from ..pdf.pipeline import parse_pdf_to_doc_ir

    if isinstance(source, (str, Path)):
        doc_ir = parse_pdf_to_doc_ir(source)
        return _collect_style_map_from_doc_ir(doc_ir)

    with TemporarySourcePath(source, suffix=".pdf") as source_path:
        doc_ir = parse_pdf_to_doc_ir(source_path)
        return _collect_style_map_from_doc_ir(doc_ir)


def extract_styles(
    source: "HwpxDocument | DocxDocument | str | Path | bytes",
    *,
    doc_type: DocType = "auto",
    include_tables: bool = True,
) -> StyleMap:
    """Extract styles for HWP/HWPX/DOCX with one interface."""
    resolved = infer_doc_type(source, doc_type)

    if resolved == "pdf":
        return extract_styles_pdf(source)

    if resolved == "hwp":
        if not isinstance(source, (str, Path)):
            raise TypeError("HWP conversion currently requires a filesystem path.")
        hwpx_bytes = convert_hwp_to_hwpx_bytes(source)
        return extract_styles_hwpx(hwpx_bytes)

    if resolved == "hwpx":
        return extract_styles_hwpx(source)

    return extract_styles_docx(source, include_tables=include_tables)


__all__ = [
    "DocType",
    "extract_styles",
    "extract_styles_docx",
    "extract_styles_hwpx",
    "extract_styles_pdf",
]
